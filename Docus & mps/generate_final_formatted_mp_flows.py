from pathlib import Path
from datetime import datetime
import importlib.util

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


HERE = Path(__file__).resolve().parent
BASE_PATH = HERE / "generate_detailed_mp_simulations.py"
spec = importlib.util.spec_from_file_location("mp_details", BASE_PATH)
base = importlib.util.module_from_spec(spec)
spec.loader.exec_module(base)  # type: ignore[union-attr]

ROOT = base.ROOT
OUT = base.OUT
REF = base.REF
FILES = base.FILES
LINES = base.LINES
PROGRAMS = base.PROGRAMS


def find_after(file_key, text, after_line):
    for line_no, line in enumerate(LINES[file_key], 1):
        if line_no >= after_line and text in line:
            return line_no
    return "?"


SOCKET = {
    "event": base.find("server", "@socketio.on('run_code')"),
    "handler": base.find("server", "def handle_run_code(data)"),
    "source": find_after("server", "source_code = data.get('source_code'", 740),
    "lex": find_after("server", "tokens, lex_errors = lex(source_code)", 740),
    "parse": find_after("server", "parse_result = parser.parse_and_build(tokens)", 740),
    "semantic": find_after("server", "semantic_result = validate_ast(parse_result['ast'], parse_result['symbol_table'])", 740),
    "session_emitter": find_after("server", "session_emitter = SessionEmitter", 740),
    "interpreter": find_after("server", "interp = Interpreter(socketio=session_emitter)", 740),
    "execute": find_after("server", "interp.interpret(ast)", 740),
}

REST = {
    "route": base.find("server", "@app.route('/api/run'"),
    "handler": base.find("server", "def run_endpoint()"),
    "source": find_after("server", "source_code = data['source_code']", 589),
    "lex": find_after("server", "tokens, lex_errors = lex(source_code)", 589),
    "parse": find_after("server", "parse_result = parser.parse_and_build(tokens)", 589),
    "semantic": find_after("server", "semantic_result = validate_ast(parse_result['ast'], parse_result['symbol_table'])", 589),
    "interpreter": find_after("server", "interp = Interpreter(socketio=collector)", 589),
    "execute": find_after("server", "interp.interpret(ast)", 589),
}

# The base script uses the first matching text. These final docs need the exact
# runtime route and exact modulo operator reference, so patch the references here.
REF["server_route"] = REST["route"]
REF["server_source"] = REST["source"]
REF["server_lex"] = REST["lex"]
REF["server_parse"] = REST["parse"]
REF["server_semantic"] = REST["semantic"]
REF["server_interpreter"] = REST["interpreter"]
REF["server_execute"] = REST["execute"]
REF["tok_mod"] = find_after("scanner", "tokens.append(Token(TT_MOD,", 1541)
REF["tok_mul"] = find_after("scanner", "tokens.append(Token(TT_MUL,", 1691)


def final_text(text):
    return (
        str(text)
        .replace("Multiplication is appended at scanner.py line 1690", f"Multiplication is appended at scanner.py line {REF['tok_mul']}")
        .replace("uses `%` at scanner.py line 1540", f"uses `%` at scanner.py line {REF['tok_mod']}")
    )


def esc(value):
    return str(value).replace("|", "\\|").replace("\n", "<br>")


def source_lines(source):
    return [(i, line) for i, line in enumerate(source.splitlines(), 1)]


def token_groups(tokens):
    groups = {}
    for token in tokens:
        if token.type in ("\n", "EOF"):
            continue
        groups.setdefault(token.line, []).append(f"{token.type}:{token.value}")
    return groups


def add_table_md(md, headers, rows):
    md.append("| " + " | ".join(headers) + " |")
    md.append("|" + "|".join("---" for _ in headers) + "|")
    for row in rows:
        md.append("| " + " | ".join(esc(cell) for cell in row) + " |")


def add_doc_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def code_para(doc, text, size=8):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    return paragraph


def exact_flow_rows():
    return [
        (
            "0",
            "Editor/UI sends code",
            f"UI run uses Socket.IO `run_code`: server.py lines {SOCKET['event']}-{SOCKET['handler']}. REST `/api/run` is server.py lines {REST['route']}-{REST['handler']}. Both use the same compiler stages.",
            "Full source code string",
        ),
        (
            "1",
            "Read source_code",
            f"Socket path stores `source_code = data.get(...)` at server.py line {SOCKET['source']}. REST path stores `source_code = data['source_code']` at line {REST['source']}.",
            "`source_code` variable",
        ),
        (
            "2",
            "Lexical analysis",
            f"`lex(source_code)` runs at socket line {SOCKET['lex']} or REST line {REST['lex']}. Lexer setup is scanner.py line {REF['lexer_init']}; `advance()` is line {REF['lexer_advance']}; main loop is line {REF['lexer_loop']}.",
            "Token list + lexical errors",
        ),
        (
            "3",
            "Syntax + AST build",
            f"`parser.parse_and_build(tokens)` runs at socket line {SOCKET['parse']} or REST line {REST['parse']}. Parser stack starts at parser.py line {REF['parser_stack']}; builder call happens at line {REF['parser_ast']}.",
            "AST + builder symbol table",
        ),
        (
            "4",
            "Semantic validation",
            f"`validate_ast(ast, symbol_table)` runs at socket line {SOCKET['semantic']} or REST line {REST['semantic']}. Analyzer walk starts at analyzer.py line {REF['semantic_walk']}.",
            "Validated AST or semantic errors",
        ),
        (
            "5",
            "Interpreter execution",
            f"Socket path creates `SessionEmitter` line {SOCKET['session_emitter']} and `Interpreter` line {SOCKET['interpreter']}; REST creates collector/interpreter at server.py line {REST['interpreter']}. Execution calls `interp.interpret(ast)` at socket line {SOCKET['execute']} or REST line {REST['execute']}.",
            "Runtime output from plant()",
        ),
    ]


def source_line_rows(program):
    rows = []
    for line_no, line in source_lines(program["source"]):
        stripped = line.strip()
        if not stripped:
            meaning = "Spacing only. Lexer may produce newline token; parser skips newlines."
        elif stripped.startswith("//"):
            meaning = "Comment. Lexer creates a comment token; parser skips it; builder removes it before AST."
        elif stripped.startswith("pollinate"):
            meaning = "Function declaration before root. Interpreter registers it first, then runs only when called."
        elif stripped.startswith("root"):
            meaning = "Main function. Interpreter automatically creates a call to root after registration."
        elif stripped.startswith("seed"):
            meaning = "Variable declaration. Builder creates a declaration AST node; interpreter creates runtime storage."
        elif stripped.startswith("cultivate"):
            meaning = "For-loop style statement: initializer, condition, update, and loop block."
        elif stripped.startswith("spring"):
            meaning = "If condition. Interpreter evaluates it first in a spring/bud/wither chain."
        elif stripped.startswith("bud"):
            meaning = "Else-if condition. Interpreter checks it only if spring was false."
        elif stripped.startswith("wither"):
            meaning = "Else branch. Interpreter runs it only when previous conditions are false."
        elif stripped.startswith("plant"):
            meaning = "Output statement. Interpreter evaluates arguments and emits text."
        elif stripped.startswith("reclaim"):
            meaning = "Function return/end statement. Interpreter exits the current function."
        elif "=" in stripped:
            meaning = "Assignment/expression statement that updates runtime variable state."
        else:
            meaning = "Block brace or structural line."
        rows.append((line_no, line, meaning))
    return rows


def stage_rows(program, data):
    return [
        ("Lexer", f"errors={data.get('lex_errors')}; token_count={data.get('token_count')}"),
        ("Parser", f"parse_ok={data.get('parse_ok')}; errors={data.get('parse_errors')}"),
        ("Builder", f"success={data.get('parse_build', {}).get('success')}; errors={data.get('parse_build', {}).get('errors')}"),
        ("Semantic", f"success={data.get('semantic', {}).get('success')}; errors={data.get('semantic', {}).get('errors')}; warnings={data.get('semantic', {}).get('warnings')}"),
        ("Interpreter", f"outputs={data.get('outputs')}; functions={data.get('functions')}; final_scopes={data.get('final_scopes')}"),
    ]


def flow_timeline(program, data):
    if program["slug"] == "MP1_TIMESTHREE":
        return [
            ("1", "Register functions", "`eval_program()` sees `timesThree` and `root`; both are stored by `eval_function_declaration()`.", "No output yet."),
            ("2", "Call root", "`eval_program()` creates `FunctionCallNode('root', [])`.", "Root starts executing."),
            ("3", "Declare variables", "`seed value = 0;` and `seed i;` execute.", "value=0, i=0."),
            ("4", "Start cultivate", "Initializer `i = 1`; condition `i <= 4` is true.", "Enter loop body."),
            ("5", "Iteration 1", "`timesThree(1)` returns 3; `3 % 2 == 0` is false.", "wither prints `Odd  product: 3`."),
            ("6", "Iteration 2", "`timesThree(2)` returns 6; condition is true.", "spring prints `Even product: 6`."),
            ("7", "Iteration 3", "`timesThree(3)` returns 9; condition is false.", "wither prints `Odd  product: 9`."),
            ("8", "Iteration 4", "`timesThree(4)` returns 12; condition is true.", "spring prints `Even product: 12`."),
            ("9", "Stop", "Update makes i=5; `5 <= 4` is false; `reclaim;` exits root.", "Execution complete."),
        ]
    if program["slug"] == "MP2_TEMPERATURE":
        return [
            ("1", "Register root", "`eval_program()` stores root then calls it.", "No output yet."),
            ("2", "Declare temperature", "`seed temperature = 50;` executes.", "temperature=50."),
            ("3", "Ignore comment", "Comment was tokenized by lexer but skipped by parser and filtered before builder.", "No runtime node."),
            ("4", "Evaluate spring", "`temperature >= 50` becomes `50 >= 50`.", "True."),
            ("5", "Run spring block", "Because spring is true, interpreter runs only the spring block.", "Prints boiling message."),
            ("6", "Skip bud/wither", "The branch chain is already satisfied.", "Execution continues to reclaim."),
            ("7", "End", "`reclaim;` exits root.", "Execution complete."),
        ]
    return [
        ("1", "Register root", "`eval_program()` stores root then calls it.", "No output yet."),
        ("2", "Declare count", "`seed count;` uses default seed value.", "count=0."),
        ("3", "Start cultivate", "Initializer `count = 1`; condition `count <= 5` is true.", "Enter loop body."),
        ("4", "Iteration 1", "`1 % 2 == 0` is false.", "wither prints `1`; count++ -> 2."),
        ("5", "Iteration 2", "`2 % 2 == 0` is true.", "spring prints `Bloom!`; count++ -> 3."),
        ("6", "Iteration 3", "`3 % 2 == 0` is false.", "wither prints `3`; count++ -> 4."),
        ("7", "Iteration 4", "`4 % 2 == 0` is true.", "spring prints `Bloom!`; count++ -> 5."),
        ("8", "Iteration 5", "`5 % 2 == 0` is false.", "wither prints `5`; count++ -> 6."),
        ("9", "Stop", "`6 <= 5` is false; `reclaim;` exits root.", "Execution complete."),
    ]


def parser_sequence(program):
    if program["slug"] == "MP1_TIMESTHREE":
        return [
            ("Lookahead `pollinate`", "`<program>` chooses the production with optional function definitions before root."),
            ("Function header", "Parser matches `pollinate seed id ( seed id ) {`."),
            ("Function body", "Parser checks `reclaim <expression>;`, where expression is `n * 3`."),
            ("Root", "Parser matches `root() {` then local declarations."),
            ("Cultivate", "Parser checks initializer `i = 1`, condition `i <= 4`, update `i++`, and body."),
            ("Spring/wither", "Parser checks `value % 2 == 0`, plant blocks, then final root `reclaim;`."),
        ]
    if program["slug"] == "MP2_TEMPERATURE":
        return [
            ("Lookahead `root`", "`<program>` skips optional globals/functions and matches root."),
            ("Declaration", "Parser checks `seed temperature = 50;`."),
            ("Comment skip", f"Parser skips comment/newline tokens at parser.py line {REF['parser_skip']}."),
            ("Spring", "Parser checks condition `temperature >= 50` and spring block."),
            ("Bud", "Parser checks optional bud condition `temperature <= 0` and bud block."),
            ("Wither", "Parser checks optional wither block and final `reclaim;`."),
        ]
    return [
        ("Lookahead `root`", "`<program>` matches root."),
        ("Declaration", "Parser checks `seed count;`."),
        ("Cultivate", "Parser checks `count = 1; count <= 5; count++`."),
        ("Spring/wither", "Parser checks `count % 2 == 0`, spring plant block, and wither plant block."),
        ("End", "Parser checks required final `reclaim;`."),
    ]


def write_markdown(path, program, data):
    group_rows = [(line_no, "`" + "  |  ".join(values) + "`") for line_no, values in sorted(token_groups(data["tokens"]).items())]
    token_rows = [(i, t.line, t.col, f"`{t.type}`", f"`{t.value}`") for i, t in enumerate(data["tokens"], 1)]
    output_rows = [(i, f"`{out}`") for i, out in enumerate(data.get("outputs", []), 1)]

    md = [
        f"# {program['title']} - Final Formatted System Flow Simulation",
        "",
        f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}",
        "",
        "## 1. Program Purpose",
        program["purpose"],
        "",
        "## 2. Source Program",
        "```gal",
        program["source"],
        "```",
        "",
        "## 3. Correct System Flow",
    ]
    add_table_md(md, ["Order", "Stage", "Exact Code Flow", "Passes To Next Stage"], exact_flow_rows())
    md += ["", "## 4. Actual Run Results"]
    add_table_md(md, ["Stage", "Result"], stage_rows(program, data))
    md += ["", "## 5. Source Line Meaning"]
    add_table_md(md, ["Line", "Code", "Meaning"], source_line_rows(program))
    md += ["", "## 6. Lexer Simulation"]
    for i, item in enumerate(program["lexer"], 1):
        md.append(f"{i}. {final_text(item)}")
    md += ["", "## 7. Tokens By Source Line"]
    add_table_md(md, ["Source Line", "Tokens Produced"], group_rows)
    md += ["", "## 8. Parser Simulation"]
    add_table_md(md, ["Parser Moment", "What Happens"], parser_sequence(program))
    md.append("")
    md.append(f"Parser mechanics: stack starts at parser.py line {REF['parser_stack']}; current stack top is read at line {REF['parser_top']}; current lookahead token is read at line {REF['parser_token']}; production is selected at line {REF['parser_prod']}; terminals match at line {REF['parser_match']}.")
    md += ["", "## 9. Builder / AST Simulation"]
    for i, item in enumerate(program["builder"], 1):
        md.append(f"{i}. {final_text(item)}")
    md += ["", "## 10. AST Shape Summary", "```text"]
    md.extend(data.get("ast_tree", []))
    md += ["```", "", "## 11. Semantic Simulation"]
    for i, item in enumerate(program["semantic"], 1):
        md.append(f"{i}. {final_text(item)}")
    md += ["", "## 12. Interpreter Simulation"]
    for i, item in enumerate(program["interpreter"], 1):
        md.append(f"{i}. {final_text(item)}")
    md += ["", "## 13. Runtime Timeline"]
    add_table_md(md, ["Step", "Runtime Moment", "What Executes", "Result"], flow_timeline(program, data))
    md += ["", "## 14. Final Output"]
    add_table_md(md, ["#", "Output"], output_rows)
    md += ["", "## 15. Full Token Table"]
    add_table_md(md, ["#", "Line", "Col", "Token Type", "Value"], token_rows)
    md += ["", "## 16. Code Reference Appendix"]
    for title, file_key, start, end, note in [
        ("Socket.IO run_code Flow", "server", SOCKET["event"], SOCKET["execute"] + 4, "This is the normal browser run path."),
        ("REST /api/run Flow", "server", REF["server_route"], REF["server_execute"] + 4, "Same stage order, used by REST run calls."),
        ("Lexer current_char Loop", "scanner", REF["lexer_init"], REF["lexer_loop"] + 4, "Character-by-character scan setup."),
        ("Parser Stack Algorithm", "parser", REF["parser_parse"], REF["parser_match"] + 5, "LL(1) stack and lookahead parser."),
        ("Builder Entry", "builder", REF["builder_build"], REF["builder_loop"] + 8, "Token-to-AST conversion starts here."),
        ("Semantic Walk", "semantic", REF["semantic_validate"], REF["semantic_handler"] + 8, "Final AST validation."),
        ("Interpreter Program Entry", "interpreter", REF["interpreter_interpret"], REF["interpreter_program"] + 18, "Runtime dispatcher and root call."),
    ]:
        md += ["", f"### {title}", f"Lines {start}-{end} in `{FILES[file_key].relative_to(ROOT)}`", "```python", base.snippet(file_key, start, end), "```", note]
    path.write_text("\n".join(md), encoding="utf-8")


def write_docx(path, program, data):
    if path.exists():
        try:
            path.unlink()
        except PermissionError:
            path = path.with_name(path.stem + f"_{datetime.now():%H%M%S}" + path.suffix)

    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    for style in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)

    heading = doc.add_heading(f"{program['title']} - Final Formatted System Flow Simulation", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.add_heading("1. Program Purpose", 1)
    doc.add_paragraph(program["purpose"])
    doc.add_heading("2. Source Program", 1)
    code_para(doc, program["source"])
    doc.add_heading("3. Correct System Flow", 1)
    add_doc_table(doc, ["Order", "Stage", "Exact Code Flow", "Passes To Next Stage"], exact_flow_rows())
    doc.add_heading("4. Actual Run Results", 1)
    add_doc_table(doc, ["Stage", "Result"], stage_rows(program, data))
    doc.add_heading("5. Source Line Meaning", 1)
    add_doc_table(doc, ["Line", "Code", "Meaning"], source_line_rows(program))
    doc.add_heading("6. Lexer Simulation", 1)
    for i, item in enumerate(program["lexer"], 1):
        doc.add_paragraph(f"{i}. {final_text(item)}")
    doc.add_heading("7. Tokens By Source Line", 1)
    add_doc_table(doc, ["Source Line", "Tokens Produced"], [(line_no, "  |  ".join(values)) for line_no, values in sorted(token_groups(data["tokens"]).items())])
    doc.add_heading("8. Parser Simulation", 1)
    add_doc_table(doc, ["Parser Moment", "What Happens"], parser_sequence(program))
    doc.add_paragraph(f"Parser mechanics: stack starts at parser.py line {REF['parser_stack']}; stack top is read at line {REF['parser_top']}; lookahead token is read at line {REF['parser_token']}; production is selected at line {REF['parser_prod']}; terminals match at line {REF['parser_match']}.")
    doc.add_heading("9. Builder / AST Simulation", 1)
    for i, item in enumerate(program["builder"], 1):
        doc.add_paragraph(f"{i}. {final_text(item)}")
    doc.add_heading("10. AST Shape Summary", 1)
    code_para(doc, "\n".join(data.get("ast_tree", [])))
    doc.add_heading("11. Semantic Simulation", 1)
    for i, item in enumerate(program["semantic"], 1):
        doc.add_paragraph(f"{i}. {final_text(item)}")
    doc.add_heading("12. Interpreter Simulation", 1)
    for i, item in enumerate(program["interpreter"], 1):
        doc.add_paragraph(f"{i}. {final_text(item)}")
    doc.add_heading("13. Runtime Timeline", 1)
    add_doc_table(doc, ["Step", "Runtime Moment", "What Executes", "Result"], flow_timeline(program, data))
    doc.add_heading("14. Final Output", 1)
    add_doc_table(doc, ["#", "Output"], [(i, out) for i, out in enumerate(data.get("outputs", []), 1)])
    doc.add_heading("15. Full Token Table", 1)
    add_doc_table(doc, ["#", "Line", "Col", "Token Type", "Value"], [(i, t.line, t.col, t.type, t.value) for i, t in enumerate(data["tokens"], 1)])
    doc.add_heading("16. Code Reference Appendix", 1)
    for title, file_key, start, end, note in [
        ("Socket.IO run_code Flow", "server", SOCKET["event"], SOCKET["execute"] + 4, "This is the normal browser run path."),
        ("REST /api/run Flow", "server", REF["server_route"], REF["server_execute"] + 4, "Same stage order, used by REST run calls."),
        ("Lexer current_char Loop", "scanner", REF["lexer_init"], REF["lexer_loop"] + 4, "Character-by-character scan setup."),
        ("Parser Stack Algorithm", "parser", REF["parser_parse"], REF["parser_match"] + 5, "LL(1) stack and lookahead parser."),
        ("Builder Entry", "builder", REF["builder_build"], REF["builder_loop"] + 8, "Token-to-AST conversion starts here."),
        ("Semantic Walk", "semantic", REF["semantic_validate"], REF["semantic_handler"] + 8, "Final AST validation."),
        ("Interpreter Program Entry", "interpreter", REF["interpreter_interpret"], REF["interpreter_program"] + 18, "Runtime dispatcher and root call."),
    ]:
        doc.add_heading(title, 2)
        doc.add_paragraph(f"Lines {start}-{end} in {FILES[file_key].relative_to(ROOT)}")
        code_para(doc, base.snippet(file_key, start, end), 7)
        doc.add_paragraph(note)
    doc.save(path)
    return path


def main():
    outputs = []
    for program in PROGRAMS:
        data = base.run_pipeline(program["source"])
        docx_path = OUT / f"{program['slug']}_FINAL_FORMATTED_SYSTEM_FLOW.docx"
        md_path = OUT / f"{program['slug']}_FINAL_FORMATTED_SYSTEM_FLOW.md"
        write_markdown(md_path, program, data)
        final_docx = write_docx(docx_path, program, data)
        outputs.append((program["title"], final_docx, md_path, data))
    for title, docx_path, md_path, data in outputs:
        print(title)
        print("  DOCX:", docx_path)
        print("  MD:", md_path)
        print("  tokens:", data.get("token_count"))
        print("  outputs:", data.get("outputs"))


if __name__ == "__main__":
    main()
