from pathlib import Path
from datetime import datetime
import re

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Docus & mps"
MD_PATH = OUT_DIR / "DETAILED_SYSTEM_CODE_EXPLANATION_GROWALANGUAGE_COMPILER.md"
DOCX_PATH = OUT_DIR / "DETAILED_SYSTEM_CODE_EXPLANATION_GROWALANGUAGE_COMPILER.docx"
PDF_PATH = OUT_DIR / "DETAILED_SYSTEM_CODE_EXPLANATION_GROWALANGUAGE_COMPILER.pdf"


def _read_lines(rel):
    path = ROOT / rel
    try:
        return path.read_text(encoding="utf-8").splitlines()
    except UnicodeDecodeError:
        return path.read_text(errors="replace").splitlines()


def snippet(rel, start, end, lang="python"):
    lines = _read_lines(rel)
    body = "\n".join(lines[start - 1:end])
    return (
        f"Source: `{rel}:{start}-{end}`\n\n"
        "Code:\n"
        f"```{lang}\n{body}\n```"
    )


def js_snippet(rel, start, end):
    return snippet(rel, start, end, lang="javascript")


SAMPLE_PROGRAM = """root() {
    seed x = 10;
    seed y = 5;
    seed sum;

    sum = x + y;

    plant("Sum: ", sum, "\\n");

    reclaim;
}"""


TOKEN_ROWS = [
    ("root", "root", "root", 1, 0),
    ("(", "(", "(", 1, 4),
    (")", ")", ")", 1, 5),
    ("{", "{", "{", 1, 7),
    ("\\n", "newline", "\\n", 1, 8),
    ("seed", "seed", "seed", 2, 4),
    ("x", "id", "x", 2, 9),
    ("=", "=", "=", 2, 11),
    ("10", "intlit", "10", 2, 13),
    (";", ";", ";", 2, 15),
    ("seed", "seed", "seed", 3, 4),
    ("y", "id", "y", 3, 9),
    ("=", "=", "=", 3, 11),
    ("5", "intlit", "5", 3, 13),
    (";", ";", ";", 3, 14),
    ("seed", "seed", "seed", 4, 4),
    ("sum", "id", "sum", 4, 9),
    (";", ";", ";", 4, 12),
    ("sum", "id", "sum", 6, 4),
    ("=", "=", "=", 6, 8),
    ("x", "id", "x", 6, 10),
    ("+", "+", "+", 6, 12),
    ("y", "id", "y", 6, 14),
    (";", ";", ";", 6, 15),
    ("plant", "plant", "plant", 8, 4),
    ("(", "(", "(", 8, 9),
    ('"Sum: "', "stringlit", '"Sum: "', 8, 10),
    (",", ",", ",", 8, 17),
    ("sum", "id", "sum", 8, 19),
    (",", ",", ",", 8, 22),
    ('"\\n"', "stringlit", '"\\n"', 8, 24),
    (")", ")", ")", 8, 28),
    (";", ";", ";", 8, 29),
    ("reclaim", "reclaim", "reclaim", 10, 4),
    (";", ";", ";", 10, 11),
    ("}", "}", "}", 11, 0),
    ("", "EOF", "", 11, 0),
]


def token_table():
    rows = ["| Lexeme | Token Type | Value | Line | Column |", "|---|---|---|---:|---:|"]
    for lex, typ, val, line, col in TOKEN_ROWS:
        rows.append(f"| `{lex}` | `{typ}` | `{val}` | {line} | {col} |")
    return "\n".join(rows)


def build_markdown():
    p = []
    p.append("# Detailed System Code Explanation of GrowALanguage Compiler")
    p.append(
        "Prepared for compiler defense. This document is based on the current source code in "
        f"`{ROOT}`.\n\nGenerated: "
        + datetime.now().strftime("%Y-%m-%d %H:%M")
    )

    p.append("""## 1. System Overview

The GrowALanguage compiler is a web-based compiler pipeline. The user writes GrowALanguage code in the editor, the frontend sends the source code to the Flask backend, and the backend processes it through lexical analysis, syntax analysis, semantic analysis, and runtime execution.

| Stage | Main Files | Main Classes / Functions | Main Responsibility |
|---|---|---|---|
| Code editor and frontend input | `UI/index.html`, `UI/main.js` | `runCode()`, `runProgram()`, `runViaREST()`, `runViaSocket()` | Reads the text from the Monaco editor and sends it to the backend API. |
| Server routing | `Backend/server.py` | `lexer_endpoint()`, `parser_endpoint()`, `semantic_endpoint()`, `run_endpoint()`, `handle_run_code()` | Receives source code and controls which compiler stages run. |
| Scanner / lexer | `Backend/lexer/scanner.py` | `Lexer`, `Lexer.make_tokens()`, `Lexer.advance()`, `lex()` | Reads source code character by character and produces tokens or lexical errors. |
| Tokens | `Backend/shared/tokens.py` | `Token`, token constants, `get_token_description()` | Defines token types and token object structure. |
| Delimiters | `Backend/lexer/delimiters.py` | `idf_delim`, `whlnum_delim`, `decim_delim`, `delim24`, `delim25`, etc. | Defines valid next characters after lexemes. |
| Parser / CFG | `Backend/parser/parser.py`, `Backend/cfg/grammar.py` | `LL1Parser`, `parse()`, `parse_and_build()`, `cfg`, `predict_sets` | Checks token order using LL(1) grammar, FIRST, FOLLOW, and PREDICT sets. |
| AST builder | `Backend/parser/builder.py` | `build_ast()`, `Parser` class methods | Converts tokens into AST nodes and performs many static checks. |
| Semantic analyzer | `Backend/semantic/analyzer.py`, `Backend/semantic/symbol_table.py` | `validate_ast()`, `ASTValidator`, `SymbolTable` | Checks meaning: declarations, scope, types, break/continue context, and related rules. |
| Runtime / interpreter | `Backend/interpreter/interpreter.py` | `Interpreter`, `interpret()`, `eval_*()` methods | Executes the AST, stores variables, runs loops/functions, and produces output. |
| Error handling | `Backend/lexer/errors.py`, `Backend/semantic/errors.py`, `Backend/interpreter/errors.py`, `Backend/parser/parser.py` | `LexicalError`, `SemanticError`, `InterpreterError`, syntax error dictionaries | Builds error messages returned to the UI. |

High-level flow:

```text
Editor source code
  -> Frontend JavaScript reads editor.getValue()
  -> Flask server receives source_code
  -> Lexer scans characters and returns tokens
  -> Parser checks tokens using CFG and PREDICT table
  -> AST builder builds syntax tree
  -> Semantic analyzer checks meaning and rules
  -> Interpreter executes AST
  -> Server returns output/errors to frontend
  -> UI displays output/errors
```""")

    p.append("""## 2. Input Flow from Editor

The system receives the whole source code string from the editor, but the lexer processes that string one character at a time. The frontend first stores the editor content in a JavaScript variable, sends it as JSON, and the Flask route stores it in a Python variable named `source_code`.""")
    p.append("### 2.1 Editor Reads the Source Code")
    p.append(js_snippet("UI/main.js", 1005, 1035))
    p.append("""Explanation:
- `editor.getValue()` reads all text currently written in the editor.
- The result is stored in the JavaScript variable `sourceCode`.
- The frontend first runs a silent lexer check.
- If the code contains `water(...)`, the frontend uses Socket.IO because input requires interactive communication.
- If the code has no `water(...)`, the frontend uses a normal REST request.""")

    p.append("### 2.2 REST Run Request")
    p.append(js_snippet("UI/main.js", 910, 916))
    p.append("""Explanation:
- `runViaREST(sourceCode)` sends the complete source string to `POST /api/run`.
- The JSON body uses the key `source_code`.
- The backend route that receives this is `run_endpoint()` in `Backend/server.py`.""")

    p.append("### 2.3 Backend Receives the Source Code")
    p.append(snippet("Backend/server.py", 323, 384))
    p.append("""Explanation:
- `run_endpoint()` receives JSON from the frontend using `request.get_json()`.
- `source_code = data['source_code']` stores the full program text.
- The server calls `lex(source_code)` first.
- If lexical errors exist, the route immediately returns a lexical error response.
- If lexing succeeds, the server calls `parser.parse_and_build(tokens)`.
- If parsing and AST building succeed, the server calls `validate_ast(ast)`.
- If semantic analysis succeeds, the server creates `Interpreter(socketio=collector)` and executes the AST.""")

    p.append("""## 3. Scanner / Lexer Detailed Explanation

The lexer is the first compiler stage. Its job is not to understand the whole program yet. Its job is to convert raw characters into meaningful tokens such as `seed`, `id`, `=`, `intlit`, `;`, `plant`, and `EOF`.""")
    p.append("### 3.1 Lexer Object Initialization and `current_char`")
    p.append(snippet("Backend/lexer/scanner.py", 18, 27))
    p.append("""Explanation:
- `self.source_code` holds the complete source code string.
- `self.pos` is a `Position` object that tracks index, line number, and column.
- `self.current_char` stores the character currently being scanned.
- `self.advance()` moves the lexer to the first character.

Example:

```gal
seed num = 10;
```

At the start:

| Variable | Meaning | Example Value |
|---|---|---|
| `source_code` | Full input text | `seed num = 10;` |
| `pos.index` | Current character index | `0` after first advance |
| `pos.ln` | Current line number | `1` |
| `pos.col` | Current column number | `0` |
| `current_char` | Character being checked | `s` |""")

    p.append("### 3.2 How `advance()` and Position Tracking Work")
    p.append(snippet("Backend/lexer/positions.py", 3, 21))
    p.append("""Explanation:
- `Position.advance()` increments the character index and column.
- If the previous character was newline, it increments the line number and resets column to 0.
- This is why lexical, syntax, semantic, and runtime errors can show line and column information.""")

    p.append("### 3.3 Main Tokenizing Loop")
    p.append(snippet("Backend/lexer/scanner.py", 29, 35))
    p.append("""Explanation:
- `make_tokens()` creates two lists: `tokens` and `errors`.
- `tokens` stores successful token objects.
- `errors` stores lexical errors.
- The loop continues while `self.current_char` is not `None`.
- `None` means the lexer reached the end of the source code.""")

    p.append("### 3.4 Reserved Words Are Checked First")
    p.append(snippet("Backend/lexer/scanner.py", 341, 379))
    p.append("""Explanation:
- When the lexer sees `r`, it tries to match `root`.
- If the next characters are exactly `o`, `o`, `t`, and the delimiter after `root` is valid, it appends `Token(TT_RW_ROOT, ident_str, line, pos.col)`.
- If the word is not exactly a reserved word, the lexer falls back to identifier scanning.

Example:

```gal
root() {
```

Token output:

| Lexeme | Token Type |
|---|---|
| `root` | `root` |
| `(` | `(` |
| `)` | `)` |
| `{` | `{` |""")

    p.append("### 3.5 How a Wrong Reserved Word Becomes an Identifier")
    p.append(snippet("Backend/lexer/scanner.py", 611, 638))
    p.append("""Explanation:
- If the user writes `roof()` instead of `root()`, the nested `root` check fails when the lexer expects `t` but sees `f`.
- There is no separate `false` branch needed. After the reserved-word checks fail, Python continues to the fallback identifier block.
- `ident_str` already contains the characters collected so far, for example `roo`.
- The fallback `while` loop collects the remaining alphanumeric characters.
- In `roof`, `self.current_char` is `f` when fallback begins, so the loop appends `f`.
- The final `ident_str` becomes `roof`.
- If the next character after `roof` is valid according to `idf_delim`, the lexer appends an identifier token.

Step-by-step:

| Step | `current_char` | `ident_str` | What Happens |
|---:|---|---|---|
| 1 | `r` | `r` | Lexer enters alphabet branch. |
| 2 | `o` | `ro` | Still trying to match `root`. |
| 3 | `o` | `roo` | Still trying to match `root`. |
| 4 | `f` | `roo` | Expected `t`, but got `f`; reserved-word path fails. |
| 5 | `f` | `roof` | Fallback identifier loop collects `f`. |
| 6 | `(` | `roof` | `(` is valid delimiter, so token becomes `id`. |""")

    p.append("### 3.6 Integer and Double Literals")
    p.append(snippet("Backend/lexer/scanner.py", 1171, 1236))
    p.append("""Explanation:
- When `current_char` is a digit, the lexer collects digits into `digits`.
- If it sees `.`, it switches to decimal scanning and counts decimal points.
- If the number has no decimal point, it creates an `intlit` token.
- If it has one decimal point with digits after it, it creates a `dblit` token.
- If it sees invalid characters after a number, it creates a lexical error.

Example:

```gal
seed x = 10;
tree y = 2.5;
```

Token output:

| Lexeme | Token Type |
|---|---|
| `10` | `intlit` |
| `2.5` | `dblit` |""")

    p.append("### 3.7 String Literals")
    p.append(snippet("Backend/lexer/scanner.py", 1334, 1388))
    p.append("""Explanation:
- A string starts with double quote `"`.
- The lexer keeps reading characters until it finds the closing double quote.
- Escape sequences such as `\\n` are handled by checking `escaped`.
- If the closing quote is missing, the lexer creates a lexical error.
- If the character after the string is not a valid delimiter, it also creates a lexical error.
- The CFG does not contain separate quotation mark terminals because the lexer turns the whole quoted text into one `stringlit` token.""")

    p.append("### 3.8 Character Literals")
    p.append(snippet("Backend/lexer/scanner.py", 1390, 1456))
    p.append("""Explanation:
- A character literal starts with single quote `'`.
- It must contain exactly one character, or one supported escaped character.
- If it has more than one character, the lexer creates a lexical error.
- If the closing quote is missing, the lexer creates a lexical error.""")

    p.append("### 3.9 Comments")
    p.append(snippet("Backend/lexer/scanner.py", 1071, 1106))
    p.append("""Explanation:
- `//` starts a single-line comment.
- `/* ... */` starts a multi-line comment.
- The lexer tokenizes comments as `comment` or `mcommentlit`.
- The LL(1) parser is configured to skip comment tokens during syntax checking.

Implementation note: syntax parsing skips comment tokens. The AST-building phase filters newline tokens, so comments should be kept away from places where the builder is strict unless the builder also handles them.""")

    p.append("### 3.10 Operators and Reserved Symbols")
    p.append(snippet("Backend/lexer/scanner.py", 811, 847))
    p.append("""Explanation:
- This block recognizes `*`, `*=`, `**`, and `**=`.
- The scanner looks ahead by calling `advance()` after seeing `*`.
- If the next character is another `*`, it may become exponent `**` or exponent assignment `**=`.
- If the next character is `=`, it becomes multiplication assignment `*=`.

More operator groups:

| Operator Group | File Logic |
|---|---|
| plus, increment, plus-equals | `Backend/lexer/scanner.py:921-963` |
| minus, decrement, minus-equals | `Backend/lexer/scanner.py:645-672` |
| equals and equality | `Backend/lexer/scanner.py:987-1013` |
| logical not and not-equal | `Backend/lexer/scanner.py:729-750` |
| modulo and modulo-equals | `Backend/lexer/scanner.py:752-770` |
| slash, slash-equals, comments | `Backend/lexer/scanner.py:1071-1106` |
| grouping and braces | `Backend/lexer/scanner.py:789-917` |
| backtick concatenation | `Backend/lexer/scanner.py:1458-1467` |""")

    p.append("### 3.11 Invalid Delimiters")
    p.append(snippet("Backend/lexer/scanner.py", 631, 637))
    p.append("""Explanation:
- If `self.current_char` is `None` or is inside `idf_delim`, the identifier is valid.
- If the next character is not inside `idf_delim`, a lexical error is added.

Example:

```gal
root(){ seed x@; reclaim; }
```

Error:

```text
LEXICAL error line 1 col 13 Invalid delimiter '@' after 'x'
```""")

    p.append("### 3.12 End of File and Lexer Return Value")
    p.append(snippet("Backend/lexer/scanner.py", 1502, 1529))
    p.append("""Explanation:
- At the end of scanning, the lexer appends an EOF token.
- `lex(source_code)` is the public function used by the server.
- It creates a `Lexer`, calls `make_tokens()`, converts lexical errors into strings, and returns `(tokens, errors)`.""")

    p.append("""## 4. Tokens Explanation

Tokens are defined in `Backend/shared/tokens.py`. A token is a small object that stores what lexeme was found, what kind of token it is, and where it appeared.""")
    p.append("### 4.1 Token Class")
    p.append(snippet("Backend/shared/tokens.py", 86, 92))
    p.append("""Explanation:
- `type` stores the token category, such as `seed`, `id`, `intlit`, or `;`.
- `value` stores the original lexeme text.
- `line` stores the line number.
- `col` stores the column number.""")
    p.append("### 4.2 Major Token Groups")
    p.append(snippet("Backend/shared/tokens.py", 3, 80))
    p.append("""| Token Group | Examples | Meaning |
|---|---|---|
| Reserved words | `seed`, `tree`, `root`, `plant`, `water`, `reclaim` | Language keywords. |
| Identifiers | `x`, `sum`, `studentName` | User-defined names. |
| Literals | `intlit`, `dblit`, `stringlit`, `chrlit`, `sunshine`, `frost` | Constant values. |
| Operators | `=`, `+=`, `+`, `**`, `==`, `&&`, `!` | Computation and logic symbols. |
| Punctuation | `;`, `,`, `(`, `)`, `[`, `]`, `{`, `}` | Structure and grouping symbols. |
| EOF | `EOF` | Marks the end of the token list. |

Example source:

```gal
seed x = 10;
```

Token sequence:

| Lexeme | Token Type | Value |
|---|---|---|
| `seed` | `seed` | `seed` |
| `x` | `id` | `x` |
| `=` | `=` | `=` |
| `10` | `intlit` | `10` |
| `;` | `;` | `;` |
| end | `EOF` | empty |""")

    p.append("""## 5. Delimiters Explanation

A delimiter set is not the same as a token. A reserved symbol token is an actual token that goes to the parser. A delimiter set is only used by the lexer to check if the character after a lexeme is valid.""")
    p.append("### 5.1 Delimiter Definitions")
    p.append(snippet("Backend/lexer/delimiters.py", 1, 60))
    p.append("""| Delimiter Set | Contains / Purpose | Example Valid | Example Invalid |
|---|---|---|---|
| `idf_delim` | Valid characters after identifiers. Includes spaces, semicolon, operators, brackets, quotes, and similar boundaries. | `seed x;` | `seed x@;` |
| `whlnum_delim` | Valid characters after integer literals. | `seed x = 10;` | `seed x = 10a;` |
| `decim_delim` | Valid characters after double literals. | `tree y = 2.5;` | `tree y = 2.5a;` |
| `comment_delim` | Characters allowed while scanning comments. | `// hello` | depends on comment form |
| `delim24` | Used for binary/assignment operators where another operand may follow. | `x + y` | operator followed by invalid punctuation |
| `delim25` | Used after increment/decrement style operators and some arithmetic operators. | `x++;` | `x++@` |
| `delim26` | Used after logical not `!`, allowing identifiers, literals, unary signs, quotes, and whitespace. | `!ready` | `!@` |

Difference:
- Reserved symbol token: a real token sent to the parser, such as `++`.
- Delimiter set: a validation set used only by the lexer.
- Valid next character: the character that is allowed to appear immediately after a lexeme.""")

    p.append("""## 6. Parser Explanation

The parser checks if the token sequence follows the CFG. Your parser is an LL(1) stack-based parser.""")
    p.append("### 6.1 Parser Setup in the Server")
    p.append(snippet("Backend/server.py", 50, 57))
    p.append("### 6.2 Parser Class Initialization")
    p.append(snippet("Backend/parser/parser.py", 37, 64))
    p.append("""Explanation:
- `self.cfg` stores all grammar productions.
- `self.predict_sets` stores the terminals that choose each production.
- `self.start_symbol` is `<program>`.
- `self.end_marker` is `EOF`.
- `self.parsing_table` is built from the PREDICT sets.""")
    p.append("### 6.3 Parsing Table Construction")
    p.append(snippet("Backend/parser/parser.py", 66, 82))
    p.append("### 6.4 Main Stack-Based Parsing Algorithm")
    p.append(snippet("Backend/parser/parser.py", 614, 715))
    p.append("""Explanation:
- The stack starts as `[EOF, <program>]`.
- The parser looks at the top of the stack and the current token.
- If the top is a nonterminal, it uses `self.parsing_table[top][lookahead]` to choose a production.
- It pops the nonterminal and pushes the chosen production in reverse order.
- If the top is a terminal and it matches the current token, the parser consumes the token.
- If no production matches or the terminal does not match, it returns a syntax error.""")
    p.append("### 6.5 AST Build After Syntax Success")
    p.append(snippet("Backend/parser/parser.py", 1183, 1240))
    p.append("""Explanation:
- `parse_and_build(tokens)` first calls `self.parse(tokens)`.
- If syntax fails, it returns syntax errors.
- If syntax succeeds, it filters newline tokens and calls `_build_ast(filtered)` from `Backend/parser/builder.py`.
- The builder returns the AST and symbol table.""")
    p.append("""### 6.6 Sample Parser Flow

Sample code:

```gal
root() {
    seed x = 10;
    reclaim;
}
```

Important CFG path:

```text
<program>
  -> <global_declaration> <function_definition> root ( ) { <local_declaration> <body_statement> reclaim ; }
<global_declaration>
  -> lambda
<function_definition>
  -> lambda
<local_declaration>
  -> <var_dec> ; <local_declaration>
<var_dec>
  -> <data_type> id <var_tail>
<data_type>
  -> seed
<body_statement>
  -> lambda
```

The current aligned grammar requires local declarations before body statements, similar to old-style C structure.""")

    p.append("""## 7. CFG, FIRST, FOLLOW, and PREDICT Explanation

CFG means Context-Free Grammar. In your system it is stored as a Python dictionary named `cfg` in `Backend/cfg/grammar.py`. The keys are nonterminals such as `<program>` and `<expression>`. The values are possible productions.""")
    p.append("### 7.1 Program and Declaration Productions")
    p.append(snippet("Backend/cfg/grammar.py", 122, 155))
    p.append("### 7.2 FIRST, FOLLOW, and PREDICT Computation")
    p.append(snippet("Backend/cfg/grammar.py", 13, 119))
    p.append("""| Set | Purpose |
|---|---|
| FIRST | Tells which token can start a nonterminal or production. |
| FOLLOW | Tells which token can legally appear after a nonterminal. |
| PREDICT | Tells the LL(1) parser which production to use for a lookahead token. |

Example for `plant("hello");`:

```text
<io_stmt> -> plant ( <arguments> ) ;
<arguments> -> <expression> <arg_next>
<expression> -> <assignment_expression>
<assignment_expression> -> <logic_or>
<factor> -> stringlit
```

Important: `"hello"` is handled by the lexer as one `stringlit` token. The CFG does not need separate quote terminals.""")
    p.append("### 7.3 Expression Grammar")
    p.append(snippet("Backend/cfg/grammar.py", 518, 631))
    p.append("""Expression order:

```text
assignment expression
  -> logical OR
  -> logical AND
  -> equality / relational
  -> arithmetic addition and subtraction
  -> multiplication / division / modulo
  -> exponent
  -> unary or factor
```

This gives operator precedence. For example, `x + y * 2` parses multiplication before addition.""")

    p.append("""## 8. Semantic Analyzer Explanation

Semantic analysis checks meaning, not just token order. In this project, semantic checks are distributed across the AST builder, symbol table, and semantic validator.""")
    p.append("### 8.1 Symbol Table")
    p.append(snippet("Backend/semantic/symbol_table.py", 3, 58))
    p.append("""Explanation:
- `variables` stores global variables.
- `scopes` stores local scopes.
- `functions` stores function declarations.
- `bundle_types` stores struct-like bundle definitions.
- `declare_variable()` detects duplicate variables.
- `lookup_variable()` detects undeclared variables.""")
    p.append("### 8.2 Variable Declaration and Duplicate Identifier Checking")
    p.append(snippet("Backend/parser/builder.py", 310, 370))
    p.append("""Valid:

```gal
root(){
    seed x = 10;
    reclaim;
}
```

Invalid:

```gal
root(){
    seed x = 10;
    seed x = 20;
    reclaim;
}
```

Logic:
- `parse_variable()` collects the data type and identifier.
- It checks whether the name conflicts with a function or existing variable through the symbol table.
- If duplicate, it raises `SemanticError`.""")
    p.append("### 8.3 Array Size Checking")
    p.append(snippet("Backend/parser/builder.py", 373, 385))
    p.append("""Valid: `seed arr[3];`

Invalid: `seed arr[2.5];`

Logic:
- The builder expects an integer literal between brackets.
- If the dimension token is `dblit`, it raises a semantic error.
- This aligns the rule that array size must be an integer.""")
    p.append("### 8.4 Type Compatibility")
    p.append(snippet("Backend/parser/builder.py", 1006, 1011))
    p.append("""Explanation:
- Exact type matches are allowed.
- `seed` and `tree` are treated as compatible numeric types in this helper.
- Other mismatches produce semantic errors.

Invalid example:

```gal
root(){ seed x = "hello"; reclaim; }
```

Error:

```text
SEMANTIC error line 1 col 17 Type mismatch: cannot assign string value '"hello"' to integer (seed) variable
```""")
    p.append("### 8.5 Assignment and Fertile Constant Reassignment")
    p.append(snippet("Backend/parser/builder.py", 1930, 1958))
    p.append("""Logic:
- The builder looks up the variable.
- If the variable is marked `is_fertile`, reassignment is rejected.
- It also checks that the assigned expression type is compatible with the variable type.""")
    p.append("### 8.6 Function Call Parameter Checking")
    p.append(snippet("Backend/parser/builder.py", 1961, 2021))
    p.append("""Logic:
- The builder verifies that the called function exists.
- It checks argument count.
- It checks each argument type against the declared parameter type.
- It also handles array/list parameter compatibility.""")
    p.append("### 8.7 Water Input Checking")
    p.append(snippet("Backend/parser/builder.py", 2024, 2112))
    p.append("""Logic:
- `water()` must target a declared variable or valid assignable location.
- The builder checks if the target variable exists.
- It rejects invalid targets based on the grammar and symbol table.""")
    p.append("### 8.8 Reclaim / Return Checking")
    p.append(snippet("Backend/parser/builder.py", 2570, 2612))
    p.append("""Logic:
- `reclaim` is the return statement.
- The builder checks if the return value matches the current function return type.
- In `root()`, the aligned grammar requires final `reclaim;`.""")
    p.append("### 8.9 Struct / Bundle Checking")
    p.append(snippet("Backend/parser/builder.py", 110, 145))
    p.append("""Logic:
- Bundle declarations are stored in `symbol_table.bundle_types`.
- Duplicate member names are rejected.
- Bundle instance declarations and member access are later checked through builder methods and interpreter member access logic.""")
    p.append("### 8.10 Scope Checking")
    p.append(snippet("Backend/semantic/symbol_table.py", 79, 85))
    p.append("### 8.11 Extra Semantic Validator Checks")
    p.append(snippet("Backend/semantic/analyzer.py", 4, 35))
    p.append(snippet("Backend/semantic/analyzer.py", 133, 141))
    p.append("""Explanation:
- After the AST is built, `validate_ast(ast)` walks the AST.
- It checks rules such as `prune` only inside loops or harvest/switch, and `skip` only inside loops.""")

    p.append("""## 9. Runtime / Execution Explanation

After parsing and semantic analysis succeed, the system interprets the AST.""")
    p.append("### 9.1 Interpreter Initialization")
    p.append(snippet("Backend/interpreter/interpreter.py", 25, 48))
    p.append("### 9.2 Central Dispatch")
    p.append(snippet("Backend/interpreter/interpreter.py", 121, 212))
    p.append("""Explanation:
- `interpret(node)` receives an AST node.
- It checks the node class or `node_type`.
- It calls the matching `eval_*` method.
- Example: `PrintNode` calls `eval_print()`, `AssignmentNode` calls `eval_assignment()`.""")
    p.append("### 9.3 Program Execution Starts at `root`")
    p.append(snippet("Backend/interpreter/interpreter.py", 214, 219))
    p.append("""Explanation:
- The interpreter first registers top-level declarations and functions.
- Then it automatically calls `root()`.
- This makes `root()` the required main function.""")
    p.append("### 9.4 Runtime Variable Storage")
    p.append(snippet("Backend/interpreter/interpreter.py", 50, 92))
    p.append("### 9.5 Variable Declaration Runtime Logic")
    p.append(snippet("Backend/interpreter/interpreter.py", 222, 298))
    p.append("### 9.6 Expression Evaluation")
    p.append(snippet("Backend/interpreter/interpreter.py", 516, 683))
    p.append("### 9.7 `plant()` Output")
    p.append(snippet("Backend/interpreter/interpreter.py", 750, 804))
    p.append("""Explanation:
- `eval_print()` evaluates each argument.
- It joins multiple arguments with spaces.
- It appends or emits output through Socket.IO or REST collector.

Example:

```gal
plant("Sum: ", sum, "\\n");
```

Runtime output in the current system:

```text
Sum:  15 \\n
```

The extra spacing happens because multiple `plant` arguments are joined with spaces.""")
    p.append("### 9.8 `water()` Input")
    p.append(snippet("Backend/interpreter/interpreter.py", 1375, 1468))
    p.append("""Explanation:
- `eval_input()` determines which variable receives input.
- It emits an input request to the frontend.
- It waits for the frontend response.
- It converts the input based on the target variable type.
- Invalid input raises `InterpreterError`.""")
    p.append("### 9.9 Loops and Conditionals")
    p.append(snippet("Backend/interpreter/interpreter.py", 1101, 1142))
    p.append(snippet("Backend/interpreter/interpreter.py", 1144, 1192))
    p.append("### 9.10 `reclaim` Stops Function Execution")
    p.append(snippet("Backend/interpreter/interpreter.py", 857, 899))
    p.append("""Explanation:
- `eval_return()` raises `ReturnValue` as a control signal.
- `eval_function_call()` catches `ReturnValue` and returns its value to the caller.
- This is how `reclaim` stops a function body early.""")

    p.append("""## 10. Error Handling Explanation

Errors are detected at different compiler stages. Each stage stops the pipeline if it fails.

| Error Type | Detected In | Error Class / Format | Example |
|---|---|---|---|
| Lexical error | `Backend/lexer/scanner.py` | `LexicalError` | Invalid delimiter, invalid identifier, unclosed string. |
| Syntax error | `Backend/parser/parser.py` | dictionary with `type: SYNTAX` | Missing semicolon, unexpected token. |
| Semantic error | `Backend/parser/builder.py`, `Backend/semantic/analyzer.py` | `SemanticError` | Undeclared variable, type mismatch, duplicate declaration. |
| Runtime error | `Backend/interpreter/interpreter.py` | `InterpreterError` | Array out of bounds, division by zero during execution. |""")
    p.append("### 10.1 Lexical Error Class")
    p.append(snippet("Backend/lexer/errors.py", 3, 11))
    p.append("""Example invalid identifier:

```gal
root(){ seed 1x; reclaim; }
```

Error:

```text
LEXICAL error line 1 col 13 Identifiers cannot start with a number: '1x...'
```

Example invalid delimiter:

```gal
root(){ seed x@; reclaim; }
```

Error:

```text
LEXICAL error line 1 col 13 Invalid delimiter '@' after 'x'
```""")
    p.append("### 10.2 Syntax Error")
    p.append("""Detected in `LL1Parser.parse()` when the current token does not match the expected terminal or no CFG production matches the lookahead.

Example:

```gal
root(){ seed x = 10 reclaim; }
```

Error:

```text
SYNTAX error line 1 col 20 Unexpected token 'reclaim'. Expected: semicolon or expression continuation tokens
```""")
    p.append("### 10.3 Semantic Error Class")
    p.append(snippet("Backend/semantic/errors.py", 6, 12))
    p.append("""Example undeclared variable:

```gal
root(){ x = 1; reclaim; }
```

Error:

```text
SEMANTIC error line 1: Variable 'x' used before declaration.
```

Example type mismatch:

```gal
root(){ seed x = "hello"; reclaim; }
```

Error:

```text
SEMANTIC error line 1 col 17 Type mismatch: cannot assign string value '"hello"' to integer (seed) variable
```""")
    p.append("### 10.4 Runtime Error Class")
    p.append(snippet("Backend/interpreter/errors.py", 6, 27))
    p.append("""Example array out of bounds:

```gal
root(){ seed a[2] = {1,2}; plant(a[3]); reclaim; }
```

Error:

```text
RUNTIME error line 1: Index '3' out of bounds for 'a'.
```""")
    p.append("### 10.5 Sending Errors Back to Frontend")
    p.append("""Errors are returned by Flask routes as JSON. Example:

```json
{
  "stage": "semantic",
  "error": "SEMANTIC error line 1: Variable 'x' used before declaration.",
  "output": ""
}
```

The frontend displays the returned `error` or `output` inside the output panel.""")

    p.append("## 11. Full Example Walkthrough")
    p.append("Sample code:\n\n```gal\n" + SAMPLE_PROGRAM + "\n```")
    p.append("""### A. Source Code Input

The user writes the code in the Monaco editor. The frontend stores it in `sourceCode` through `editor.getValue()` and sends it to `/api/run` as `source_code`.

### B. Lexer Scanning

The lexer receives the whole source string, but scans it character by character using `current_char` and `advance()`.

### C. Token List Generated

This token table was generated from the current system pipeline.""")
    p.append(token_table())
    p.append("""### D. Parser Checking With CFG

Important derivation:

```text
<program>
  -> <global_declaration> <function_definition> root ( ) { <local_declaration> <body_statement> reclaim ; }
<local_declaration>
  -> <var_dec> ; <local_declaration>
  -> seed id <var_tail> ; <local_declaration>
  -> seed id = <init_val> ; <local_declaration>
<body_statement>
  -> <non_reclaim_stmt> <body_statement>
  -> <id_stmt> <body_statement>
  -> id <id_next> <id_stmt_tail> <body_statement>
  -> id = <expression> ; <body_statement>
  -> <non_reclaim_stmt> <body_statement>
  -> <io_stmt> <body_statement>
  -> plant ( <arguments> ) ; <body_statement>
<body_statement>
  -> lambda
```

### E. Semantic Checking

| Check | Result |
|---|---|
| `x`, `y`, and `sum` are declared before use | Passed |
| `x` and `y` are `seed` integers | Passed |
| `sum = x + y` assigns numeric result to `seed` | Passed |
| `plant("Sum: ", sum, "\\n")` uses valid output arguments | Passed |
| `reclaim;` exists at the end of root | Passed |

### F. Runtime Execution

| Step | Action | Runtime Effect |
|---:|---|---|
| 1 | Enter `root()` | New function scope is created. |
| 2 | `seed x = 10;` | Store `x` as seed value `10`. |
| 3 | `seed y = 5;` | Store `y` as seed value `5`. |
| 4 | `seed sum;` | Store `sum` as seed default value `0`. |
| 5 | `sum = x + y;` | Evaluate `10 + 5`, store `sum = 15`. |
| 6 | `plant("Sum: ", sum, "\\n");` | Output text and value. |
| 7 | `reclaim;` | Stop root function execution. |

Runtime variable table while inside `root()`:

| Variable | Type | Value |
|---|---|---:|
| `x` | `seed` | `10` |
| `y` | `seed` | `5` |
| `sum` | `seed` | `15` |

### G. Final Output

Actual current-system output:

```text
Sum:  15 \\n
```""")

    p.append("""## 12. Code-Level Explanation Format

For defense, use this format whenever explaining an important code block:

```text
Code:
[paste the exact code snippet]

Explanation:
- What this code does
- When it runs
- What input it receives
- What output it returns
- What part of the compiler depends on it
```""")
    p.append(snippet("Backend/lexer/scanner.py", 1513, 1529))
    p.append("""Explanation:
- This code is the public lexer entry point.
- It runs whenever the server calls `lex(source_code)`.
- It receives the full source code string.
- It returns the generated token list and lexical error list.
- The parser depends on the tokens returned by this function.""")

    p.append("""## 13. Defense Script Section

Use this short Taglish explanation during defense:

```text
Sa system namin, unang nangyayari is kinukuha muna ng frontend yung code na nilagay ng user sa editor gamit yung editor.getValue(). Then ipinapasa siya sa backend Flask server as source_code.

Pagdating sa backend, unang stage is lexer. Yung lexer hindi niya agad iniintindi yung buong program as meaning. Binabasa niya muna yung source code character by character gamit yung current_char and advance(). Habang nagbabasa siya, kino-convert niya yung raw text into tokens, like seed, id, intlit, equals, semicolon, plant, and EOF.

After lexing, yung token list pinapasa sa parser. Yung parser namin is LL(1), so gumagamit siya ng CFG, FIRST, FOLLOW, and PREDICT sets. May stack siya, then tinitingnan niya yung current token or lookahead para malaman kung anong production rule ang gagamitin. Kapag hindi tugma yung token sa grammar, syntax error ang lalabas.

Kapag pasado sa syntax, bubuo naman yung builder ng AST. Then semantic checking happens. Dito chine-check kung declared ba yung variables, tama ba yung type, may duplicate ba, valid ba yung array size, tama ba yung function call arguments, and valid ba yung reclaim.

Kapag pasado lahat, interpreter naman yung tatakbo. Yung interpreter ang nag-e-execute ng AST. Siya yung nag-i-store ng variables sa scopes, nag-e-evaluate ng expressions, nagpapatakbo ng loops and conditionals, nagha-handle ng plant output, and nagre-request ng input kapag may water().

Kung may error sa kahit anong stage, titigil yung pipeline doon. For example lexical error kapag invalid character, syntax error kapag mali grammar, semantic error kapag undeclared variable or type mismatch, and runtime error kapag execution problem like array out of bounds. Then yung error message ibabalik ng server sa frontend para makita ng user sa output panel.
```""")

    p.append("""## 14. Formatting Requirements Checklist

This document is prepared as a PDF-ready technical explanation. It includes:

- Clear numbered sections.
- Tables for pipeline, tokens, delimiters, semantic checks, and runtime variables.
- Code snippets from the actual Python and JavaScript files.
- Exact file names, class names, function names, and variable names.
- A full example walkthrough from editor input to final output.
- A defense script in simple Taglish.""")

    p.append("""## Appendix C. Deep Code Trace With Exact Line References

This appendix is added for defense questions that ask, "What exact code runs next?"

### C.1 Server Receives Code Before Lexing

Exact flow:

| Step | File and Lines | Code Object | What Happens |
|---:|---|---|---|
| 1 | `UI/main.js:1005-1035` | `runProgram()` | Gets the full editor text using `editor.getValue()` and stores it in `sourceCode`. |
| 2 | `UI/main.js:910-916` | `runViaREST(sourceCode)` | Sends JSON body `{ source_code: sourceCode }` to `/api/run`. |
| 3 | `Backend/server.py:323-384` | `run_endpoint()` | Reads `data = request.get_json()` then stores `source_code = data['source_code']`. |
| 4 | `Backend/server.py:340` | `lex(source_code)` | Passes the full source string to the lexer. |

Key point: the server passes the whole source code string, but the lexer reads it one character at a time.

### C.2 Lexer Reserved Word Process

Example:

```gal
root()
```

Exact flow:

| Step | File and Lines | Code Object | What Happens |
|---:|---|---|---|
| 1 | `Backend/lexer/scanner.py:18-27` | `Lexer.__init__()` and `advance()` | Stores the source in `self.source_code`, creates `self.pos`, sets `self.current_char`, then moves to the first character. |
| 2 | `Backend/lexer/scanner.py:29-35` | `make_tokens()` | Starts the main `while self.current_char != None` loop. |
| 3 | `Backend/lexer/scanner.py:37-40` | alphabet branch | Because `current_char` is `r`, the lexer enters the alphabet/reserved-word path and starts `ident_str`. |
| 4 | `Backend/lexer/scanner.py:341-379` | `root` nested check | Checks `r -> o -> o -> t`, then checks if the next character is valid for `root`, usually `(`. |
| 5 | `Backend/lexer/scanner.py:375` | `tokens.append(...)` | Appends `Token(TT_RW_ROOT, ident_str, line, pos.col)`. |

Important variables:

| Variable | Meaning |
|---|---|
| `self.current_char` | The character currently being checked. Example: `r`, then `o`, then `o`, then `t`. |
| `ident_str` | The word being built. Example: `r`, `ro`, `roo`, `root`. |
| `pos` | Copy of the starting position of the lexeme, used for token line/column. |
| `tokens` | List where successful tokens are stored. |

### C.3 Lexer Fallback From Reserved Word to Identifier

Example:

```gal
roof()
```

Exact flow:

| Step | File and Lines | Code Object | What Happens |
|---:|---|---|---|
| 1 | `Backend/lexer/scanner.py:341-367` | `root` attempt | The lexer collects `r`, `o`, `o`. |
| 2 | `Backend/lexer/scanner.py:368` | `if self.current_char == "t":` | This condition is false because the current character is `f`. |
| 3 | `Backend/lexer/scanner.py:611-638` | identifier fallback | Since the reserved-word path did not append a token, the code continues to the fallback identifier block. |
| 4 | `Backend/lexer/scanner.py:614-615` | `while self.current_char is not None and self.current_char in ALPHANUM:` | The loop collects the remaining character `f`, so `ident_str` becomes `roof`. |
| 5 | `Backend/lexer/scanner.py:631-632` | delimiter check and token append | If the next character is valid, such as `(`, the lexer appends `Token(TT_IDENTIFIER, ident_str, line, pos.col)`. |

Why there is no `else false` code:

```text
Python does not need a special false program here.
If the nested reserved-word condition fails and no token is appended,
execution naturally continues to the fallback identifier code below it.
```

### C.4 Lexer Passes Tokens Back to the Server

Exact flow:

| Step | File and Lines | Code Object | What Happens |
|---:|---|---|---|
| 1 | `Backend/lexer/scanner.py:1502-1505` | EOF append and return | Adds `Token(TT_EOF, "", line, pos.col)` and returns `tokens, errors`. |
| 2 | `Backend/lexer/scanner.py:1513-1529` | `lex(source_code)` | Creates `Lexer(source_code)`, calls `make_tokens()`, converts lexical errors to strings, and returns tokens/errors. |
| 3 | `Backend/server.py:340-344` | `run_endpoint()` | If `lex_errors` exists, returns a lexical error immediately. Otherwise proceeds to parsing. |

### C.5 Parser Detailed Process After Lexer

Exact flow:

| Step | File and Lines | Code Object | What Happens |
|---:|---|---|---|
| 1 | `Backend/server.py:348` | `parser.parse_and_build(tokens)` | Sends the token list to the parser. |
| 2 | `Backend/parser/parser.py:1183-1191` | `parse_and_build()` | Calls `self.parse(tokens)` first for syntax checking. |
| 3 | `Backend/parser/parser.py:614-626` | `parse()` setup | Converts tokens into `_TokView`, creates stack `[EOF, <program>]`, and starts reading lookahead tokens. |
| 4 | `Backend/parser/parser.py:665-698` | nonterminal expansion | If stack top is a nonterminal, uses `self.parsing_table[top][token_type]` to choose a production. |
| 5 | `Backend/parser/parser.py:717-860` | terminal matching | If stack top is a terminal and matches the current token, consumes token by increasing `index`. |
| 6 | `Backend/parser/parser.py:702-715` and `863-1181` | syntax error creation | If no production or terminal mismatch occurs, creates a detailed syntax error. |
| 7 | `Backend/parser/parser.py:1196-1198` | builder call | After syntax success, filters newline tokens and calls `_build_ast(filtered)`. |

Parser defense sentence:

```text
After lexing, the parser does not read raw characters anymore.
It reads token types. It compares the current lookahead token with the LL(1)
PREDICT table to know which grammar production to use.
```

### C.6 CFG and PREDICT Decision Code

Exact flow:

| Step | File and Lines | Code Object | What Happens |
|---:|---|---|---|
| 1 | `Backend/cfg/grammar.py:122-631` | `cfg` dictionary | Stores productions such as `<program>`, `<statement>`, and `<expression>`. |
| 2 | `Backend/cfg/grammar.py:13-48` | `compute_first()` | Computes FIRST sets from the grammar. |
| 3 | `Backend/cfg/grammar.py:51-85` | `compute_follow()` | Computes FOLLOW sets from the grammar. |
| 4 | `Backend/cfg/grammar.py:88-119` | `compute_predict()` | Combines FIRST and FOLLOW to produce PREDICT sets. |
| 5 | `Backend/parser/parser.py:66-82` | `construct_parsing_table()` | Converts PREDICT sets into the parser table. |

### C.7 Semantic Process After Parser

Exact flow:

| Step | File and Lines | Code Object | What Happens |
|---:|---|---|---|
| 1 | `Backend/parser/parser.py:1196-1198` | `_build_ast(filtered)` | Builds AST using `Backend/parser/builder.py`. |
| 2 | `Backend/parser/builder.py:21-32` | `build_ast(tokens)` | Creates the builder parser, starts parsing, and returns AST plus symbol table. |
| 3 | `Backend/parser/builder.py:310-453` | `parse_variable()` | Handles variable declarations and declaration-related semantic checks. |
| 4 | `Backend/semantic/symbol_table.py:13-48` | `declare_variable()` | Stores variables and catches duplicates. |
| 5 | `Backend/semantic/symbol_table.py:50-58` | `lookup_variable()` | Finds variables or reports undeclared identifiers. |
| 6 | `Backend/server.py:359` | `validate_ast(ast)` | Runs the semantic validator after AST build. |
| 7 | `Backend/semantic/analyzer.py:4-23` | `ASTValidator.validate()` | Walks the AST and collects semantic errors/warnings. |
| 8 | `Backend/semantic/analyzer.py:133-141` | `_check_Break()` / `_check_Continue()` | Checks `prune` and `skip` context rules. |

Semantic defense sentence:

```text
Semantic analysis checks meaning. The grammar may say a statement shape is valid,
but semantic analysis checks if variables exist, types match, constants are not reassigned,
arrays are valid, and control statements are used in the correct place.
```

### C.8 Interpreter Process After Semantic Success

Exact flow:

| Step | File and Lines | Code Object | What Happens |
|---:|---|---|---|
| 1 | `Backend/server.py:370` | `Interpreter(socketio=collector)` | Creates the runtime interpreter. |
| 2 | `Backend/server.py:372` | `interpreter.interpret(ast)` | Starts execution using the AST. |
| 3 | `Backend/interpreter/interpreter.py:121-212` | `interpret(node)` | Dispatches AST nodes to matching `eval_*` methods. |
| 4 | `Backend/interpreter/interpreter.py:214-219` | `eval_program()` | Registers top-level declarations and automatically calls `root()`. |
| 5 | `Backend/interpreter/interpreter.py:50-92` | variable runtime methods | Stores, looks up, and updates runtime variables. |
| 6 | `Backend/interpreter/interpreter.py:516-683` | `eval_binary()` | Evaluates arithmetic, relational, logical, and concatenation expressions. |
| 7 | `Backend/interpreter/interpreter.py:750-804` | `eval_print()` | Handles `plant()` output. |
| 8 | `Backend/interpreter/interpreter.py:1375-1468` | `eval_input()` | Handles `water()` interactive input. |
| 9 | `Backend/interpreter/interpreter.py:857-899` | `eval_return()` / function call | Uses `ReturnValue` to stop function execution when `reclaim` runs. |

Interpreter defense sentence:

```text
The interpreter does not check grammar anymore. At this point, the code is already
valid. The interpreter walks the AST and performs the actual actions: assign values,
evaluate expressions, execute loops, print output, and request input.
```
""")

    p.append("""## Appendix A. Important File Interaction Map

```text
UI/main.js or UI/index.html
  -> sends source_code to Backend/server.py
Backend/server.py
  -> calls Backend/lexer/scanner.py lex(source_code)
Backend/lexer/scanner.py
  -> returns Token objects from Backend/shared/tokens.py
Backend/server.py
  -> calls LL1Parser.parse_and_build(tokens)
Backend/parser/parser.py
  -> uses Backend/cfg/grammar.py cfg and predict_sets
Backend/parser/parser.py
  -> calls Backend/parser/builder.py build_ast(tokens)
Backend/parser/builder.py
  -> uses Backend/shared/ast_nodes.py and Backend/semantic/symbol_table.py
Backend/server.py
  -> calls Backend/semantic/analyzer.py validate_ast(ast)
Backend/server.py
  -> calls Backend/interpreter/interpreter.py Interpreter.interpret(ast)
Backend/interpreter/interpreter.py
  -> returns output or InterpreterError
Backend/server.py
  -> sends JSON output/error to frontend
```

## Appendix B. Common Defense Questions and Direct Answers

| Question | Short Answer |
|---|---|
| Does the lexer scan the whole code at once? | The server passes the whole string to the lexer, but the lexer processes it one character at a time using `current_char` and `advance()`. |
| Where does `roof` become an identifier instead of `root`? | In `Backend/lexer/scanner.py`, after reserved-word matching fails, the fallback identifier loop collects remaining alphanumeric characters and appends `Token(TT_IDENTIFIER, ident_str, ...)`. |
| Why are quotes not in the CFG? | Because the lexer converts the whole quoted text into one `stringlit` token before parsing. |
| What chooses the CFG production? | The LL(1) parser uses the current lookahead token and the PREDICT set in `self.parsing_table`. |
| What catches undeclared variables? | `SymbolTable.lookup_variable()` and builder semantic checks raise `SemanticError`. |
| What executes the program? | `Interpreter.interpret(ast)` dispatches AST nodes to `eval_*` methods and automatically calls `root()`. |
| How is output displayed? | `plant()` becomes a `PrintNode`, `eval_print()` emits output, the server returns it as JSON, and the frontend displays it. |""")

    p.append("""## Appendix D. Removed CGMA Built-In Names

The old CGMA-style names `ts` and `taper` are no longer part of the backend compiler system.

Current system behavior:

| Old Name | Current Status | Why |
|---|---|---|
| `ts` / `.ts` | Removed | It came from the CGMA reference material, not the GrowALanguage reserved-word set. |
| `taper` / `.taper` | Removed | It came from the CGMA reference material, not the GrowALanguage reserved-word set. |

Important implementation note:
- These names are not lexer reserved words.
- The previous hidden builder/interpreter support was removed.
- The backend now has no `TaperNode`, `TSNode`, `eval_taper()`, or `eval_ts()` logic.

Recommended future GAL-style names, if this feature is added later:

| Feature | Recommended GAL Name | Example |
|---|---|---|
| array/string size | `.count` | `seed size = arr.count;` |
| split vine into leaf characters | `.leaves` | `leaf letters[3] = word.leaves;` |

These recommended names are documentation guidance only unless the backend is updated to implement them.""")

    p.append("""## Appendix E. Concrete Source Code Logic Notes

This section explains the main non-lexer source files more directly. It is meant for defense questions like "what does this file really do line by line?" or "how does the data move from this function to the next one?"

### E.1 `Backend/server.py` Route Logic

The server is the coordinator. It does not tokenize, parse, or execute by itself. Instead, it receives source code from the frontend and calls each compiler stage in order.""")

    p.append(snippet("Backend/server.py", 323, 384))
    p.append("""Concrete logic:

| Code Part | What It Means | Why It Matters |
|---|---|---|
| `@app.route('/api/run', methods=['POST'])` | Connects the `/api/run` URL to the Python function below it. | This is the endpoint used when the user presses Run without interactive input. |
| `def run_endpoint():` | Function that handles a run request. | This is the backend entry point for full compilation and execution. |
| `data = request.get_json()` | Reads the JSON body sent by JavaScript. | This is how Python receives the editor text. |
| `source_code = data['source_code']` | Stores the user program in a Python variable. | This exact variable is passed to the lexer. |
| `tokens, lex_errors = lex(source_code)` | Calls the lexer. | The raw string becomes a token list. |
| `if lex_errors:` | Checks if lexing failed. | The pipeline stops early for lexical errors. |
| `parse_result = parser.parse_and_build(tokens)` | Calls syntax parser and AST builder. | Tokens become an AST if grammar and builder checks pass. |
| `sem_result = validate_ast(ast, symbol_table)` | Runs semantic validation. | Checks meaning after syntax succeeds. |
| `Interpreter(socketio=collector)` | Creates runtime executor. | The interpreter will execute the AST and collect `plant()` output. |
| `interpreter.interpret(ast)` | Executes the program. | This is where actual runtime behavior happens. |
| `return jsonify(...)` | Sends output or error back to frontend. | The UI displays this response in the output panel. |

Important defense sentence:

```text
The server is the controller of the compiler pipeline. It receives source_code, then calls lexer, parser/builder, semantic validator, and interpreter. If any stage returns an error, the server stops and returns that error to the frontend.
```""")

    p.append("""### E.2 `OutputCollector` Logic

`OutputCollector` is used by the REST execution route so `plant()` output can be collected without needing a real Socket.IO client.""")
    p.append(snippet("Backend/server.py", 306, 316))
    p.append("""Concrete logic:

| Code Part | What It Means |
|---|---|
| `self.outputs = []` | Stores every output string produced during execution. |
| `emit(self, event, data=None)` | Mimics the Socket.IO `emit()` method. |
| `if event == 'output'` | Only output events are collected as printed output. |
| `self.outputs.append(...)` | Saves the output text so `/api/run` can return it as JSON. |

This is why the same interpreter can work for both REST execution and Socket.IO execution.""")

    p.append("""### E.3 `Backend/parser/parser.py` LL(1) Logic

The parser uses a stack. It does not read characters. It reads token types from the lexer.""")
    p.append(snippet("Backend/parser/parser.py", 614, 715))
    p.append("""Concrete logic:

| Variable | Meaning |
|---|---|
| `toks` | Normalized token list. Every token has `type`, `value`, `line`, and `col`. |
| `stack` | Parser work list. Starts with `EOF` and `<program>`. |
| `index` | Current position in the token list. |
| `current_token()` | Helper that returns the next non-skipped token. |
| `top` | Top symbol of the parser stack. Can be a terminal or nonterminal. |
| `token_type` | Current lookahead token type. |
| `row = self.parsing_table.get(top)` | Gets possible productions for the current nonterminal. |
| `production = row.get(token_type)` | Chooses the grammar production using the lookahead token. |

Decision logic:

| Situation | Parser Action |
|---|---|
| `top` is a nonterminal and PREDICT has a production | Pop `top`, push the selected production in reverse. |
| `top` is a terminal and matches current token | Pop terminal and move `index` to the next token. |
| No production exists for lookahead | Return syntax error. |
| Terminal does not match current token | Return syntax error. |

Example:

```text
Stack top: <program>
Current token: root
Parser checks parsing_table["<program>"]["root"]
It finds the <program> production and pushes its symbols onto the stack.
```

Defense sentence:

```text
The parser is table-driven. The current token does not directly call a function. Instead, the parser uses stack top plus lookahead token to find the correct grammar production in the parsing table.
```""")

    p.append("""### E.4 `Backend/parser/builder.py` AST Build Logic

After the LL(1) parser accepts the token order, the builder reads the same tokens again to create AST nodes and perform many semantic checks.""")
    p.append(snippet("Backend/parser/builder.py", 21, 32))
    p.append(snippet("Backend/parser/builder.py", 39, 108))
    p.append("""Concrete logic:

| Code Part | What It Means |
|---|---|
| `build_ast(tokens)` | Entry point called by `parse_and_build()`. |
| `parser = Parser(tokens)` | Creates a builder parser object that will walk through tokens. |
| `ast = parser.parse_program()` | Starts building the AST from the program level. |
| `symbol_table` | Stores declarations discovered while building. |
| `parse_program()` | Looks at top-level tokens and decides if each part is a variable, constant, bundle, function, or root. |

Important difference from the LL(1) parser:

```text
LL(1) parser checks if token order is allowed by CFG.
Builder creates meaning: declarations, assignments, function nodes, bundle nodes, and output/input nodes.
```""")

    p.append("""### E.5 `parse_statement()` Routing Logic

`parse_statement()` is the builder's dispatcher for statements inside functions and root.""")
    p.append(snippet("Backend/parser/builder.py", 442, 520))
    p.append("""Concrete logic:

| Token Seen | Builder Decision |
|---|---|
| data type like `seed`, `tree`, `vine` | Parse a variable declaration. |
| `fertile` | Parse a constant declaration. |
| `id` | Could be assignment, function call, array access, member access, increment, or list operation. |
| `plant` | Parse output statement. |
| `water` | Parse input statement. |
| `spring` | Parse conditional statement. |
| `cultivate`, `grow`, `tend` | Parse loop statement. |
| `harvest` | Parse switch-like statement. |
| `reclaim` | Parse return statement. |

The function is called repeatedly while building a block. Each successful parse returns:

```text
(node, new_index)
```

`node` is the AST node that represents the statement. `new_index` tells the builder where to continue reading tokens.""")

    p.append("""### E.6 `parse_variable()` Declaration Logic

This function handles declarations such as `seed x = 10;`, `tree y;`, `vine name = "GAL";`, and arrays.""")
    p.append(snippet("Backend/parser/builder.py", 310, 390))
    p.append("""Concrete logic:

| Step | Logic |
|---:|---|
| 1 | Read the declared type, such as `seed`. |
| 2 | Read the identifier name, such as `x`. |
| 3 | Check if the name conflicts with an existing function or variable. |
| 4 | If `[` appears, parse array dimension and reject invalid double sizes. |
| 5 | If `=` appears, parse initializer expression or initializer list. |
| 6 | Infer initializer type and compare it to declared type. |
| 7 | Store the variable in the symbol table. |
| 8 | Return a `VariableDeclarationNode`. |

This is why this code catches duplicate identifiers and type mismatch during building, before runtime.""")

    p.append("""### E.7 `parse_factor()` Expression Atom Logic

`parse_factor()` handles the smallest parts of expressions: literals, identifiers, grouped expressions, function calls, list access, and member access.""")
    p.append(snippet("Backend/parser/builder.py", 1352, 1438))
    p.append("""Concrete logic:

| Token Pattern | Meaning |
|---|---|
| `(` expression `)` | Parenthesized expression. |
| `intlit`, `dblit`, `chrlit`, `stringlit`, `sunshine`, `frost` | Literal value. |
| `id (` | Function call. |
| `id . wilt` or `id . bloom` | Supported vine member operation. |
| `id . id` | Bundle member access, only valid if the left identifier has bundle type. |
| `id [` | Array or vine indexing. |
| plain `id` | Variable value. |

Because `ts` and `taper` were removed, `arr.ts` and `c.taper` no longer have hidden special handling. They fall into the normal `id . id` member-access path and fail unless the left side is a real bundle with that member.""")

    p.append("""### E.8 `Backend/semantic/symbol_table.py` Logic

The symbol table is the compiler's memory during semantic checking. It remembers what names exist and what they mean.""")
    p.append(snippet("Backend/semantic/symbol_table.py", 13, 58))
    p.append("""Concrete logic:

| Function | Purpose |
|---|---|
| `declare_variable()` | Adds a variable to the current scope or global scope. |
| `lookup_variable()` | Searches current local scopes first, then global variables. |
| duplicate checks | Prevent declaring the same variable/function name in invalid places. |
| `is_list` | Records whether the variable is an array/list. |
| `is_fertile` | Records whether the variable is constant and cannot be reassigned. |

Defense sentence:

```text
The symbol table is how the compiler knows if an identifier was declared before use. Without it, the parser could accept the grammar but the compiler would not know whether a variable actually exists.
```""")

    p.append("""### E.9 `Backend/semantic/analyzer.py` Logic

The semantic analyzer walks the AST after the builder has created it.""")
    p.append(snippet("Backend/semantic/analyzer.py", 4, 35))
    p.append(snippet("Backend/semantic/analyzer.py", 133, 141))
    p.append("""Concrete logic:

| Code Part | Meaning |
|---|---|
| `validate(ast, symbol_table_data)` | Starts semantic validation. |
| `self._walk(ast)` | Recursively visits every AST node. |
| `getattr(self, f'_check_{node.node_type}', None)` | Dynamically finds a checker method based on node type. |
| `_check_Break()` | Ensures `prune` is inside a loop or harvest/switch. |
| `_check_Continue()` | Ensures `skip` is inside a loop. |

This file is a second semantic safety pass. The builder catches many type/declaration errors. The analyzer catches AST-level rules such as invalid control-flow usage.""")

    p.append("""### E.10 `Backend/interpreter/interpreter.py` Runtime Logic

The interpreter executes the AST. It no longer checks grammar. It performs actions.""")
    p.append(snippet("Backend/interpreter/interpreter.py", 121, 212))
    p.append("""Concrete logic:

| Code Part | Meaning |
|---|---|
| `interpret(node)` | Main dispatcher for runtime execution. |
| `isinstance(node, ProgramNode)` | If the node is a program, run the program. |
| `VariableDeclarationNode` | Create runtime variable storage. |
| `AssignmentNode` | Update runtime variable value. |
| `PrintNode` | Execute `plant()`. |
| `InputNode` | Execute `water()`. |
| `FunctionCallNode` | Call a function such as `root()`. |
| loop/if nodes | Execute control flow. |

Defense sentence:

```text
The interpreter is a dispatcher. It receives an AST node, checks what kind of node it is, then calls the matching eval method.
```""")

    p.append("""### E.11 Runtime Variables, Assignment, and Expressions""")
    p.append(snippet("Backend/interpreter/interpreter.py", 50, 92))
    p.append(snippet("Backend/interpreter/interpreter.py", 346, 430))
    p.append(snippet("Backend/interpreter/interpreter.py", 510, 585))
    p.append("""Concrete runtime flow:

| Runtime Action | Function | Logic |
|---|---|---|
| Declare variable | `declare_variable()` | Store name, type, value, list flag, and fertile flag in current scope. |
| Read variable | `lookup_variable()` | Search local scopes from newest to oldest, then global variables. |
| Assign variable | `eval_assignment()` | Evaluate right side, convert/check value, then update target variable. |
| Evaluate expression | `eval_binary_op()` | Evaluate left and right nodes, then apply operator such as `+`, `*`, `==`, `&&`, or backtick. |

Example:

```gal
sum = x + y;
```

Runtime logic:

```text
1. eval_assignment() receives AssignmentNode.
2. It interprets the right side BinaryOpNode.
3. eval_binary_op() looks up x and y.
4. It computes x + y.
5. eval_assignment() stores the result into sum.
```""")

    p.append("""### E.12 Function Call and Reclaim Logic""")
    p.append(snippet("Backend/interpreter/interpreter.py", 856, 899))
    p.append("""Concrete logic:

| Code Part | Meaning |
|---|---|
| `eval_function_call()` | Finds the function declaration at runtime. |
| argument binding | Evaluates arguments and stores them as parameter variables in a new scope. |
| `self.eval_block(...)` | Runs the function body statement by statement. |
| `ReturnValue` exception | Internal signal used by `reclaim` to stop the function body. |
| `finally: self.exit_scope()` | Ensures local variables are removed after the function ends. |

This is why `reclaim` can immediately stop a function even if there are more statements after it.""")

    p.append("""### E.13 `plant()` and `water()` Runtime Logic""")
    p.append(snippet("Backend/interpreter/interpreter.py", 752, 804))
    p.append(snippet("Backend/interpreter/interpreter.py", 1348, 1440))
    p.append("""Concrete logic:

| Feature | Runtime Function | Logic |
|---|---|---|
| `plant()` | `eval_print()` | Evaluates all print arguments and emits output text. |
| multiple plant arguments | `eval_print()` | Joins evaluated arguments with spaces. |
| `water()` | `eval_input()` | Requests input, waits for the frontend response, validates type, and stores converted value. |
| REST output | `OutputCollector.emit()` | Collects `output` events into a Python list. |
| Socket output | Socket.IO `emit()` | Sends live output/input events to the UI. |

This explains why output/errors are visible in the frontend output panel after execution.""")

    return "\n\n".join(p)


def set_cell_shading(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_docx(md_text):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    doc.styles["Normal"].font.size = Pt(10)

    def add_code_block(code):
        for line in code.splitlines() or [""]:
            para = doc.add_paragraph()
            run = para.add_run(line if line else " ")
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(8)
            para.paragraph_format.left_indent = Inches(0.22)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

    def clean(text):
        return text.replace("`", "")

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block("\n".join(code_lines))
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                rows = []
                for idx, tl in enumerate(table_lines):
                    if idx == 1:
                        continue
                    rows.append([clean(c.strip()) for c in tl.strip("|").split("|")])
                max_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(max_cols):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = row[c_idx] if c_idx < len(row) else ""
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        if r_idx == 0:
                            set_cell_shading(cell, "D9EAD3")
                doc.add_paragraph()
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = clean(line[level:].strip())
            if level == 1:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                doc.add_heading(text, 1)
            elif level == 3:
                doc.add_heading(text, 2)
            else:
                doc.add_heading(text, 3)
            i += 1
            continue

        if line.startswith("- "):
            doc.add_paragraph(clean(line[2:].strip()), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
            i += 1
            continue

        para = doc.add_paragraph(clean(line.strip()))
        para.paragraph_format.space_after = Pt(4)
        i += 1

    doc.save(DOCX_PATH)


def export_pdf():
    try:
        import win32com.client

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DOCX_PATH.resolve()))
        doc.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return True, str(PDF_PATH)
    except Exception as exc:
        return False, str(exc)


if __name__ == "__main__":
    markdown = build_markdown()
    MD_PATH.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    ok, detail = export_pdf()
    print(f"markdown={MD_PATH}")
    print(f"docx={DOCX_PATH}")
    print(f"pdf={'ok ' + detail if ok else 'failed ' + detail}")
