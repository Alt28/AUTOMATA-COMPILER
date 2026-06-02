from pathlib import Path
from datetime import datetime
import re

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Docus & mps"
MD_PATH = OUT_DIR / "GCD_INTERPRETER_SIMULATION.md"
DOCX_PATH = OUT_DIR / "GCD_INTERPRETER_SIMULATION.docx"
PDF_PATH = OUT_DIR / "GCD_INTERPRETER_SIMULATION.pdf"


PROGRAM = """pollinate seed gcd(seed a, seed b) {
    spring (b == 0) {
        reclaim a;
    }
    reclaim gcd(b, a % b);
}

root() {
    seed a;
    seed b;
    seed result;

    plant("Enter first number: ");
    a = water();
    plant("Enter second number: ");
    b = water();

    result = gcd(a, b);
    plant("GCD of", a,"and ",b, "is", result);

    reclaim;
}"""


def read_lines(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8", errors="replace").splitlines()


def code(rel_path, start, end, lang="python"):
    lines = read_lines(rel_path)
    last = min(end, len(lines))
    body = "\n".join(f"{idx:04d}: {lines[idx - 1]}" for idx in range(start, last + 1))
    return f"Source: `{rel_path}:{start}-{last}`\n\n```{lang}\n{body}\n```"


def table(headers, rows):
    out = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join(["---"] * len(headers)) + "|",
    ]
    for row in rows:
        out.append("| " + " | ".join(str(col).replace("\n", "<br>") for col in row) + " |")
    return "\n".join(out)


def build_markdown():
    parts = []

    parts.append("# GCD Program Interpreter Simulation")
    parts.append(
        "Detailed step-by-step explanation of how the GrowALanguage interpreter executes the recursive GCD program.\n\n"
        f"Project path: `{ROOT}`\n\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    parts.append("## 1. Program Under Simulation")
    parts.append("This is the exact GAL program being simulated:")
    parts.append(f"```gal\n{PROGRAM}\n```")
    parts.append(
        "Sample input used in this simulation:\n\n"
        "```text\n"
        "first number = 48\n"
        "second number = 18\n"
        "```\n\n"
        "Expected mathematical result:\n\n"
        "```text\n"
        "gcd(48, 18) = 6\n"
        "```"
    )

    parts.append("## 2. Important Interpreter View")
    parts.append(
        "By the time the interpreter runs, the source code is no longer plain text. "
        "The lexer already created tokens, the parser checked the CFG, the builder created AST nodes, "
        "and semantic validation already passed. The interpreter executes AST nodes such as "
        "`FunctionDeclarationNode`, `AssignmentNode`, `FunctionCallNode`, `IfStatementNode`, "
        "`ReturnNode`, and `PrintNode`."
    )

    parts.append("## 3. Backend Execution Path For This Program")
    parts.append(
        "Because this program uses `water()`, the interactive UI normally uses the Socket.IO run path. "
        "That path is in `Backend/server.py`."
    )
    parts.append(code("Backend/server.py", 411, 489))
    parts.append(
        "Explanation:\n\n"
        "- Lines 411-414 receive the source code from the UI.\n"
        "- Line 415 sends the complete source code to the lexer.\n"
        "- Lines 424-430 run parser plus AST builder using `parser.parse_and_build(tokens)`.\n"
        "- Lines 434-439 run semantic validation.\n"
        "- Lines 457-463 create an `Interpreter` object and execute the AST with `interp.interpret(ast)`.\n"
        "- Lines 482-489 receive input from the UI when `water()` asks for a value."
    )

    parts.append("## 4. AST Nodes Used By This Program")
    parts.append(code("Backend/shared/ast_nodes.py", 3, 65))
    parts.append(code("Backend/shared/ast_nodes.py", 83, 106))
    parts.append(
        "Main AST node meanings for the GCD program:\n\n"
        + table(
            ["GAL construct", "AST node", "Purpose"],
            [
                ["`pollinate seed gcd(...) { ... }`", "`FunctionDeclarationNode`", "Stores function name, return type, parameters, and body block."],
                ["`root() { ... }`", "`FunctionDeclarationNode`", "Stored as the main function that `eval_program()` calls."],
                ["`seed a;`", "`VariableDeclarationNode`", "Creates a variable record."],
                ["`a = water();`", "`AssignmentNode` + `Input` node", "Gets input then assigns it to `a`."],
                ["`result = gcd(a, b);`", "`AssignmentNode` + `FunctionCallNode`", "Calls `gcd`, receives returned value, stores it in `result`."],
                ["`spring (b == 0)`", "`IfStatementNode`", "Checks base case of recursion."],
                ["`reclaim a;`", "`ReturnNode`", "Returns a value from the function."],
                ["`plant(...)`", "`PrintNode`", "Sends output to UI."],
            ],
        )
    )

    parts.append("## 5. Builder Creates Function And Statement Nodes")
    parts.append(code("Backend/parser/builder.py", 195, 300))
    parts.append(
        "How this applies to the program:\n\n"
        "- `pollinate seed gcd(seed a, seed b)` becomes a `FunctionDeclarationNode` named `gcd`.\n"
        "- The return type is `seed`.\n"
        "- The parameter list stores two parameters: `seed a` and `seed b`.\n"
        "- `root()` is also built as a `FunctionDeclarationNode`, but with no parameters.\n"
        "- The block between `{` and `}` becomes the function body stored in the node."
    )
    parts.append(code("Backend/parser/builder.py", 1909, 1968))
    parts.append(
        "This code builds calls like `gcd(a, b)` and recursive calls like `gcd(b, a % b)`. "
        "It checks the function exists, counts arguments, checks argument types, and creates a `FunctionCallNode`."
    )
    parts.append(code("Backend/parser/builder.py", 1972, 2057))
    parts.append(
        "`water()` and `water(variable)` are turned into `Input` nodes. "
        "For `a = water();`, the assignment parser creates an `AssignmentNode` whose right side is an `Input` node."
    )
    parts.append(code("Backend/parser/builder.py", 2063, 2304))
    parts.append(
        "`plant(\"GCD of\", a,\"and \",b, \"is\", result);` becomes a `PrintNode` with multiple children. "
        "Since the first string has no `{}` placeholder, the interpreter uses the multi-argument join behavior."
    )
    parts.append(code("Backend/parser/builder.py", 2412, 2451))
    parts.append(
        "`spring (b == 0)` becomes an `IfStatementNode` with a condition wrapper and a block."
    )
    parts.append(code("Backend/parser/builder.py", 2518, 2560))
    parts.append(
        "`reclaim a;` and `reclaim gcd(b, a % b);` become `ReturnNode` objects. "
        "During execution, these nodes stop the current function call by raising a `ReturnValue` object."
    )

    parts.append("## 6. Interpreter Dispatch: How It Chooses What To Run")
    parts.append(code("Backend/interpreter/interpreter.py", 121, 190))
    parts.append(
        "The interpreter uses `interpret(node)` as a dispatcher. It checks the node class or node type, then sends it to the correct method. "
        "For this GCD program, the important branches are:\n\n"
        + table(
            ["Node seen by `interpret()`", "Line", "Method called"],
            [
                ["`ProgramNode`", "122-123", "`eval_program()`"],
                ["`FunctionDeclarationNode`", "140-141", "`eval_function_declaration()`"],
                ["`VariableDeclarationNode`", "130-131", "`eval_variable_declaration()`"],
                ["`AssignmentNode`", "132-133", "`eval_assignment()`"],
                ["`BinaryOpNode`", "134-139", "`eval_binary_op()`"],
                ["`PrintNode`", "142-143", "`eval_print()`"],
                ["`ReturnNode`", "148-149", "`eval_return()`"],
                ["`FunctionCallNode`", "150-151", "`eval_function_call()`"],
                ["`IfStatementNode`", "168-169", "`eval_if_statement()`"],
                ["`Input` node", "182-183", "`eval_input()`"],
                ["`Identifier` node", "187-190", "Looks variable up by name"],
            ],
        )
    )

    parts.append("## 7. Program Start: Register Functions Then Call root()")
    parts.append(code("Backend/interpreter/interpreter.py", 210, 215))
    parts.append(code("Backend/interpreter/interpreter.py", 716, 733))
    parts.append(
        "Execution starts with the whole `ProgramNode`.\n\n"
        "Step-by-step:\n\n"
        "1. `eval_program()` loops through the program children on lines 211-212.\n"
        "2. First child is the `gcd` function declaration, so `eval_function_declaration()` stores it in `self.functions`.\n"
        "3. Second child is the `root` function declaration, so it is also stored in `self.functions`.\n"
        "4. After all declarations are registered, line 214 creates `FunctionCallNode(\"root\", [])`.\n"
        "5. Line 215 executes that root call.\n\n"
        "After registration, the function table is conceptually:\n\n"
        + table(
            ["Function", "Return type", "Parameters", "Stored body"],
            [
                ["`gcd`", "`seed`", "`seed a`, `seed b`", "spring base case and recursive reclaim"],
                ["`root`", "`empty` or main/root type depending builder", "none", "input, function call, print, reclaim"],
            ],
        )
    )

    parts.append("## 8. Calling root()")
    parts.append(code("Backend/interpreter/interpreter.py", 856, 890))
    parts.append(
        "When `root()` is called:\n\n"
        "- Line 857 gets the function name: `root`.\n"
        "- Line 858 evaluates arguments. Root has no arguments, so `args = []`.\n"
        "- Line 860 finds `root` in `self.functions`.\n"
        "- Line 873 enters a new function scope.\n"
        "- Lines 876-881 declare parameters if any. Root has none.\n"
        "- Line 884 executes the root body block.\n"
        "- If `reclaim;` happens, line 886 catches the return and line 887 returns its value."
    )

    parts.append("## 9. Root Variable Declarations")
    parts.append(code("Backend/interpreter/interpreter.py", 218, 275))
    parts.append(code("Backend/interpreter/interpreter.py", 50, 61))
    parts.append(
        "Root executes:\n\n"
        "```gal\n"
        "seed a;\n"
        "seed b;\n"
        "seed result;\n"
        "```\n\n"
        "For each declaration:\n\n"
        "- Lines 219-221 read type, name, and optional value node.\n"
        "- Lines 224-230 define default values for each GAL type.\n"
        "- Because `seed a;`, `seed b;`, and `seed result;` have no explicit value, they receive default `0`.\n"
        "- Lines 55-61 store the variable record in the current scope.\n\n"
        "Root scope after declarations:\n\n"
        + table(
            ["Variable", "Type", "Initial value", "Stored in"],
            [
                ["`a`", "`seed`", "0", "`self.scopes[-1]`"],
                ["`b`", "`seed`", "0", "`self.scopes[-1]`"],
                ["`result`", "`seed`", "0", "`self.scopes[-1]`"],
            ],
        )
    )

    parts.append("## 10. First plant() And water()")
    parts.append(code("Backend/interpreter/interpreter.py", 752, 798))
    parts.append(code("Backend/interpreter/interpreter.py", 800, 805))
    parts.append(
        "For:\n\n"
        "```gal\n"
        "plant(\"Enter first number: \");\n"
        "```\n\n"
        "- Line 752 enters `eval_print()`.\n"
        "- Line 756 gets the first print argument.\n"
        "- Line 758 interprets the string node.\n"
        "- Lines 800-805 remove quotes and process escape basics in `eval_formatted_string()`.\n"
        "- Line 798 sends the string to `self.plant()`.\n"
        "- Line 745 emits output to the UI."
    )
    parts.append(code("Backend/interpreter/interpreter.py", 1312, 1346))
    parts.append(code("Backend/interpreter/interpreter.py", 1348, 1396))
    parts.append(code("Backend/interpreter/interpreter.py", 346, 358))
    parts.append(code("Backend/interpreter/interpreter.py", 489, 505))
    parts.append(
        "For:\n\n"
        "```gal\n"
        "a = water();\n"
        "```\n\n"
        "The interpreter executes an assignment node.\n\n"
        "Detailed flow:\n\n"
        "1. `eval_assignment()` starts at line 346.\n"
        "2. Line 347 gets target node `a`.\n"
        "3. Line 348 gets value node `Input/water()`.\n"
        "4. Line 356 calls `self.interpret(value_node)`, which routes to `eval_input()`.\n"
        "5. In `eval_input()`, lines 1354-1364 detect that the parent is an `AssignmentNode` and target is `a`.\n"
        "6. Line 1364 looks up `a` and discovers its type is `seed`.\n"
        "7. Lines 1374-1380 emit an input request and wait for the UI.\n"
        "8. Server receives the user input through `capture_input` in `server.py` lines 482-489.\n"
        "9. `provide_input()` lines 1315-1324 wakes the waiting input event.\n"
        "10. Since type is `seed`, lines 1385-1395 convert the text input `\"48\"` into integer `48`.\n"
        "11. Back in `eval_assignment()`, lines 489-505 store `a = 48`.\n\n"
        "After the first input:\n\n"
        + table(
            ["Variable", "Value"],
            [["`a`", "48"], ["`b`", "0"], ["`result`", "0"]],
        )
    )

    parts.append("## 11. Second plant() And water()")
    parts.append(
        "The second input follows the same flow:\n\n"
        "```gal\n"
        "plant(\"Enter second number: \");\n"
        "b = water();\n"
        "```\n\n"
        "With sample input `18`, `eval_input()` converts `\"18\"` into integer `18`, then `eval_assignment()` stores it in `b`.\n\n"
        "Root scope after both inputs:\n\n"
        + table(
            ["Variable", "Value"],
            [["`a`", "48"], ["`b`", "18"], ["`result`", "0"]],
        )
    )

    parts.append("## 12. result = gcd(a, b)")
    parts.append(
        "Now root executes:\n\n"
        "```gal\n"
        "result = gcd(a, b);\n"
        "```\n\n"
        "This is an assignment whose right side is a `FunctionCallNode`."
    )
    parts.append(code("Backend/interpreter/interpreter.py", 75, 103))
    parts.append(code("Backend/interpreter/interpreter.py", 856, 890))
    parts.append(
        "Flow for the first `gcd(a, b)` call:\n\n"
        "1. `eval_assignment()` line 356 interprets the right side.\n"
        "2. `interpret()` sees `FunctionCallNode` and calls `eval_function_call()` line 856.\n"
        "3. Line 857 reads function name `gcd`.\n"
        "4. Line 858 evaluates arguments `a` and `b`. It uses identifier lookup to get `48` and `18`.\n"
        "5. Line 860 uses `lookup_function()` to find the stored `gcd` declaration.\n"
        "6. Line 873 creates a new function scope.\n"
        "7. Lines 876-881 declare local parameter variables for this call: `a = 48`, `b = 18`.\n"
        "8. Line 884 executes the body of `gcd`."
    )

    parts.append("## 13. How spring(b == 0) Runs Inside gcd")
    parts.append(code("Backend/interpreter/interpreter.py", 1074, 1115))
    parts.append(code("Backend/interpreter/interpreter.py", 510, 600))
    parts.append(
        "Inside each `gcd` call, the first statement is:\n\n"
        "```gal\n"
        "spring (b == 0) {\n"
        "    reclaim a;\n"
        "}\n"
        "```\n\n"
        "The condition `b == 0` is a `BinaryOpNode`.\n\n"
        "- `eval_if_statement()` line 1075 calls `self.interpret()` on the condition.\n"
        "- `eval_binary_op()` lines 511-513 evaluate left, right, and operator.\n"
        "- Line 599 checks `==`.\n"
        "- Line 600 returns `left == right`.\n\n"
        "For the first call:\n\n"
        "```text\n"
        "b = 18\n"
        "18 == 0 -> False\n"
        "```\n\n"
        "Because it is false, line 1083 goes to the else path. Since there is no `bud` or `wither` inside the function, the if statement ends and execution continues to the next statement."
    )

    parts.append("## 14. How reclaim gcd(b, a % b) Runs")
    parts.append(code("Backend/interpreter/interpreter.py", 851, 853))
    parts.append(code("Backend/interpreter/interpreter.py", 586, 598))
    parts.append(
        "The next statement in `gcd` is:\n\n"
        "```gal\n"
        "reclaim gcd(b, a % b);\n"
        "```\n\n"
        "`eval_return()` line 852 interprets the return expression first. The return expression is another function call.\n\n"
        "For the first call, current local values are:\n\n"
        "```text\n"
        "a = 48\n"
        "b = 18\n"
        "```\n\n"
        "The recursive call is:\n\n"
        "```text\n"
        "gcd(b, a % b)\n"
        "gcd(18, 48 % 18)\n"
        "```\n\n"
        "Modulo is handled in `eval_binary_op()`:\n\n"
        "- Line 586 checks operator `%`.\n"
        "- Lines 596-597 prevent modulo by zero.\n"
        "- Line 598 returns `left % right`.\n\n"
        "So:\n\n"
        "```text\n"
        "48 % 18 = 12\n"
        "```\n\n"
        "The next recursive call becomes:\n\n"
        "```text\n"
        "gcd(18, 12)\n"
        "```"
    )

    parts.append("## 15. Complete Recursive Call Stack")
    parts.append(
        table(
            ["Depth", "Call", "`b == 0`?", "`a % b`", "Next action"],
            [
                ["1", "`gcd(48, 18)`", "`18 == 0` -> false", "`48 % 18 = 12`", "return `gcd(18, 12)`"],
                ["2", "`gcd(18, 12)`", "`12 == 0` -> false", "`18 % 12 = 6`", "return `gcd(12, 6)`"],
                ["3", "`gcd(12, 6)`", "`6 == 0` -> false", "`12 % 6 = 0`", "return `gcd(6, 0)`"],
                ["4", "`gcd(6, 0)`", "`0 == 0` -> true", "not needed", "run `reclaim a;`, return `6`"],
            ],
        )
    )
    parts.append(
        "Each recursive call has its own function scope. That means the parameter names `a` and `b` are reused, but each call has its own local values.\n\n"
        + table(
            ["Call depth", "Local `a`", "Local `b`", "Scope meaning"],
            [
                ["1", "48", "18", "First gcd call from root"],
                ["2", "18", "12", "Recursive call from depth 1"],
                ["3", "12", "6", "Recursive call from depth 2"],
                ["4", "6", "0", "Base case call"],
            ],
        )
    )

    parts.append("## 16. Base Case: reclaim a")
    parts.append(
        "At depth 4:\n\n"
        "```text\n"
        "gcd(6, 0)\n"
        "```\n\n"
        "The condition is true:\n\n"
        "```text\n"
        "b == 0\n"
        "0 == 0 -> true\n"
        "```\n\n"
        "`eval_if_statement()` line 1080 enters the true branch, and line 1081 evaluates the spring block. "
        "That block contains:\n\n"
        "```gal\n"
        "reclaim a;\n"
        "```\n\n"
        "`eval_return()` line 852 interprets `a`, which looks up the current local parameter value `6`. "
        "Then line 853 raises `ReturnValue(6)`. That is how the interpreter stops the current function body immediately."
    )
    parts.append(code("Backend/interpreter/interpreter.py", 883, 887))
    parts.append(
        "The raised return value is caught by `eval_function_call()`:\n\n"
        "- Line 884 is executing the function body.\n"
        "- Line 886 catches `ReturnValue`.\n"
        "- Line 887 returns `ret.value` to the caller.\n\n"
        "Return unwinding:\n\n"
        "```text\n"
        "gcd(6, 0) returns 6\n"
        "gcd(12, 6) returns 6\n"
        "gcd(18, 12) returns 6\n"
        "gcd(48, 18) returns 6\n"
        "```"
    )

    parts.append("## 17. Storing The Result")
    parts.append(
        "After `gcd(a, b)` returns `6`, the interpreter goes back to:\n\n"
        "```gal\n"
        "result = gcd(a, b);\n"
        "```\n\n"
        "`eval_assignment()` now stores the returned value in `result`."
    )
    parts.append(code("Backend/interpreter/interpreter.py", 489, 505))
    parts.append(
        "At this moment, root scope becomes:\n\n"
        + table(
            ["Variable", "Value"],
            [["`a`", "48"], ["`b`", "18"], ["`result`", "6"]],
        )
    )

    parts.append("## 18. Final plant() Output")
    parts.append(
        "The final print statement is:\n\n"
        "```gal\n"
        "plant(\"GCD of\", a,\"and \",b, \"is\", result);\n"
        "```\n\n"
        "In your interpreter, this uses the multi-argument print path because the first string `\"GCD of\"` does not contain `{}` placeholders."
    )
    parts.append(code("Backend/interpreter/interpreter.py", 752, 798))
    parts.append(
        "Detailed print behavior:\n\n"
        "- Line 758 evaluates the first argument: `\"GCD of\"`.\n"
        "- Line 764 checks if the first string contains `{}`. It does not.\n"
        "- Line 786 sees there are multiple arguments.\n"
        "- Line 787 starts `parts = [\"GCD of\"]`.\n"
        "- Lines 788-794 evaluate and append `a`, `\"and \"`, `b`, `\"is\"`, and `result`.\n"
        "- Line 795 prints `' '.join(parts)`.\n\n"
        "Because `\"and \"` already contains a trailing space and line 795 adds spaces between all parts, the actual output has two spaces before `18`:\n\n"
        "```text\n"
        "GCD of 48 and  18 is 6\n"
        "```\n\n"
        "Cleaner version if you want single spaces:\n\n"
        "```gal\n"
        "plant(\"GCD of\", a, \"and\", b, \"is\", result);\n"
        "```\n\n"
        "Or using placeholders:\n\n"
        "```gal\n"
        "plant(\"GCD of {} and {} is: {}\\n\", a, b, result);\n"
        "```"
    )

    parts.append("## 19. Final reclaim In root")
    parts.append(
        "The last statement is:\n\n"
        "```gal\n"
        "reclaim;\n"
        "```\n\n"
        "This becomes a `ReturnNode` with no value. `eval_return()` line 852 sets value to `None`, then line 853 raises `ReturnValue(None)`. "
        "`eval_function_call()` catches it at line 886 and returns from `root()`. Program execution is complete."
    )

    parts.append("## 20. Actual Simulated Output")
    parts.append(
        "With user inputs `48` and `18`, the backend test produced these output events:\n\n"
        "```text\n"
        "Enter first number: \n"
        "[input_required: Input for a]\n"
        "Enter second number: \n"
        "[input_required: Input for b]\n"
        "GCD of 48 and  18 is 6\n"
        "```\n\n"
        "The `input_required` events are not normal printed text from `plant()`. They are UI events emitted by `eval_input()` so the frontend can ask the user for a value."
    )

    parts.append("## 21. Whole Interpreter Simulation Summary")
    parts.append(
        table(
            ["Step", "GAL code", "Interpreter method", "Result"],
            [
                ["1", "`pollinate seed gcd(...)`", "`eval_function_declaration()` line 716", "Stores `gcd` in `self.functions`."],
                ["2", "`root() { ... }`", "`eval_function_declaration()` line 716", "Stores `root` in `self.functions`."],
                ["3", "Program start", "`eval_program()` line 210", "Creates and executes `FunctionCallNode(\"root\")`."],
                ["4", "`seed a; seed b; seed result;`", "`eval_variable_declaration()` line 218", "Creates variables with default `0`."],
                ["5", "`plant(\"Enter first number: \")`", "`eval_print()` line 752", "Outputs prompt."],
                ["6", "`a = water();`", "`eval_input()` line 1348 + `eval_assignment()` line 346", "Stores input `48` in `a`."],
                ["7", "`b = water();`", "`eval_input()` line 1348 + `eval_assignment()` line 346", "Stores input `18` in `b`."],
                ["8", "`result = gcd(a, b);`", "`eval_function_call()` line 856", "Runs recursive GCD and returns `6`."],
                ["9", "`spring (b == 0)`", "`eval_if_statement()` line 1074", "Controls recursion base case."],
                ["10", "`a % b`", "`eval_binary_op()` line 586", "Computes remainder for next recursive call."],
                ["11", "`reclaim a;`", "`eval_return()` line 851", "Returns final answer `6`."],
                ["12", "Final `plant(...)`", "`eval_print()` line 752", "Outputs `GCD of 48 and  18 is 6`."],
                ["13", "`reclaim;`", "`eval_return()` line 851", "Stops `root()` and completes program."],
            ],
        )
    )

    parts.append("## 22. Defense Script Taglish")
    parts.append(
        "Pwede mong sabihin sa defense:\n\n"
        "```text\n"
        "Sa interpreter part, hindi na raw source code ang binabasa ng system. AST nodes na siya. "
        "Una, sa eval_program, nireregister muna niya yung pollinate function na gcd at yung root function. "
        "Pag tapos na maregister, automatic niyang tatawagin yung root.\n\n"
        "Sa root, yung seed a, seed b, at seed result ay ginagawa niyang variables sa current scope. "
        "Then yung plant ay dumadaan sa eval_print para mag-output sa UI. "
        "Yung water naman ay dumadaan sa eval_input. Dahil naka-assign siya sa variable a or b, tinitingnan ng interpreter yung type ng variable. "
        "Since seed siya, kino-convert niya yung input string into integer.\n\n"
        "Pag dumating sa result = gcd(a, b), tatawagin niya yung eval_function_call. "
        "Doon gumagawa siya ng bagong scope para sa parameters ng gcd. "
        "Sa unang call, a is 48 and b is 18. Chine-check niya yung spring b == 0. "
        "Kapag false, gagawin niya yung recursive reclaim gcd(b, a % b). "
        "Kaya magiging gcd(18, 12), then gcd(12, 6), then gcd(6, 0).\n\n"
        "Pag gcd(6, 0), true na yung b == 0 kaya reclaim a, meaning return 6. "
        "Yung return value na 6 babalik pataas sa lahat ng recursive calls hanggang mastore siya sa result. "
        "Finally, ipiprint niya yung GCD message, then reclaim sa root to end the program.\n"
        "```"
    )

    parts.append("## 23. Quick Line Reference Table")
    parts.append(
        table(
            ["File", "Line(s)", "Purpose"],
            [
                ["`Backend/server.py`", "411-489", "Interactive run path and input capture."],
                ["`Backend/parser/builder.py`", "195-300", "Builds function declarations."],
                ["`Backend/parser/builder.py`", "1909-1968", "Builds function calls such as `gcd(a, b)`."],
                ["`Backend/parser/builder.py`", "1972-2057", "Builds `water()` input nodes."],
                ["`Backend/parser/builder.py`", "2063-2304", "Builds `plant()` print nodes."],
                ["`Backend/parser/builder.py`", "2412-2451", "Builds `spring` if nodes."],
                ["`Backend/parser/builder.py`", "2518-2560", "Builds `reclaim` return nodes."],
                ["`Backend/interpreter/interpreter.py`", "121-190", "Main node dispatcher."],
                ["`Backend/interpreter/interpreter.py`", "210-215", "Registers declarations and calls `root`."],
                ["`Backend/interpreter/interpreter.py`", "716-733", "Stores function declarations."],
                ["`Backend/interpreter/interpreter.py`", "856-890", "Executes function calls and catches returns."],
                ["`Backend/interpreter/interpreter.py`", "218-275", "Executes variable declarations."],
                ["`Backend/interpreter/interpreter.py`", "346-505", "Executes assignments."],
                ["`Backend/interpreter/interpreter.py`", "510-600", "Executes binary operations including `==`, `%`, and arithmetic."],
                ["`Backend/interpreter/interpreter.py`", "752-798", "Executes `plant()`."],
                ["`Backend/interpreter/interpreter.py`", "851-853", "Executes `reclaim`."],
                ["`Backend/interpreter/interpreter.py`", "1074-1115", "Executes `spring/bud/wither`."],
                ["`Backend/interpreter/interpreter.py`", "1312-1396", "Handles interactive input and conversion."],
            ],
        )
    )

    return "\n\n".join(parts)


def clean(text):
    return text.replace("`", "")


def set_cell_shading(cell, fill):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    except Exception:
        pass


def build_docx(markdown):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(13)

    def add_code_block(text):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        for line in text.splitlines():
            run = paragraph.add_run(line + "\n")
            run.font.name = "Consolas"
            run.font.size = Pt(7.5)

    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block("\n".join(code_lines))
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1

            rows = []
            for row_idx, table_line in enumerate(table_lines):
                if row_idx == 1:
                    continue
                rows.append([clean(cell.strip()) for cell in table_line.strip("|").split("|")])

            if rows:
                max_cols = max(len(row) for row in rows)
                tbl = doc.add_table(rows=len(rows), cols=max_cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(max_cols):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = row[c_idx] if c_idx < len(row) else ""
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        if r_idx == 0:
                            set_cell_shading(cell, "D9EAD3")
                doc.add_paragraph()
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = clean(line[level:].strip())
            if level == 1:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                doc.add_heading(text, 1)
            elif level == 3:
                doc.add_heading(text, 2)
            else:
                doc.add_heading(text, 3)
            i += 1
            continue

        if line.startswith("- "):
            doc.add_paragraph(clean(line[2:].strip()), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
            i += 1
            continue

        para = doc.add_paragraph(clean(line.strip()))
        para.paragraph_format.space_after = Pt(4)
        i += 1

    doc.save(DOCX_PATH)


def export_pdf():
    try:
        import win32com.client

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DOCX_PATH.resolve()))
        doc.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return True, str(PDF_PATH)
    except Exception as exc:
        return False, str(exc)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = build_markdown()
    MD_PATH.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    ok, detail = export_pdf()
    print(f"markdown={MD_PATH}")
    print(f"docx={DOCX_PATH}")
    print(f"pdf={'ok ' + detail if ok else 'failed ' + detail}")
