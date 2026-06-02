from pathlib import Path
from datetime import datetime
import re

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Docus & mps"
MD_PATH = OUT_DIR / "LINE_BY_LINE_CODE_STUDY_GUIDE_GROWALANGUAGE_COMPILER.md"
DOCX_PATH = OUT_DIR / "LINE_BY_LINE_CODE_STUDY_GUIDE_GROWALANGUAGE_COMPILER.docx"
PDF_PATH = OUT_DIR / "LINE_BY_LINE_CODE_STUDY_GUIDE_GROWALANGUAGE_COMPILER.pdf"


def read_lines(rel):
    path = ROOT / rel
    try:
        return path.read_text(encoding="utf-8").splitlines()
    except UnicodeDecodeError:
        return path.read_text(errors="replace").splitlines()


def code(rel, start, end, lang="python"):
    lines = read_lines(rel)
    body = "\n".join(f"{idx:04d}: {lines[idx - 1]}" for idx in range(start, min(end, len(lines)) + 1))
    return f"Source: `{rel}:{start}-{end}`\n\n```{lang}\n{body}\n```"


def js_code(rel, start, end):
    return code(rel, start, end, "javascript")


def table(headers, rows):
    out = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in rows:
        out.append("| " + " | ".join(str(c).replace("\n", "<br>") for c in row) + " |")
    return "\n".join(out)


def build_markdown():
    parts = []
    parts.append("# Line-by-Line Code Study Guide of GrowALanguage Compiler")
    parts.append(
        "This file is a defense study guide. It explains the actual source code by exact file, "
        "line number, function, variable, and logic.\n\n"
        f"Project path: `{ROOT}`\n\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    parts.append("""## How To Use This Guide

The compiler has very large files, especially `builder.py`, `parser.py`, and `interpreter.py`. So this guide explains the code line-by-line by important function or block. That means each section shows:

- exact source file and line numbers
- numbered source code snippet
- meaning of each important line or group of lines
- what input the code receives
- what output it returns
- what compiler stage depends on it

This is the best way to study the whole system without drowning in repeated helper lines.""")

    parts.append("""## 0. Python Terms You Must Understand First

| Term | Meaning In Your Compiler |
|---|---|
| `self` | The current object. Example: inside `Lexer`, `self.current_char` means the current character owned by this lexer object. |
| `self.variable` | An instance variable. It keeps its value while the object is running. Example: `self.pos`, `self.source_code`, `self.current_char`. |
| `def function_name(...)` | Defines a function or method. Example: `def advance(self):` defines how the lexer moves forward. |
| `return` | Sends a result back to the caller. Example: `return tokens, errors`. |
| `list` | Ordered collection. Example: `tokens = []` stores generated token objects. |
| `dict` | Key-value collection. Example: `cfg` maps nonterminals to grammar productions. |
| `class` | Blueprint for objects. Example: `class Lexer:` creates lexer objects. |
| `append()` | Adds one item to the end of a list. Example: `tokens.append(Token(...))`. |
| `None` | Means no value. In the lexer, `current_char = None` means end of source code. |

Defense shortcut:

```text
`self.current_char` is not a random variable. It is the character currently being scanned by the Lexer object.
Every time `advance()` runs, the position moves forward and `self.current_char` changes to the next character.
```""")

    parts.append("""## 1. Frontend To Backend Input Flow""")
    parts.append("### 1.1 Frontend Reads Editor Code")
    parts.append(js_code("UI/main.js", 1005, 1035))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1005", "`window.runProgram` is the main frontend function for running code."),
            ("1007", "`const sourceCode = editor.getValue();` gets the whole text from the Monaco editor."),
            ("1010-1012", "Clears previous visual error highlights so old errors do not stay on screen."),
            ("1014", "Runs the lexer silently to update the token table before execution."),
            ("1026", "Checks if the code contains `water(...)`. This matters because `water()` needs interactive input."),
            ("1028-1032", "If input is needed, use Socket.IO. Otherwise, use REST `/api/run`."),
        ],
    ))

    parts.append("### 1.2 Frontend Sends REST Request")
    parts.append(js_code("UI/main.js", 910, 916))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("910", "`runViaREST(sourceCode)` receives the full source program string."),
            ("911", "Calls `fetch()` to send a POST request to `/api/run`."),
            ("913", "Sends JSON, so Flask can read it with `request.get_json()`."),
            ("915", "The key is `source_code`; the value is the full text from the editor."),
        ],
    ))

    parts.append("### 1.3 Backend Receives Source Code")
    parts.append(code("Backend/server.py", 323, 384))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("323", "`run_endpoint()` is the Flask route function for `/api/run`."),
            ("325", "`data = request.get_json()` reads JSON sent by the frontend."),
            ("326", "`source_code = data['source_code']` stores the whole program string in Python."),
            ("340", "`tokens, lex_errors = lex(source_code)` sends the source to the lexer."),
            ("341-344", "If lexer errors exist, the pipeline stops and returns lexical error JSON."),
            ("348", "`parser.parse_and_build(tokens)` sends lexer tokens to the parser and AST builder."),
            ("349-353", "If parsing fails, the pipeline stops and returns syntax/parse error JSON."),
            ("359", "`validate_ast(ast)` performs semantic analysis after parsing succeeds."),
            ("360-364", "If semantic errors exist, the pipeline stops and returns semantic error JSON."),
            ("370-372", "Creates an `Interpreter` and executes the AST."),
            ("377-384", "Returns final runtime output or runtime error to the frontend."),
        ],
    ))

    parts.append("""## 2. Lexer Core: Character-by-Character Scanning""")
    parts.append("### 2.1 Lexer Initialization")
    parts.append(code("Backend/lexer/scanner.py", 18, 27))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("18", "`class Lexer:` creates the scanner object type."),
            ("19", "`__init__` runs automatically when `Lexer(source_code)` is created."),
            ("20", "`source_code.replace('\\r', '')` removes Windows carriage returns, keeping newline handling consistent."),
            ("21", "`self.pos = Position(-1, 1, -1)` starts before the first character. Index and column are `-1` because `advance()` will move to 0."),
            ("22", "`self.current_char = None` means no character is selected yet."),
            ("23", "`self.advance()` immediately moves to the first source character."),
            ("25-27", "`advance()` updates position, then sets `current_char` to the character at the new index. If past the end, it becomes `None`."),
        ],
    ))

    parts.append("""Plain explanation:

```text
self.source_code = the whole code text
self.pos = where the lexer is currently located
self.current_char = the exact character the lexer is looking at right now
advance() = move one character forward
```

Example with `seed`:

| advance call | pos.index | current_char |
|---:|---:|---|
| before advance | -1 | None |
| 1st | 0 | `s` |
| 2nd | 1 | `e` |
| 3rd | 2 | `e` |
| 4th | 3 | `d` |
| 5th | 4 | None, because source ended |""")

    parts.append("### 2.2 Position Tracking")
    parts.append(code("Backend/lexer/positions.py", 3, 21))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("3-7", "`Position` stores filename, index, line number, and column number."),
            ("9-11", "Every advance moves index and column forward."),
            ("12-14", "If the previous character was newline, line number increases and column resets to 0."),
            ("16-21", "`copy()` creates a snapshot of the current position. The lexer uses this as the starting position of a token."),
        ],
    ))

    parts.append("### 2.3 Main Lexer Loop")
    parts.append(code("Backend/lexer/scanner.py", 29, 35))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("30", "`tokens = []` is where token objects are stored."),
            ("31", "`line = 1` stores the current line counter used in many token creations."),
            ("32", "`errors = []` stores lexical error objects."),
            ("34", "`while self.current_char != None:` means keep scanning until the end of the source."),
            ("35", "If the current character is alphabetic, the lexer enters reserved-word/identifier logic."),
        ],
    ))

    parts.append("### 2.4 Reserved Word Recognition: `root`")
    parts.append(code("Backend/lexer/scanner.py", 341, 379))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("341", "This branch runs when the first character collected is `r`."),
            ("342-344", "Adds `r`, moves forward, checks if next character is `o`."),
            ("345-348", "Adds first `o`, moves forward, checks if next character is second `o`."),
            ("349-352", "Adds second `o`, moves forward, checks if next character is `t`."),
            ("353-356", "Adds `t`, moves forward, now `ident_str` is `root`."),
            ("357-358", "Checks if `root` is followed by end/space/valid root delimiter `(`."),
            ("360-373", "If invalid delimiter exists after `root`, create lexical error."),
            ("375", "If delimiter is valid, append `Token(TT_RW_ROOT, ident_str, line, pos.col)`."),
            ("377", "`break` exits the alphabet branch because the token is finished."),
        ],
    ))

    parts.append("""Example:

```gal
root()
```

Process:

| Character | ident_str after reading | What the lexer checks |
|---|---|---|
| `r` | `r` | maybe a word starting with r |
| `o` | `ro` | maybe `root` |
| `o` | `roo` | still maybe `root` |
| `t` | `root` | exact reserved word found |
| `(` | not added | valid delimiter, so token becomes `root` |""")

    parts.append("### 2.5 Fallback to Identifier")
    parts.append(code("Backend/lexer/scanner.py", 611, 638))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("611", "This block runs if the alphabet text was not accepted as a reserved word."),
            ("613", "Identifier length limit is 15."),
            ("614-615", "While the current character is alphanumeric, add it to `ident_str` and advance."),
            ("616-628", "If the final identifier is too long, split/report length error."),
            ("630", "`remaining = ident_str[i:]` gets the valid remaining identifier text after length handling."),
            ("631-632", "If the next character is a valid identifier delimiter, append an `id` token."),
            ("633-637", "If the next character is not a valid delimiter, create lexical error."),
            ("638", "Break out because this lexeme is finished."),
        ],
    ))

    parts.append("""Important example: `roof()`

```text
The lexer tries to match root.
It reads r, o, o.
Then it expects t, but sees f.
The reserved-word condition fails.
No token was appended yet, so Python naturally continues to the fallback identifier block.
The fallback while-loop collects f.
The final ident_str becomes roof.
Because the next character is (, and ( is valid in idf_delim, it creates Token(TT_IDENTIFIER, "roof", ...).
```

The code that collects `f` is:

```python
while self.current_char is not None and self.current_char in ALPHANUM:
    ident_str += self.current_char
    self.advance()
```

`self.current_char` is `f`, so `ident_str += self.current_char` changes `roo` into `roof`.""")

    parts.append("### 2.6 Number Literals")
    parts.append(code("Backend/lexer/scanner.py", 1171, 1236))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1171", "Runs when current character is a digit."),
            ("1173-1175", "Starts collecting number characters and saves starting position."),
            ("1177-1182", "Collects all digits into `digits`."),
            ("1184-1199", "If a dot appears, scan decimal part and validate decimal structure."),
            ("1201-1214", "If the character after the number is another invalid digit/letter/dot pattern, create lexical error."),
            ("1216-1227", "If there was no decimal point and delimiter is valid, append `intlit`."),
            ("1228-1236", "If there is a decimal point and delimiter is valid, append `dblit`."),
        ],
    ))

    parts.append("### 2.7 String Literal")
    parts.append(code("Backend/lexer/scanner.py", 1334, 1388))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1334", "Runs when current character is double quote."),
            ("1336-1341", "Initializes string buffer and moves past the opening quote."),
            ("1343-1358", "Reads string contents until closing quote or end of file. Handles escape sequences."),
            ("1360-1366", "If no closing quote exists, create lexical error."),
            ("1368-1372", "If closing quote exists, include it and advance."),
            ("1374-1386", "Check valid delimiter after string, then append `stringlit`."),
        ],
    ))

    parts.append("### 2.8 Character Literal")
    parts.append(code("Backend/lexer/scanner.py", 1390, 1456))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1390", "Runs when current character is single quote."),
            ("1392-1400", "Starts collecting character literal content."),
            ("1402-1422", "Reads characters and escape sequences until closing quote."),
            ("1424-1430", "Reports unclosed character literal."),
            ("1432-1439", "Counts actual character length, allowing valid escapes."),
            ("1441-1454", "Checks delimiter and appends `chrlit` if valid."),
        ],
    ))

    parts.append("### 2.9 Comments")
    parts.append(code("Backend/lexer/scanner.py", 1071, 1106))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1071-1073", "When scanner sees `/`, it checks if next char creates comment or operator."),
            ("1075-1084", "`//` comment: keep advancing until newline or end."),
            ("1086-1106", "`/* ... */` comment: keep advancing until closing `*/`; error if not closed."),
        ],
    ))

    parts.append("### 2.10 Operator Example: `**=`")
    parts.append(code("Backend/lexer/scanner.py", 811, 847))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("811", "Runs when current character is `*`."),
            ("813-815", "Starts symbol buffer with `*` and advances."),
            ("817", "If next char is another `*`, it may be exponent operator."),
            ("818-820", "Builds `**`."),
            ("821-825", "If next char is `=`, builds `**=` and appends exponent-assignment token."),
            ("826-832", "Otherwise, appends exponent token `**`."),
            ("835-840", "If after first `*` is `=`, append `*=` token."),
            ("842-846", "Otherwise, append normal multiplication token `*`."),
        ],
    ))

    parts.append("### 2.11 Lexer Wrapper Return")
    parts.append(code("Backend/lexer/scanner.py", 1502, 1529))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1502-1505", "When scanning ends, append EOF token and return tokens/errors."),
            ("1513", "`lex(source_code)` is the function imported by the server."),
            ("1515", "Creates a `Lexer` object from the source string."),
            ("1516", "Runs `make_tokens()` to scan all characters."),
            ("1523-1528", "Converts lexical errors into strings."),
            ("1529", "Returns token list and error list to the server."),
        ],
    ))

    parts.append("""## 3. Tokens and Delimiters""")
    parts.append("### 3.1 Token Class")
    parts.append(code("Backend/shared/tokens.py", 86, 92))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("86", "`class Token:` defines the object used to store each token."),
            ("87", "`type_` is the category, such as `seed`, `id`, `intlit`, or `;`."),
            ("88", "`value` is the actual lexeme text, such as `x` or `10`."),
            ("89", "`line` stores where the token appears."),
            ("90", "`col` stores the starting column."),
            ("92", "`__repr__` controls how tokens print during debugging."),
        ],
    ))

    parts.append("### 3.2 Token Constants")
    parts.append(code("Backend/shared/tokens.py", 3, 80))
    parts.append("""These constants prevent hardcoding random strings everywhere. Instead of writing `"seed"` many times, the lexer can use `TT_DT_INT = "seed"`. The parser receives these token types and compares them against CFG terminals.""")

    parts.append("### 3.3 Delimiter Sets")
    parts.append(code("Backend/lexer/delimiters.py", 1, 60))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1-8", "Defines character groups: letters, digits, zero digit, alphanumeric."),
            ("13-20", "Defines general delimiter groups like whitespace, data-type delimiters, brackets, and string delimiters."),
            ("22-53", "Defines numbered delimiter sets used by different reserved words/symbols."),
            ("54", "`idf_delim` is valid next character after an identifier."),
            ("55", "`whlnum_delim` is valid next character after an integer literal."),
            ("56", "`decim_delim` is valid next character after a decimal literal."),
            ("60", "`comment_delim` is for comment scanning."),
        ],
    ))

    parts.append("""Delimiter defense shortcut:

```text
A token is the thing the parser receives.
A delimiter set is only a lexer validation rule for the next character.
For example, ++ is a token, but delim25 only checks what can legally come after ++.
```""")

    parts.append("""### 3.4 Removed CGMA Names: `ts` and `taper`

The old CGMA-style names `ts` and `taper` were removed from the backend system.

What changed:

| Old Code Item | Current Status |
|---|---|
| `TaperNode` | Removed from `Backend/shared/ast_nodes.py`. |
| `TSNode` | Removed from `Backend/shared/ast_nodes.py`. |
| `eval_taper()` | Removed from `Backend/interpreter/interpreter.py`. |
| `eval_ts()` | Removed from `Backend/interpreter/interpreter.py`. |
| builder checks for `.taper` | Removed from `Backend/parser/builder.py`. |
| builder checks for `.ts` | Removed from `Backend/parser/builder.py`. |

Plain explanation:

```text
The lexer never treated ts or taper as reserved words.
Before cleanup, the builder gave them hidden special meaning when it saw id . ts or id . taper.
Now that hidden support is removed, so they behave like unsupported member access.
```

Recommended future names if the feature is added again:

```gal
seed size = arr.count;
leaf letters[3] = word.leaves;
```

Important: `.count` and `.leaves` are recommended names only. They are not implemented until backend support is added.""")

    parts.append("""## 4. Parser Detailed Code Flow""")
    parts.append("### 4.1 Token View Normalization")
    parts.append(code("Backend/parser/parser.py", 13, 34))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("13-19", "`_TokView` is a simple normalized token format used by the parser."),
            ("22-34", "`_as_tok()` converts real `Token` objects or dictionaries into `_TokView`."),
            ("29-33", "If the token is a `Token` object, it reads `type`, `value`, `line`, and `col`."),
        ],
    ))

    parts.append("### 4.2 Parser Initialization")
    parts.append(code("Backend/parser/parser.py", 37, 64))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("37", "`LL1Parser` is the main parser class."),
            ("38-47", "Constructor receives CFG, PREDICT sets, FIRST sets, start symbol, EOF marker, and skip token types."),
            ("48-53", "Stores grammar data into instance variables."),
            ("54", "Defines aliases, for example newline token mapping."),
            ("55", "Builds the LL(1) parsing table."),
            ("56-64", "Creates state variables used for better error messages."),
        ],
    ))

    parts.append("### 4.3 Parsing Table")
    parts.append(code("Backend/parser/parser.py", 66, 82))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("67", "Creates empty dictionary for parsing table."),
            ("68-69", "Loops through each nonterminal and its productions in the CFG."),
            ("70", "Creates row for current nonterminal."),
            ("71", "Gets the PREDICT set for the production."),
            ("72-79", "For every terminal in the PREDICT set, store which production should be used."),
            ("74-78", "If a table cell is already filled, that means LL(1) conflict."),
            ("81-82", "Returns the completed parsing table."),
        ],
    ))

    parts.append("### 4.4 Main Parse Loop")
    parts.append(code("Backend/parser/parser.py", 614, 715))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("615", "Converts token list into normalized `_TokView` tokens."),
            ("618-620", "Ensures EOF exists at the end."),
            ("624", "Initial stack is `[EOF, <program>]`."),
            ("626", "`index = 0` means parser starts at the first token."),
            ("627-631", "`current_token()` skips newline/comment tokens if configured."),
            ("641-644", "If stack becomes empty, parsing is finished."),
            ("646-664", "If top is EOF, check if current token is EOF."),
            ("665", "If stack top is a nonterminal, use parsing table."),
            ("666-669", "Get current lookahead token type."),
            ("672-690", "If production exists, pop nonterminal and push production symbols in reverse order."),
            ("702-715", "If no production exists, create syntax error and return failure."),
        ],
    ))

    parts.append("### 4.5 Terminal Match and Consume")
    parts.append(code("Backend/parser/parser.py", 717, 760))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("717", "If stack top is terminal and matches current token type, it is accepted."),
            ("718", "Stores actual token for state/error tracking."),
            ("719-760", "Updates parser state depending on important tokens like `pollinate`, `root`, braces, and declaration keywords."),
            ("Later in this block", "After handling state, the parser pops the terminal and moves to next token."),
        ],
    ))

    parts.append("### 4.6 Parse and Build")
    parts.append(code("Backend/parser/parser.py", 1183, 1240))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1184", "Runs syntax parsing first."),
            ("1185-1194", "If syntax fails, return error and no AST."),
            ("1196-1198", "If syntax succeeds, remove newline tokens and call AST builder."),
            ("1199-1205", "Return success, no errors, AST, and symbol table."),
            ("1206-1240", "Catch semantic or unexpected builder errors and format them."),
        ],
    ))

    parts.append("""Parser defense shortcut:

```text
The parser no longer reads characters. It reads token types.
It uses stack top plus current token lookahead to choose grammar production from the parsing table.
If the token matches the expected terminal, it consumes it and moves forward.
```""")

    parts.append("""## 5. CFG, FIRST, FOLLOW, PREDICT Code""")
    parts.append("### 5.1 FIRST Set Function")
    parts.append(code("Backend/cfg/grammar.py", 13, 48))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("13", "`compute_first(cfg)` receives the grammar dictionary."),
            ("14", "Creates empty FIRST set for every nonterminal."),
            ("16-17", "`first_of(symbol)` is a helper function."),
            ("18-20", "If the symbol is a terminal, its FIRST set is itself."),
            ("22-23", "If FIRST was already computed, return it."),
            ("25-39", "Loops through productions and adds possible starting terminals."),
            ("41-47", "Computes FIRST for all nonterminals."),
            ("48", "Returns FIRST sets."),
        ],
    ))

    parts.append("### 5.2 FOLLOW Set Function")
    parts.append(code("Backend/cfg/grammar.py", 51, 85))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("51", "`compute_follow` receives CFG, FIRST sets, and start symbol."),
            ("52", "Creates empty FOLLOW set for every nonterminal."),
            ("53", "Adds EOF to FOLLOW of start symbol."),
            ("56-83", "Repeats until no FOLLOW set changes."),
            ("61-81", "When a nonterminal appears inside a production, add what can legally follow it."),
            ("85", "Returns FOLLOW sets."),
        ],
    ))

    parts.append("### 5.3 PREDICT Set Function")
    parts.append(code("Backend/cfg/grammar.py", 88, 119))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("88", "`compute_predict` builds PREDICT sets from CFG, FIRST, and FOLLOW."),
            ("91-94", "Loops through every production."),
            ("96-110", "Computes FIRST of that production."),
            ("112-116", "If production can become lambda, add FOLLOW of the nonterminal."),
            ("117", "Stores the computed PREDICT set."),
            ("119", "Returns all PREDICT sets."),
        ],
    ))

    parts.append("""## 6. AST Builder and Semantic Checks""")
    parts.append("### 6.1 Builder Entry Point")
    parts.append(code("Backend/parser/builder.py", 21, 32))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("21", "`build_ast(tokens)` is called by `parse_and_build()` after syntax success."),
            ("22", "Creates a builder `Parser` object."),
            ("23", "Calls `parse_program()` to build the AST root."),
            ("24-30", "Catches semantic errors and unexpected errors."),
            ("32", "Returns AST and symbol table."),
        ],
    ))

    parts.append("### 6.2 Program-Level Build Flow")
    parts.append(code("Backend/parser/builder.py", 39, 108))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("39-46", "If token is a data type, parse variable declaration."),
            ("48-61", "If token is `fertile`, parse constant declaration."),
            ("63-90", "If token is `pollinate`, parse function definition."),
            ("93-100", "If token is `bundle`, parse bundle declaration."),
            ("102-108", "If token is `root`, parse main/root function."),
        ],
    ))

    parts.append("### 6.3 Variable Declaration")
    parts.append(code("Backend/parser/builder.py", 310, 453))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("310", "`parse_variable()` handles data-type declarations like `seed x;`."),
            ("314-320", "Reads data type token and variable identifier."),
            ("322-335", "Rejects duplicate function/variable conflicts."),
            ("337-370", "Checks if variable name exists already in current scope."),
            ("373-385", "Handles array dimensions and rejects invalid double size."),
            ("389-425", "Handles initializers including array initializer braces."),
            ("432-445", "Infers initializer type and checks compatibility."),
            ("448-453", "Declares variable into symbol table and creates AST declaration node."),
        ],
    ))

    parts.append("### 6.4 Symbol Table Variable Storage")
    parts.append(code("Backend/semantic/symbol_table.py", 13, 58))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("13", "`declare_variable()` stores a variable in the correct scope."),
            ("14-20", "Rejects duplicate names against functions and existing variables."),
            ("22-28", "If inside a function, store the variable in function-local records."),
            ("30-40", "If inside current block scope, store it in the latest scope dictionary."),
            ("42-48", "Otherwise store it globally."),
            ("50-58", "`lookup_variable()` searches local scopes first, then globals, then reports undeclared variable."),
        ],
    ))

    parts.append("### 6.5 Type Compatibility")
    parts.append(code("Backend/parser/builder.py", 1006, 1011))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1007", "Exact same type is compatible."),
            ("1009", "`seed` and `tree` are both numeric, so helper treats them as compatible."),
            ("1011", "Everything else is incompatible."),
        ],
    ))

    parts.append("### 6.6 Assignment Semantic Check")
    parts.append(code("Backend/parser/builder.py", 1930, 1958))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1930", "`validate_assignment()` checks if assigning to a variable is legal."),
            ("1931", "Looks up variable in symbol table."),
            ("1932-1934", "If variable does not exist, raise semantic error."),
            ("1936-1938", "If variable is `fertile`, reject reassignment."),
            ("1945-1955", "Infers assigned expression type and checks compatibility."),
            ("1958", "Returns success if assignment is valid."),
        ],
    ))

    parts.append("### 6.7 Function Call Check")
    parts.append(code("Backend/parser/builder.py", 1961, 2021))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("1961", "Validates a function call."),
            ("1962-1964", "Looks up function declaration."),
            ("1966-1969", "Checks argument count."),
            ("1971-2016", "Checks each argument type against each parameter type."),
            ("2021", "Returns success if call is valid."),
        ],
    ))

    parts.append("### 6.8 Water, Plant, Reclaim")
    parts.append(code("Backend/parser/builder.py", 2024, 2112))
    parts.append(code("Backend/parser/builder.py", 2115, 2188))
    parts.append(code("Backend/parser/builder.py", 2570, 2612))
    parts.append("""Meaning:

- `parse_water_statement()` checks that input goes into a valid assignable variable/location.
- `parse_plant_statement()` parses output arguments.
- `parse_reclaim_statement()` checks function return type and creates the return/reclaim node.""")

    parts.append("### 6.9 Semantic Validator")
    parts.append(code("Backend/semantic/analyzer.py", 4, 35))
    parts.append(code("Backend/semantic/analyzer.py", 133, 141))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("4-23", "`ASTValidator.validate()` walks the AST and collects errors/warnings."),
            ("26-35", "`_walk()` dispatches each AST node to a specific checker."),
            ("133-136", "`prune` is valid only inside loops or harvest/switch."),
            ("138-141", "`skip` is valid only inside loops."),
        ],
    ))

    parts.append("""## 7. Interpreter Runtime Code""")
    parts.append("### 7.1 Interpreter Initialization")
    parts.append(code("Backend/interpreter/interpreter.py", 25, 48))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("25", "`Interpreter` executes the AST after semantic success."),
            ("26-27", "`output` stores printed output; `socketio` is used for frontend communication."),
            ("28", "`variables` stores global runtime variables."),
            ("29", "`functions` stores runtime function declarations."),
            ("30", "`scopes` stores local runtime scopes."),
            ("31-36", "Loop/control flags handle break, continue, and input state."),
            ("37-48", "Input fields support `water()` waiting for user input."),
        ],
    ))

    parts.append("### 7.2 Runtime Variable Methods")
    parts.append(code("Backend/interpreter/interpreter.py", 50, 92))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("50-57", "`declare_variable()` stores a variable in current local scope or global storage."),
            ("59-70", "`lookup_variable()` searches local scopes first, then globals."),
            ("72-92", "`set_variable()` updates existing variable value in the nearest valid scope."),
        ],
    ))

    parts.append("### 7.3 Central Interpreter Dispatcher")
    parts.append(code("Backend/interpreter/interpreter.py", 121, 212))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("121", "`interpret(node)` receives one AST node."),
            ("122-124", "If the node is a list, interpret each item."),
            ("126-127", "If node is empty, return nothing."),
            ("129-212", "Checks node type/class and calls the correct `eval_*` method."),
        ],
    ))

    parts.append("### 7.4 Program Starts at Root")
    parts.append(code("Backend/interpreter/interpreter.py", 214, 219))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("214", "`eval_program()` executes the whole AST program."),
            ("215-216", "Interprets top-level declarations first."),
            ("218-219", "Automatically calls `root()` as the main function."),
        ],
    ))

    parts.append("### 7.5 Variable Declaration Runtime")
    parts.append(code("Backend/interpreter/interpreter.py", 222, 298))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("222", "Runs when executing a variable declaration node."),
            ("225-248", "Handles list/array initialization."),
            ("250-266", "Handles single-value initialization or default values."),
            ("269-295", "Converts runtime value based on declared type."),
            ("297", "Stores final variable in runtime scope."),
        ],
    ))

    parts.append("### 7.6 Assignment Runtime")
    parts.append(code("Backend/interpreter/interpreter.py", 352, 513))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("352", "Runs when executing an assignment AST node."),
            ("354-381", "Handles array/list element assignment."),
            ("382-414", "Handles bundle/member assignment."),
            ("416-470", "Handles compound assignment operators."),
            ("472-511", "Converts value according to declared variable type."),
            ("513", "Stores the new value."),
        ],
    ))

    parts.append("### 7.7 Expression Runtime")
    parts.append(code("Backend/interpreter/interpreter.py", 516, 683))
    parts.append(table(
        ["Lines", "Meaning"],
        [
            ("516", "Runs when evaluating binary expressions."),
            ("517-523", "Evaluates left and right operands; backtick concatenates strings."),
            ("525-572", "Handles arithmetic operators."),
            ("574-621", "Handles comparison operators."),
            ("623-647", "Handles logical operators."),
            ("649-683", "Handles unary-like operation forms and error checks."),
        ],
    ))

    parts.append("### 7.8 Output and Input Runtime")
    parts.append(code("Backend/interpreter/interpreter.py", 750, 804))
    parts.append(code("Backend/interpreter/interpreter.py", 1375, 1468))
    parts.append("""Meaning:

- `eval_print()` executes `plant(...)`.
- It evaluates every print argument, joins them, and emits output.
- `eval_input()` executes `water(...)`.
- It asks the frontend for input, waits for it, validates the input type, then stores the converted value.""")

    parts.append("### 7.9 Loops and Reclaim")
    parts.append(code("Backend/interpreter/interpreter.py", 1101, 1192))
    parts.append(code("Backend/interpreter/interpreter.py", 857, 899))
    parts.append("""Meaning:

- `eval_if()` handles `bud` / `wither` branches.
- `eval_for()` handles `cultivate`.
- `eval_return()` raises `ReturnValue`.
- `eval_function_call()` catches `ReturnValue`, so `reclaim` stops function execution properly.""")

    parts.append("""## 8. Error Classes and How Errors Return To UI""")
    parts.append("### 8.1 Lexical Error")
    parts.append(code("Backend/lexer/errors.py", 3, 11))
    parts.append("### 8.2 Semantic Error")
    parts.append(code("Backend/semantic/errors.py", 6, 12))
    parts.append("### 8.3 Runtime Error")
    parts.append(code("Backend/interpreter/errors.py", 6, 27))
    parts.append(table(
        ["Error Type", "Detected By", "Returned By Server"],
        [
            ("Lexical", "`scanner.py` using `LexicalError`", "`run_endpoint()` returns stage `lexical`."),
            ("Syntax", "`parser.py` inside `LL1Parser.parse()`", "`run_endpoint()` returns stage `parse` or `syntax`."),
            ("Semantic", "`builder.py` or `analyzer.py` using `SemanticError`", "`run_endpoint()` returns stage `semantic`."),
            ("Runtime", "`interpreter.py` using `InterpreterError`", "`run_endpoint()` returns stage `runtime`."),
        ],
    ))

    parts.append("""## 9. Full Mini Simulation: `seed num = 10;`

This simulation shows how the code is processed through the lexer and later stages.

### 9.1 Lexer Simulation

Input:

```gal
seed num = 10;
```

| Step | current_char | Code Area | Result |
|---:|---|---|---|
| 1 | `s` | `scanner.py` alphabet branch | Lexer tries reserved words starting with `s`. |
| 2 | `s e e d` | reserved word check | `ident_str` becomes `seed`; token `seed` is appended. |
| 3 | space | whitespace branch | Lexer skips whitespace. |
| 4 | `n` | alphabet branch | Not a reserved word, fallback identifier loop collects `num`. |
| 5 | space | identifier delimiter check | Space is valid in `idf_delim`, so token `id` with value `num` is appended. |
| 6 | `=` | operator branch | Token `=` is appended. |
| 7 | `1` | number branch | Digit loop collects `10`; token `intlit` is appended. |
| 8 | `;` | semicolon branch | Token `;` is appended. |
| 9 | end | EOF append | Token `EOF` is appended. |

Final token sequence:

```text
seed id = intlit ; EOF
```

### 9.2 Parser Simulation

The parser does not see characters anymore. It sees:

```text
seed id = intlit ;
```

If this appears inside root local declaration, it matches:

```text
<local_declaration> -> <var_dec> ; <local_declaration>
<var_dec> -> <data_type> id <var_tail>
<data_type> -> seed
<var_tail> -> = <init_val>
<init_val> -> <expression>
<factor> -> intlit
```

### 9.3 Semantic Simulation

The builder/semantic stage checks:

| Check | Result |
|---|---|
| Is `num` already declared in same scope? | No, so allowed. |
| Is `10` compatible with `seed`? | Yes, `10` is `intlit`, so type is `seed`. |
| Should `num` be stored in symbol table? | Yes. |

### 9.4 Interpreter Simulation

At runtime:

| Step | Runtime Effect |
|---:|---|
| 1 | Interpreter evaluates initializer `10`. |
| 2 | Converts/stores value as seed integer. |
| 3 | Runtime scope stores `num = 10`. |""")

    parts.append("""## 10. Quick Defense Script For Code-Level Questions

```text
Kapag tinanong po kung paano binabasa ng lexer yung code, ang sagot is:
Yung buong source code po ay pinapasa muna sa Lexer object. Pero hindi niya ito binabasa as one whole meaning agad. Meron siyang self.current_char, which means current character na tinitingnan niya. Every time na tatawag siya ng advance(), gagalaw yung position and mag-uupdate yung current_char.

For example, sa word na root, chine-check niya character by character kung r, then o, then o, then t. Kapag complete and valid yung delimiter after root, gagawa siya ng Token root.

Kapag hindi exact reserved word, like roof, hindi siya error agad. Magfa-fall back siya sa identifier block. Yung while loop na may ALPHANUM ang magko-collect ng remaining letters, kaya magiging identifier token yung roof.

After lexer, token list na ang pinapasa sa parser. Parser naman uses stack and PREDICT table. Hindi na raw characters ang binabasa niya, token types na.

After parser, builder creates AST and checks semantics using symbol table. Then kapag pasado, interpreter executes the AST and handles variables, expressions, plant output, water input, loops, and reclaim.
```""")

    return "\n\n".join(parts)


def set_cell_shading(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_docx(md_text):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    doc.styles["Normal"].font.size = Pt(10)

    def clean(text):
        return text.replace("`", "")

    def add_code_block(text):
        for line in text.splitlines() or [""]:
            para = doc.add_paragraph()
            run = para.add_run(line if line else " ")
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(7.5)
            para.paragraph_format.left_indent = Inches(0.18)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block("\n".join(code_lines))
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for idx, table_line in enumerate(table_lines):
                if idx == 1:
                    continue
                rows.append([clean(c.strip()) for c in table_line.strip("|").split("|")])
            if rows:
                max_cols = max(len(row) for row in rows)
                tbl = doc.add_table(rows=len(rows), cols=max_cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(max_cols):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = row[c_idx] if c_idx < len(row) else ""
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                        if r_idx == 0:
                            set_cell_shading(cell, "D9EAD3")
                doc.add_paragraph()
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = clean(line[level:].strip())
            if level == 1:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                doc.add_heading(text, 1)
            elif level == 3:
                doc.add_heading(text, 2)
            else:
                doc.add_heading(text, 3)
            i += 1
            continue
        if line.startswith("- "):
            doc.add_paragraph(clean(line[2:].strip()), style="List Bullet")
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
            i += 1
            continue
        para = doc.add_paragraph(clean(line.strip()))
        para.paragraph_format.space_after = Pt(4)
        i += 1

    doc.save(DOCX_PATH)


def export_pdf():
    try:
        import win32com.client

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DOCX_PATH.resolve()))
        doc.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return True, str(PDF_PATH)
    except Exception as exc:
        return False, str(exc)


if __name__ == "__main__":
    markdown = build_markdown()
    MD_PATH.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    ok, detail = export_pdf()
    print(f"markdown={MD_PATH}")
    print(f"docx={DOCX_PATH}")
    print(f"pdf={'ok ' + detail if ok else 'failed ' + detail}")
