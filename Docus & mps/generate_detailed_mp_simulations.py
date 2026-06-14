from pathlib import Path
from datetime import datetime
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Docus & mps"

FILES = {
    "server": ROOT / "Backend/server.py",
    "scanner": ROOT / "Backend/lexer/scanner.py",
    "parser": ROOT / "Backend/parser/parser.py",
    "builder": ROOT / "Backend/parser/builder.py",
    "semantic": ROOT / "Backend/semantic/analyzer.py",
    "interpreter": ROOT / "Backend/interpreter/interpreter.py",
}
LINES = {key: path.read_text(encoding="utf-8").splitlines() for key, path in FILES.items()}


def find(file_key, text):
    for line_no, line in enumerate(LINES[file_key], 1):
        if text in line:
            return line_no
    return "?"


def snippet(file_key, start, end):
    if not isinstance(start, int) or not isinstance(end, int):
        return "Line reference not found."
    file_lines = LINES[file_key]
    return "\n".join(
        f"{line_no}: {file_lines[line_no - 1]}"
        for line_no in range(start, min(end, len(file_lines)) + 1)
    )


def esc(value):
    return str(value).replace("|", "\\|").replace("\n", "<br>")


REF = {
    "server_route": find("server", "@app.route('/api/run'"),
    "server_json": find("server", "data = request.get_json()"),
    "server_source": find("server", "source_code = data['source_code']"),
    "server_lex": find("server", "tokens, lex_errors = lex(source_code)"),
    "server_parse": find("server", "parse_result = parser.parse_and_build(tokens)"),
    "server_semantic": find("server", "semantic_result = validate_ast(parse_result['ast'], parse_result['symbol_table'])"),
    "server_interpreter": find("server", "interp = Interpreter(socketio=collector)"),
    "server_execute": find("server", "interp.interpret(ast)"),
    "lex_public": find("scanner", "def lex(source_code)"),
    "lexer_init": find("scanner", "def __init__(self, source_code)"),
    "lexer_source": find("scanner", "self.source_code = source_code.replace"),
    "lexer_pos": find("scanner", "self.pos = Position"),
    "lexer_current": find("scanner", "self.current_char = None"),
    "lexer_advance": find("scanner", "def advance(self)"),
    "lexer_advance_pos": find("scanner", "self.pos.advance(self.current_char)"),
    "lexer_advance_load": find("scanner", "self.current_char = self.source_code"),
    "lexer_make": find("scanner", "def make_tokens(self)"),
    "lexer_tokens": find("scanner", "tokens = []"),
    "lexer_errors": find("scanner", "errors = []"),
    "lexer_loop": find("scanner", "while self.current_char != None"),
    "lexer_return": find("scanner", "return tokens, errors"),
    "tok_pollinate": find("scanner", "Token(TT_RW_POLLINATE"),
    "tok_seed": find("scanner", "Token(TT_RW_SEED"),
    "tok_root": find("scanner", "Token(TT_RW_ROOT"),
    "tok_cultivate": find("scanner", "Token(TT_RW_CULTIVATE"),
    "tok_spring": find("scanner", "Token(TT_RW_SPRING"),
    "tok_bud": find("scanner", "Token(TT_RW_BUD"),
    "tok_wither": find("scanner", "Token(TT_RW_WITHER"),
    "tok_plant": find("scanner", "Token(TT_RW_PLANT"),
    "tok_reclaim": find("scanner", "Token(TT_RW_RECLAIM"),
    "tok_id": find("scanner", "Token(TT_IDENTIFIER, ident_str"),
    "tok_int": find("scanner", "Token(TT_INTEGERLIT, ident_str"),
    "tok_string": find("scanner", "Token(TT_STRINGLIT"),
    "tok_comment": find("scanner", "Token(TT_COMMENT"),
    "tok_lte": find("scanner", "Token(TT_LTEQ"),
    "tok_gte": find("scanner", "Token(TT_GTEQ"),
    "tok_eq": find("scanner", "Token(TT_EQTO"),
    "tok_mod": find("scanner", "tokens.append(Token(TT_MOD"),
    "tok_inc": find("scanner", "Token(TT_INCREMENT"),
    "tok_mul": find("scanner", "Token(TT_MUL"),
    "parser_parse_and_build": find("parser", "def parse_and_build"),
    "parser_calls_parse": find("parser", "syntax_ok, syntax_errors = self.parse(tokens)"),
    "parser_parse": find("parser", "def parse(self, tokens"),
    "parser_tokview": find("parser", "toks = [_as_tok"),
    "parser_eof": find("parser", "toks = self._ensure_eof"),
    "parser_stack": find("parser", "stack: List[str]"),
    "parser_loop": find("parser", "while stack:"),
    "parser_top": find("parser", "top = stack[-1]"),
    "parser_token": find("parser", "tok = current_token()"),
    "parser_skip": find("parser", "if token_type in self.skip_token_types"),
    "parser_row": find("parser", "row = self.parsing_table[top]"),
    "parser_prod": find("parser", "production = row[token_type]"),
    "parser_push": find("parser", "stack.extend(reversed(production))"),
    "parser_match": find("parser", "if top == token_type"),
    "parser_filter": find("parser", "filtered = [t for t in tokens"),
    "parser_ast": find("parser", "ast = _build_ast(filtered)"),
    "builder_build": find("builder", "def build_ast(tokens)"),
    "builder_program": find("builder", "root = ProgramNode()"),
    "builder_reset": find("builder", "symbol_table.variables = {}"),
    "builder_loop": find("builder", "while index < len(tokens)"),
    "builder_function": find("builder", "def parse_function(tokens"),
    "builder_variable": find("builder", "def parse_variable"),
    "builder_statement": find("builder", "def parse_statement"),
    "builder_call": find("builder", "def parse_function_call"),
    "builder_if": find("builder", "def parse_if"),
    "builder_for": find("builder", "def parse_for"),
    "builder_return": find("builder", "def parse_return"),
    "semantic_validate_ast": find("semantic", "def validate_ast"),
    "semantic_validate": find("semantic", "def validate(self, ast"),
    "semantic_walk": find("semantic", "def _walk"),
    "semantic_handler": find("semantic", "handler = getattr"),
    "semantic_for": find("semantic", "def _check_ForLoop"),
    "interpreter_init": find("interpreter", "def __init__(self, socketio=None)"),
    "interpreter_interpret": find("interpreter", "def interpret"),
    "interpreter_program": find("interpreter", "def eval_program"),
    "interpreter_main": find("interpreter", "main_call = FunctionCallNode"),
    "interpreter_function_declaration": find("interpreter", "def eval_function_declaration"),
    "interpreter_function_call": find("interpreter", "def eval_function_call"),
    "interpreter_args": find("interpreter", "args = [self.interpret"),
    "interpreter_enter_scope": find("interpreter", "self.enter_scope()"),
    "interpreter_run_body": find("interpreter", "self.eval_block(function_node.children[2])"),
    "interpreter_variable": find("interpreter", "def eval_variable_declaration"),
    "interpreter_assignment": find("interpreter", "def eval_assignment"),
    "interpreter_binary": find("interpreter", "def eval_binary_op"),
    "interpreter_mul": find("interpreter", "elif operator == '*':"),
    "interpreter_mod": find("interpreter", "elif operator == '%':"),
    "interpreter_eq": find("interpreter", "elif operator == '==':"),
    "interpreter_lte": find("interpreter", "elif operator == '<=':"),
    "interpreter_gte": find("interpreter", "elif operator == '>=':"),
    "interpreter_unary": find("interpreter", "def eval_unaryop"),
    "interpreter_if": find("interpreter", "def eval_if_statement"),
    "interpreter_for": find("interpreter", "def eval_for_loop"),
    "interpreter_for_condition": find("interpreter", "condition_result = self.interpret(condition_node)"),
    "interpreter_for_update": find("interpreter", "for update_expr in node.children[2].children"),
    "interpreter_print": find("interpreter", "def eval_print"),
    "interpreter_print_join": find("interpreter", 'self.plant(" ".join(parts))'),
    "interpreter_return": find("interpreter", "def eval_return"),
}


PROGRAMS = [
    {
        "slug": "MP1_TIMESTHREE",
        "title": "MP1 - Function + Cultivate + Spring/Wither",
        "source": """pollinate seed timesThree(seed n) {
    reclaim n * 3;
}
\t
root() {
    seed value = 0;
\tseed i;
    cultivate  (i = 1; i <= 4; i++) {
        value = timesThree(i);
        spring (value % 2 == 0) {
            plant("Even product:",value);
        } wither {
            plant("Odd  product:",value);
        }
    }

    reclaim;
}""",
        "purpose": "This program shows the complete path for a user-defined function, a cultivate loop, a function call inside assignment, modulo/equality checking, spring/wither branching, and plant output.",
        "lexer": [
            f"The first real word is `pollinate`. The lexer recognizes it as a reserved word and appends a pollinate token at scanner.py line {REF['tok_pollinate']}.",
            f"`seed` is recognized as a data type reserved word at scanner.py line {REF['tok_seed']}.",
            f"`timesThree` is not a reserved word. After the lexer collects its letters, it becomes an identifier token at scanner.py line {REF['tok_id']}.",
            f"`reclaim n * 3;` becomes reclaim, id, multiplication, integer literal, and semicolon. Multiplication is appended at scanner.py line {REF['tok_mul']}; integer literals are appended at line {REF['tok_int']}.",
            f"`i <= 4` uses the `<=` token at scanner.py line {REF['tok_lte']}; `i++` uses increment at line {REF['tok_inc']}.",
            f"`value % 2 == 0` uses `%` at scanner.py line {REF['tok_mod']} and `==` at line {REF['tok_eq']}.",
            f"The strings inside plant are collected as whole string literal tokens at scanner.py line {REF['tok_string']}. The quote symbols are not separate parser terminals.",
        ],
        "parser": [
            "`<program>` starts with lookahead `pollinate`, so the parser accepts a function definition before root.",
            "The function header `pollinate seed timesThree(seed n) {` is matched token by token.",
            "`reclaim n * 3;` is parsed as a reclaim statement with an expression.",
            "After the function closes, parser matches the required `root() {` structure.",
            "Inside root, parser accepts local declarations first: `seed value = 0;` and `seed i;`.",
            "The parser then accepts `cultivate (i = 1; i <= 4; i++) { ... }` as a for-loop style statement.",
            "Inside the loop, `value = timesThree(i);` is parsed as assignment with a function call expression.",
            "Then parser checks the spring/wither chain and final `reclaim;`.",
        ],
        "builder": [
            f"`build_ast()` at builder.py line {REF['builder_build']} creates the ProgramNode.",
            f"When builder sees `pollinate`, `parse_function()` at line {REF['builder_function']} creates a FunctionDeclarationNode for `timesThree`.",
            f"`parse_return()` at line {REF['builder_return']} creates the ReturnNode for `reclaim n * 3;`.",
            f"`parse_variable()` at line {REF['builder_variable']} creates VariableDeclarationNodes for `value` and `i`.",
            f"`parse_for()` at line {REF['builder_for']} creates the ForLoopNode.",
            f"`parse_function_call()` at line {REF['builder_call']} checks that `timesThree(i)` matches the function declaration.",
            f"`parse_if()` at line {REF['builder_if']} creates the IfStatementNode for spring/wither.",
        ],
        "semantic": [
            "Semantic validation walks the AST made by builder. It does not scan the source again.",
            f"`validate_ast()` starts at analyzer.py line {REF['semantic_validate_ast']}; `_walk()` at line {REF['semantic_walk']} visits each AST node.",
            f"The cultivate loop activates `_check_ForLoop()` at analyzer.py line {REF['semantic_for']}, which marks that the validator is inside a loop.",
            "No semantic errors are found because `value` and `i` are declared, `timesThree(i)` has the correct argument type, and the spring condition produces branch/boolean.",
        ],
        "interpreter": [
            f"`eval_program()` at interpreter.py line {REF['interpreter_program']} registers function declarations first.",
            f"`eval_function_declaration()` at line {REF['interpreter_function_declaration']} stores `timesThree` and `root` in `self.functions`.",
            f"After registration, line {REF['interpreter_main']} creates a runtime call to `root()`.",
            f"`eval_function_call()` at line {REF['interpreter_function_call']} runs root in a new scope.",
            f"`eval_variable_declaration()` at line {REF['interpreter_variable']} creates `value = 0` and `i = 0`.",
            f"`eval_for_loop()` at line {REF['interpreter_for']} runs the cultivate loop: initializer, condition, body, update, repeat.",
            f"`value = timesThree(i)` runs through `eval_assignment()` at line {REF['interpreter_assignment']}. The RHS calls `timesThree`, binds parameter `n`, and evaluates `n * 3` using `eval_binary_op()` line {REF['interpreter_binary']}.",
            f"The spring condition uses `%` branch at line {REF['interpreter_mod']} and `==` branch at line {REF['interpreter_eq']}.",
            f"`eval_print()` at line {REF['interpreter_print']} emits each output line.",
        ],
        "runtime": [
            ("Register", "timesThree and root", "Both functions are stored before root executes."),
            ("Root declaration", "value = 0, i = 0", "value has initializer 0; i gets default seed value 0."),
            ("Loop init", "i = 1", "cultivate initializer runs once."),
            ("Iteration 1", "i=1, timesThree(1)=3, 3%2==0 is false", "wither prints `Odd  product: 3`."),
            ("Iteration 2", "i=2, timesThree(2)=6, 6%2==0 is true", "spring prints `Even product: 6`."),
            ("Iteration 3", "i=3, timesThree(3)=9, false", "wither prints `Odd  product: 9`."),
            ("Iteration 4", "i=4, timesThree(4)=12, true", "spring prints `Even product: 12`."),
            ("Stop", "i becomes 5; 5 <= 4 is false", "Loop ends, then reclaim exits root."),
        ],
    },
    {
        "slug": "MP2_TEMPERATURE",
        "title": "MP2 - Temperature Spring/Bud/Wither",
        "source": """root() {
    seed temperature = 50;

    // First check: Is it boiling?
    spring (temperature >= 50) {
        plant("Too hot! The water is boiling!");
    } 
        bud (temperature <= 0) {
        plant("Too cold! The water turned into ice!");
    } 
    wither {
        plant("The water is just right for watering.");
    }

    reclaim;
}""",
        "purpose": "This program shows declaration, comment tokenization, spring/bud/wither control flow, comparison operators, and string output.",
        "lexer": [
            f"`root` is recognized at scanner.py line {REF['tok_root']}.",
            f"`seed temperature = 50;` produces seed at line {REF['tok_seed']}, identifier at line {REF['tok_id']}, integer literal at line {REF['tok_int']}, plus assignment and semicolon tokens.",
            f"The comment `// First check: Is it boiling?` becomes a comment token at scanner.py line {REF['tok_comment']}. It appears in the token table but it is skipped later.",
            f"`spring`, `bud`, and `wither` are tokenized at scanner.py lines {REF['tok_spring']}, {REF['tok_bud']}, and {REF['tok_wither']}.",
            f"`>=` is tokenized at line {REF['tok_gte']}; `<=` is tokenized at line {REF['tok_lte']}.",
            f"Each message in plant is one string literal token from scanner.py line {REF['tok_string']}.",
        ],
        "parser": [
            "Parser starts from `<program>` and sees lookahead `root`, so it matches the required root function structure.",
            "It accepts local declaration `seed temperature = 50;`.",
            f"The comment token is skipped by parser.py line {REF['parser_skip']}. The grammar does not need a comment production.",
            "The next meaningful token is `spring`, so parser accepts a spring condition and block.",
            "After the spring block, parser sees `bud`, so it accepts the optional else-if branch.",
            "After bud, parser sees `wither`, so it accepts the else branch.",
            "Finally parser matches required `reclaim;` and closes root.",
        ],
        "builder": [
            f"`parse_variable()` at builder.py line {REF['builder_variable']} creates the temperature declaration.",
            f"Before builder runs, parser filters comments/newlines at parser.py line {REF['parser_filter']}. That is why the comment does not become an AST statement.",
            f"`parse_if()` at builder.py line {REF['builder_if']} builds one IfStatementNode containing spring, bud, and wither parts.",
            "The conditions `temperature >= 50` and `temperature <= 0` become BinaryOpNode expressions.",
        ],
        "semantic": [
            "Semantic validation checks that `temperature` exists before it is used.",
            "It accepts `temperature >= 50` and `temperature <= 0` because comparison expressions return branch/boolean.",
            "The plant statements are valid because they output string literals.",
            "No semantic errors are found.",
        ],
        "interpreter": [
            f"`eval_program()` at interpreter.py line {REF['interpreter_program']} registers root and then calls root.",
            f"`eval_variable_declaration()` at line {REF['interpreter_variable']} stores `temperature = 50`.",
            f"`eval_if_statement()` at line {REF['interpreter_if']} evaluates the spring condition first.",
            f"`temperature >= 50` uses `eval_binary_op()` at line {REF['interpreter_binary']} and the `>=` branch at line {REF['interpreter_gte']}.",
            "Because 50 >= 50 is true, only the spring block runs.",
            f"`eval_print()` at line {REF['interpreter_print']} emits `Too hot! The water is boiling!`.",
            "The bud and wither blocks are skipped because the spring branch already succeeded.",
        ],
        "runtime": [
            ("Declare", "temperature = 50", "Runtime stores seed temperature."),
            ("Spring condition", "50 >= 50 is true", "Spring block is selected."),
            ("Output", "plant message", "`Too hot! The water is boiling!` is emitted."),
            ("Bud", "temperature <= 0", "Not evaluated because spring is already true."),
            ("Wither", "else branch", "Not executed."),
            ("End", "reclaim", "root exits."),
        ],
    },
    {
        "slug": "MP3_COUNT_LOOP",
        "title": "MP3 - Count Loop Odd/Even Output",
        "source": """root() {
\tseed count;
    cultivate ( count = 1; count <= 5; count++) {
        spring (count % 2  == 0) {
            plant("Bloom!");
        } wither {
            plant(count);
        }
    }

    reclaim;
}""",
        "purpose": "This program shows a cultivate loop using an existing variable, a comparison condition, postfix increment, modulo/equality condition, and alternating spring/wither output.",
        "lexer": [
            f"`root` becomes a root reserved word token at scanner.py line {REF['tok_root']}.",
            f"`seed count;` produces seed at line {REF['tok_seed']} and identifier `count` at line {REF['tok_id']}.",
            f"`cultivate` is tokenized at scanner.py line {REF['tok_cultivate']}.",
            f"The loop condition `count <= 5` uses the `<=` token at scanner.py line {REF['tok_lte']}.",
            f"The update `count++` uses the increment token at scanner.py line {REF['tok_inc']}.",
            f"The spring condition uses `%` at scanner.py line {REF['tok_mod']} and `==` at line {REF['tok_eq']}.",
            f"`Bloom!` is tokenized as a string literal at scanner.py line {REF['tok_string']}.",
        ],
        "parser": [
            "Parser matches root and local declaration `seed count;`.",
            "It sees `cultivate` and parses the header as initializer, condition, and update.",
            "Initializer is `count = 1`.",
            "Condition is `count <= 5`.",
            "Update is `count++`.",
            "Inside the loop block, parser accepts spring/wither with condition `count % 2 == 0`.",
            "Root ends with required `reclaim;`.",
        ],
        "builder": [
            f"`parse_variable()` at builder.py line {REF['builder_variable']} builds the `count` declaration.",
            f"`parse_for()` at line {REF['builder_for']} builds the cultivate loop node.",
            f"`parse_if()` at line {REF['builder_if']} builds the spring/wither node inside the loop.",
            "The expression `count % 2 == 0` becomes nested BinaryOpNodes.",
        ],
        "semantic": [
            "Semantic validation checks that `count` is declared before being assigned, compared, incremented, and printed.",
            "The cultivate condition `count <= 5` is valid because it returns branch/boolean.",
            "The spring condition `count % 2 == 0` is valid because `%` produces a number and `==` produces branch/boolean.",
            f"`_check_ForLoop()` at analyzer.py line {REF['semantic_for']} marks loop context. No semantic errors are found.",
        ],
        "interpreter": [
            f"`eval_variable_declaration()` at interpreter.py line {REF['interpreter_variable']} creates `count`. Because it has no initializer, seed default is 0.",
            f"`eval_for_loop()` at line {REF['interpreter_for']} runs the cultivate flow.",
            "The initializer assigns `count = 1` before the first loop check.",
            f"`count <= 5` uses `eval_binary_op()` line {REF['interpreter_binary']} and the `<=` branch at line {REF['interpreter_lte']}.",
            f"`count % 2 == 0` uses modulo branch line {REF['interpreter_mod']} and equality branch line {REF['interpreter_eq']}.",
            f"`eval_if_statement()` at line {REF['interpreter_if']} chooses spring for even values and wither for odd values.",
            f"`count++` uses `eval_unaryop()` at line {REF['interpreter_unary']} after each loop body.",
            f"`eval_print()` at line {REF['interpreter_print']} emits either `Bloom!` or the numeric count.",
        ],
        "runtime": [
            ("Declare", "count = 0", "Default seed value."),
            ("Init", "count = 1", "cultivate initializer."),
            ("Iteration 1", "1 <= 5 true; 1 % 2 == 0 false", "wither prints `1`; count becomes 2."),
            ("Iteration 2", "2 <= 5 true; 2 % 2 == 0 true", "spring prints `Bloom!`; count becomes 3."),
            ("Iteration 3", "3 <= 5 true; 3 % 2 == 0 false", "wither prints `3`; count becomes 4."),
            ("Iteration 4", "4 <= 5 true; 4 % 2 == 0 true", "spring prints `Bloom!`; count becomes 5."),
            ("Iteration 5", "5 <= 5 true; 5 % 2 == 0 false", "wither prints `5`; count becomes 6."),
            ("Stop", "6 <= 5 false", "Loop exits and root reclaims."),
        ],
    },
]


sys.path.insert(0, str(ROOT / "Backend"))
from lexer import lex
from parser import LL1Parser
from cfg.grammar import cfg, first_sets, predict_sets
from semantic import validate_ast
from interpreter.interpreter import Interpreter


class Collector:
    def __init__(self):
        self.outputs = []

    def emit(self, event, data=None):
        if event == "output":
            self.outputs.append((data or {}).get("output", ""))


def ast_tree(node, max_depth=4, depth=0, out=None, limit=120):
    if out is None:
        out = []
    if node is None or len(out) >= limit:
        return out
    node_type = getattr(node, "node_type", type(node).__name__)
    value = getattr(node, "value", None)
    line = getattr(node, "line", "")
    label = "  " * depth + f"- {node_type}"
    if value not in (None, ""):
        label += f" value={value}"
    if line not in (None, ""):
        label += f" line={line}"
    out.append(label)
    if depth < max_depth:
        for child in getattr(node, "children", []):
            ast_tree(child, max_depth, depth + 1, out, limit)
            if len(out) >= limit:
                break
    return out


def run_pipeline(source):
    tokens, lex_errors = lex(source)
    result = {"tokens": tokens, "lex_errors": lex_errors, "token_count": len(tokens)}
    if lex_errors:
        return result
    parser = LL1Parser(cfg, predict_sets, first_sets)
    parse_ok, parse_errors = parser.parse(tokens)
    result["parse_ok"] = parse_ok
    result["parse_errors"] = parse_errors
    parse_result = parser.parse_and_build(tokens)
    result["parse_build"] = parse_result
    if not parse_result.get("success"):
        return result
    result["ast_tree"] = ast_tree(parse_result.get("ast"))
    semantic_result = validate_ast(parse_result["ast"], parse_result["symbol_table"])
    result["semantic"] = semantic_result
    if not semantic_result.get("success"):
        return result
    collector = Collector()
    interpreter = Interpreter(socketio=collector)
    result["runtime_result"] = interpreter.interpret(semantic_result["ast"])
    result["outputs"] = collector.outputs
    result["functions"] = list(interpreter.functions.keys())
    result["final_scopes"] = str(interpreter.scopes)
    return result


def source_line_meaning(line):
    stripped = line.strip()
    if not stripped:
        return "Blank line used only for spacing."
    if stripped.startswith("//"):
        return "Comment. Lexer records it as a comment token; parser/builder skip it."
    if "pollinate" in stripped:
        return "Function declaration header before root."
    if stripped.startswith("root"):
        return "Main function header. Interpreter automatically calls root after registering functions."
    if stripped.startswith("seed"):
        return "Variable declaration. Builder creates a VariableDeclarationNode; interpreter creates runtime variable storage."
    if stripped.startswith("cultivate"):
        return "For-loop style statement: initializer, condition, update, and block."
    if stripped.startswith("spring"):
        return "If branch. Interpreter runs this block when condition is true."
    if stripped.startswith("bud"):
        return "Else-if branch. Interpreter checks it only if spring is false."
    if stripped.startswith("wither"):
        return "Else branch. Interpreter runs it only if all previous conditions are false."
    if stripped.startswith("plant"):
        return "Output statement. Interpreter evaluates arguments and emits text."
    if stripped.startswith("reclaim"):
        return "Return/end statement. Interpreter exits the current function."
    if "=" in stripped:
        return "Assignment or expression statement."
    if stripped in {"}", "};"}:
        return "Block closing brace."
    return "Program structure line scanned by lexer and checked by parser."


def token_groups(tokens):
    groups = {}
    for token in tokens:
        if token.type in ("\n", "EOF"):
            continue
        groups.setdefault(token.line, []).append(f"{token.type}:{token.value}")
    return groups


def add_table_md(md, headers, rows):
    md.append("| " + " | ".join(headers) + " |")
    md.append("|" + "|".join("---" for _ in headers) + "|")
    for row in rows:
        md.append("| " + " | ".join(esc(cell) for cell in row) + " |")


def add_doc_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for i, cell in enumerate(row):
            cells[i].text = str(cell)
    return table


def safe_docx_path(slug):
    base = OUT / f"{slug}_VERY_DETAILED_STEP_BY_STEP_SIMULATION.docx"
    if not base.exists():
        return base
    try:
        base.unlink()
        return base
    except PermissionError:
        return OUT / f"{slug}_VERY_DETAILED_STEP_BY_STEP_SIMULATION_{datetime.now():%H%M%S}.docx"


def write_program_docs(program):
    data = run_pipeline(program["source"])
    md_path = OUT / f"{program['slug']}_VERY_DETAILED_STEP_BY_STEP_SIMULATION.md"
    docx_path = safe_docx_path(program["slug"])

    stage_rows = [
        ("Lexer", f"errors={data.get('lex_errors')}; token_count={data.get('token_count')}"),
        ("Parser", f"parse_ok={data.get('parse_ok')}; errors={data.get('parse_errors')}"),
        ("Builder", f"success={data.get('parse_build', {}).get('success')}; errors={data.get('parse_build', {}).get('errors')}"),
        ("Semantic", f"success={data.get('semantic', {}).get('success')}; errors={data.get('semantic', {}).get('errors')}; warnings={data.get('semantic', {}).get('warnings')}"),
        ("Interpreter", f"outputs={data.get('outputs')}; functions={data.get('functions')}; final_scopes={data.get('final_scopes')}"),
    ]
    line_rows = [
        (line_no, f"`{line}`", source_line_meaning(line))
        for line_no, line in enumerate(program["source"].splitlines(), 1)
    ]
    group_rows = [
        (line_no, "`" + "  |  ".join(values) + "`")
        for line_no, values in sorted(token_groups(data["tokens"]).items())
    ]
    token_rows = [
        (i, token.line, token.col, f"`{token.type}`", f"`{token.value}`")
        for i, token in enumerate(data["tokens"], 1)
    ]
    output_rows = [
        (i, f"`{output}`")
        for i, output in enumerate(data.get("outputs", []), 1)
    ]

    md = [
        f"# {program['title']} - Very Detailed Step-by-Step Simulation",
        "",
        f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}",
        "",
        "## Program Purpose",
        program["purpose"],
        "",
        "## Source Program",
        "```gal",
        program["source"],
        "```",
        "",
        "## 1. Full System Pipeline For This MP",
    ]
    add_table_md(md, ["Stage", "Detailed Explanation"], [
        ("Input", f"Backend `/api/run` receives the whole editor text. `server.py` line {REF['server_json']} reads JSON, and line {REF['server_source']} stores the code in `source_code`."),
        ("Lexer", f"`server.py` line {REF['server_lex']} calls `lex(source_code)`. Lexer setup is in `scanner.py` line {REF['lexer_init']}; `advance()` is line {REF['lexer_advance']}; `make_tokens()` is line {REF['lexer_make']}."),
        ("Parser", f"`parser.parse_and_build(tokens)` is called at `server.py` line {REF['server_parse']}. `parse()` starts at `parser.py` line {REF['parser_parse']}; the stack starts at line {REF['parser_stack']}."),
        ("Builder", f"After syntax succeeds, comments/newlines are filtered at `parser.py` line {REF['parser_filter']}, and `_build_ast(filtered)` is called at line {REF['parser_ast']}. `build_ast()` starts at `builder.py` line {REF['builder_build']}."),
        ("Semantic", f"`validate_ast()` is called at `server.py` line {REF['server_semantic']}. Semantic wrapper is `analyzer.py` line {REF['semantic_validate_ast']}; recursive walk is line {REF['semantic_walk']}."),
        ("Interpreter", f"`Interpreter` is created at `server.py` line {REF['server_interpreter']}; execution starts with `interp.interpret(ast)` at line {REF['server_execute']}. Dispatcher is `interpreter.py` line {REF['interpreter_interpret']}."),
    ])
    md += ["", "## 2. Actual Run Results"]
    add_table_md(md, ["Stage", "Actual Result"], stage_rows)
    md += ["", "## 3. Source Line Meaning"]
    add_table_md(md, ["Line", "Code", "Meaning"], line_rows)
    for key, title in [
        ("lexer", "## 4. Lexer Stage - Detailed Explanation"),
        ("parser", "## 5. Parser Stage - Detailed Explanation"),
        ("builder", "## 6. Builder / AST Stage - Detailed Explanation"),
        ("semantic", "## 7. Semantic Stage - Detailed Explanation"),
        ("interpreter", "## 8. Interpreter Stage - Detailed Explanation"),
    ]:
        md += ["", title]
        for i, item in enumerate(program[key], 1):
            md.append(f"{i}. {item}")
    md += ["", "## 9. Runtime Trace"]
    add_table_md(md, ["Moment", "State / Expression", "Result"], program["runtime"])
    md += ["", "## 10. Tokens By Source Line"]
    add_table_md(md, ["Source Line", "Tokens Produced"], group_rows)
    md += ["", "## 11. AST Shape Summary", "```text"]
    md.extend(data.get("ast_tree", []))
    md += ["```", "", "## 12. Final Output"]
    add_table_md(md, ["#", "Output"], output_rows)
    md += ["", "## 13. Full Token Table"]
    add_table_md(md, ["#", "Line", "Col", "Token Type", "Value"], token_rows)
    md += ["", "## 14. Important Code Snippets"]
    for title, file_key, start, end, note in [
        ("Server Pipeline", "server", REF["server_route"], REF["server_execute"] + 8, "Controls the whole compile/run flow."),
        ("Lexer `current_char` Flow", "scanner", REF["lexer_init"], REF["lexer_loop"] + 4, "Shows source storage, position, current_char, advance, and token loop."),
        ("Parser Stack Flow", "parser", REF["parser_parse"], REF["parser_match"] + 5, "Shows LL(1) stack, lookahead, production choice, and terminal matching."),
        ("Builder Entry", "builder", REF["builder_build"], REF["builder_loop"] + 8, "Shows AST root creation and token-to-AST loop."),
        ("Semantic Walk", "semantic", REF["semantic_validate"], REF["semantic_handler"] + 8, "Shows final AST validation logic."),
        ("Interpreter Program Entry", "interpreter", REF["interpreter_interpret"], REF["interpreter_program"] + 18, "Shows dispatcher and automatic root call."),
    ]:
        md += ["", f"### {title}", f"Lines {start}-{end} in `{FILES[file_key].relative_to(ROOT)}`", "```python", snippet(file_key, start, end), "```", note]
    md_path.write_text("\n".join(md), encoding="utf-8")

    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    for style in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)

    heading = doc.add_heading(f"{program['title']} - Very Detailed Step-by-Step Simulation", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.add_heading("Program Purpose", 1)
    doc.add_paragraph(program["purpose"])
    doc.add_heading("Source Program", 1)
    para = doc.add_paragraph()
    run = para.add_run(program["source"])
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    doc.add_heading("1. Full System Pipeline For This MP", 1)
    add_doc_table(doc, ["Stage", "Detailed Explanation"], [
        ("Input", f"Backend /api/run receives the whole editor text. server.py line {REF['server_json']} reads JSON, and line {REF['server_source']} stores the code in source_code."),
        ("Lexer", f"server.py line {REF['server_lex']} calls lex(source_code). Lexer setup is scanner.py line {REF['lexer_init']}; advance() is line {REF['lexer_advance']}; make_tokens() is line {REF['lexer_make']}."),
        ("Parser", f"parser.parse_and_build(tokens) is called at server.py line {REF['server_parse']}. parse() starts at parser.py line {REF['parser_parse']}; stack starts at line {REF['parser_stack']}."),
        ("Builder", f"After syntax succeeds, comments/newlines are filtered at parser.py line {REF['parser_filter']}, and _build_ast(filtered) is called at line {REF['parser_ast']}. build_ast() starts at builder.py line {REF['builder_build']}."),
        ("Semantic", f"validate_ast() is called at server.py line {REF['server_semantic']}. Semantic wrapper is analyzer.py line {REF['semantic_validate_ast']}; recursive walk is line {REF['semantic_walk']}."),
        ("Interpreter", f"Interpreter is created at server.py line {REF['server_interpreter']}; execution starts with interp.interpret(ast) at line {REF['server_execute']}. Dispatcher is interpreter.py line {REF['interpreter_interpret']}."),
    ])
    doc.add_heading("2. Actual Run Results", 1)
    add_doc_table(doc, ["Stage", "Actual Result"], stage_rows)
    doc.add_heading("3. Source Line Meaning", 1)
    add_doc_table(doc, ["Line", "Code", "Meaning"], line_rows)
    for key, title in [
        ("lexer", "4. Lexer Stage - Detailed Explanation"),
        ("parser", "5. Parser Stage - Detailed Explanation"),
        ("builder", "6. Builder / AST Stage - Detailed Explanation"),
        ("semantic", "7. Semantic Stage - Detailed Explanation"),
        ("interpreter", "8. Interpreter Stage - Detailed Explanation"),
    ]:
        doc.add_heading(title, 1)
        for i, item in enumerate(program[key], 1):
            doc.add_paragraph(f"{i}. {item}")
    doc.add_heading("9. Runtime Trace", 1)
    add_doc_table(doc, ["Moment", "State / Expression", "Result"], program["runtime"])
    doc.add_heading("10. Tokens By Source Line", 1)
    add_doc_table(doc, ["Source Line", "Tokens Produced"], group_rows)
    doc.add_heading("11. AST Shape Summary", 1)
    para = doc.add_paragraph()
    run = para.add_run("\n".join(data.get("ast_tree", [])))
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    doc.add_heading("12. Final Output", 1)
    add_doc_table(doc, ["#", "Output"], output_rows)
    doc.add_heading("13. Full Token Table", 1)
    add_doc_table(doc, ["#", "Line", "Col", "Token Type", "Value"], token_rows)
    doc.add_heading("14. Important Code Snippets", 1)
    for title, file_key, start, end, note in [
        ("Server Pipeline", "server", REF["server_route"], REF["server_execute"] + 8, "Controls the whole compile/run flow."),
        ("Lexer current_char Flow", "scanner", REF["lexer_init"], REF["lexer_loop"] + 4, "Shows source storage, position, current_char, advance, and token loop."),
        ("Parser Stack Flow", "parser", REF["parser_parse"], REF["parser_match"] + 5, "Shows LL(1) stack, lookahead, production choice, and terminal matching."),
        ("Builder Entry", "builder", REF["builder_build"], REF["builder_loop"] + 8, "Shows AST root creation and token-to-AST loop."),
        ("Semantic Walk", "semantic", REF["semantic_validate"], REF["semantic_handler"] + 8, "Shows final AST validation logic."),
        ("Interpreter Program Entry", "interpreter", REF["interpreter_interpret"], REF["interpreter_program"] + 18, "Shows dispatcher and automatic root call."),
    ]:
        doc.add_heading(title, 2)
        doc.add_paragraph(f"Lines {start}-{end} in {FILES[file_key].relative_to(ROOT)}")
        para = doc.add_paragraph()
        run = para.add_run(snippet(file_key, start, end))
        run.font.name = "Consolas"
        run.font.size = Pt(7)
        doc.add_paragraph(note)
    doc.save(docx_path)
    return docx_path, md_path, data


def main():
    outputs = []
    for program in PROGRAMS:
        outputs.append((program["title"],) + write_program_docs(program))
    for title, docx_path, md_path, data in outputs:
        print(title)
        print("  DOCX:", docx_path)
        print("  MD:", md_path)
        print("  tokens:", data.get("token_count"))
        print("  outputs:", data.get("outputs"))


if __name__ == "__main__":
    main()
