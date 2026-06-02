from pathlib import Path
from datetime import datetime
import re

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Docus & mps"
SOURCE_REL = "Backend/interpreter/interpreter.py"
SOURCE_PATH = ROOT / SOURCE_REL

MD_PATH = OUT_DIR / "INTERPRETER_LINE_BY_LINE_EXPLANATION.md"
DOCX_PATH = OUT_DIR / "INTERPRETER_LINE_BY_LINE_EXPLANATION.docx"
PDF_PATH = OUT_DIR / "INTERPRETER_LINE_BY_LINE_EXPLANATION.pdf"


FUNCTION_PURPOSE = {
    "__init__": "Initializes the interpreter runtime memory: output channel, loop flags, input handling, variable scopes, function table, and bundle table.",
    "declare_variable": "Creates a variable record in the current runtime scope.",
    "lookup_variable": "Searches the active scopes and variable table for a variable name.",
    "set_variable": "Updates an existing variable value inside the nearest active scope.",
    "declare_function": "Stores a function declaration so it can be called later.",
    "lookup_function": "Finds a stored function declaration by name.",
    "enter_scope": "Creates a new local scope for functions, blocks, loops, or conditionals.",
    "exit_scope": "Removes the current local scope after leaving a function/block.",
    "interpret": "Main dispatcher. It receives any AST node and routes it to the correct evaluator method.",
    "eval_program": "Executes the whole program node: registers top-level declarations and automatically calls root().",
    "eval_variable_declaration": "Executes variable declarations and stores default or evaluated initial values.",
    "eval_bundle_definition": "Registers bundle/struct member definitions.",
    "_build_bundle_defaults": "Creates default runtime values for all members of a bundle type.",
    "eval_member_access": "Reads a bundle member value like student.age.",
    "eval_array_member_access": "Reads a member from a bundle stored inside an array/list.",
    "eval_sturdy_declaration": "Executes fertile/constant declarations.",
    "eval_assignment": "Executes assignments, including normal variables, lists, bundle members, and input assignments.",
    "eval_binary_op": "Evaluates binary expressions such as +, -, *, /, %, ==, &&, and ||.",
    "_parse_literal": "Converts literal text or token values into Python runtime values.",
    "eval_function_declaration": "Stores function declarations into the interpreter function table.",
    "eval_block": "Executes every statement inside a block one by one.",
    "plant": "Sends output text to the frontend through Socket.IO or the output collector.",
    "plant_out": "Sends and stores output text for collector-style execution.",
    "eval_print": "Executes plant() statements, including placeholder formatting and multi-argument printing.",
    "eval_formatted_string": "Removes string quotes and converts escape sequences such as \\n and \\t.",
    "eval_list": "Evaluates list/array literal values.",
    "eval_list_access": "Reads a value from a list or string using an index.",
    "eval_return": "Executes reclaim by raising ReturnValue to stop the current function.",
    "eval_function_call": "Executes a function call: evaluates arguments, creates scope, binds parameters, runs body, and catches reclaim.",
    "eval_append": "Executes list append behavior.",
    "eval_insert": "Executes list insertion behavior.",
    "eval_remove": "Executes list removal behavior.",
    "eval_unaryop": "Executes unary operators like ++, --, !, and ~.",
    "eval_cast": "Converts values to a requested GAL type.",
    "eval_soil": "Converts a string variable to lowercase.",
    "eval_bloom": "Converts a string variable to uppercase.",
    "eval_if_statement": "Executes spring/bud/wither conditional logic.",
    "eval_for_loop": "Executes cultivate loops.",
    "eval_while_loop": "Executes grow loops.",
    "eval_do_while_loop": "Executes tend/wither style do-while loops.",
    "eval_break": "Executes prune by setting the break flag.",
    "trigger_break": "Turns on the runtime break flag.",
    "break_triggered": "Checks if a break/prune was triggered.",
    "enter_loop": "Records that the interpreter is currently inside a loop.",
    "exit_loop": "Removes the current loop marker and resets loop flags.",
    "eval_continue": "Executes skip by setting the continue flag.",
    "continue_triggered": "Checks if skip/continue was triggered.",
    "trigger_continue": "Turns on the runtime continue flag.",
    "eval_switch": "Executes switch/variety style branching.",
    "emit_input_request": "Asks the frontend UI for input when water() is executed.",
    "provide_input": "Receives user input from the frontend and wakes the waiting interpreter.",
    "wait_for_input": "Pauses execution until the frontend supplies an input value.",
    "eval_input": "Executes water(), determines expected type, waits for input, and converts it.",
}


STATE_EXPLANATIONS = {
    "self.output": "runtime output list used by collector-style execution",
    "self.socketio": "frontend communication object used for plant() and water() events",
    "self.loop_stack": "stack/list that records active loops for prune and skip validation",
    "self.break_flag": "flag set when prune/break is triggered",
    "self.continue_flag": "flag set when skip/continue is triggered",
    "self.input_required": "flag showing the program is waiting for water() input",
    "self.input_events": "dictionary of waiting input events by variable name",
    "self.input_values": "dictionary of input values already received from the frontend",
    "self.current_node": "optional tracker for the current AST node",
    "self.current_parent": "optional tracker for the current parent AST node",
    "self.variables": "extra variable table used for global/runtime variables",
    "self.global_variables": "table for global variables",
    "self.functions": "function table storing pollinate/root declarations",
    "self.scopes": "scope stack; each scope is a dictionary of variable records",
    "self.current_func_name": "name of the function currently being executed",
    "self.function_variables": "per-function variable helper storage",
    "self.bundle_types": "bundle/struct type definition table",
}


def read_source():
    return SOURCE_PATH.read_text(encoding="utf-8", errors="replace").splitlines()


def extract_function_name(line):
    match = re.match(r"\s*def\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(", line)
    return match.group(1) if match else None


def extract_blocks(lines):
    starts = []
    if lines:
        starts.append((1, "Module Setup"))
    for idx, line in enumerate(lines, start=1):
        name = extract_function_name(line)
        if name:
            starts.append((idx, name))
        elif re.match(r"\s*class\s+Interpreter", line):
            starts.append((idx, "Interpreter Class"))

    starts = sorted(set(starts))
    blocks = []
    for i, (start, name) in enumerate(starts):
        end = starts[i + 1][0] - 1 if i + 1 < len(starts) else len(lines)
        if start <= end:
            blocks.append((name, start, end))
    return blocks


def escape_md(text):
    return text.replace("|", "\\|").replace("\n", "<br>")


def table(headers, rows):
    out = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join(["---"] * len(headers)) + "|",
    ]
    for row in rows:
        out.append("| " + " | ".join(escape_md(str(col)) for col in row) + " |")
    return "\n".join(out)


def code_line_meaning(line, current_function=None):
    stripped = line.strip()
    if stripped == "":
        return "Blank line used to separate code sections for readability."
    if stripped.startswith("#"):
        return "Comment that documents the purpose of the nearby code."
    if stripped.startswith("from ") or stripped.startswith("import "):
        return "Imports modules/classes needed by the interpreter."
    if stripped.startswith("sys.setrecursionlimit"):
        return "Raises Python recursion limit so recursive GAL functions can run deeper before hitting Python's default limit."
    if stripped.startswith("try:"):
        return "Starts a protected block where errors can be handled by except/finally."
    if stripped.startswith("except ImportError"):
        return "Handles the case where eventlet is not installed."
    if stripped.startswith("except ReturnValue"):
        return "Catches a reclaim return value from a called function."
    if stripped.startswith("except"):
        return "Handles an error or special control-flow case."
    if stripped.startswith("finally:"):
        return "Cleanup block that runs even if the try block returns or raises an error."
    if stripped.startswith("class Interpreter"):
        return "Defines the Interpreter class, the runtime engine that executes AST nodes."
    func_name = extract_function_name(line)
    if func_name:
        return FUNCTION_PURPOSE.get(func_name, f"Defines interpreter method {func_name}().")
    if stripped.startswith("if isinstance(node,"):
        return "Dispatcher check: if the current AST node has this class, route it to the matching eval method."
    if stripped.startswith("elif isinstance(node,"):
        return "Next dispatcher check for another possible AST node class."
    if stripped.startswith("elif node.node_type"):
        return "Dispatcher check for AST nodes identified by their node_type string."
    if stripped.startswith("if ") or stripped.startswith("elif "):
        if "node.children" in stripped:
            return "Checks a condition using values stored inside the AST node children."
        if "var_type" in stripped:
            return "Checks the declared GAL type so the runtime can validate or convert the value."
        if "operator ==" in stripped:
            return "Chooses which binary/unary operator logic to execute."
        return "Conditional branch that decides which runtime path should run."
    if stripped.startswith("else:"):
        return "Fallback branch used when the previous if/elif conditions were false."
    if stripped.startswith("for "):
        if "node.children" in stripped:
            return "Loops through child AST nodes and processes each one in order."
        return "Loops through a collection of values/items."
    if stripped.startswith("while "):
        return "Repeats a runtime action while the condition remains true."
    if stripped.startswith("return self."):
        method_match = re.search(r"self\.([A-Za-z_][A-Za-z0-9_]*)", stripped)
        method_name = method_match.group(1) if method_match else "method"
        return f"Calls {method_name}() on the same interpreter object and returns its result to the caller."
    if stripped.startswith("return "):
        return "Returns the computed value/result from this method to its caller."
    if stripped == "return":
        return "Stops this method without returning a meaningful value."
    if stripped.startswith("raise ReturnValue"):
        return "Implements reclaim: raises a special return object so function execution stops immediately."
    if stripped.startswith("raise InterpreterError"):
        return "Creates a runtime error message and stops execution for an invalid runtime situation."
    if stripped.startswith("raise SemanticError"):
        return "Raises a semantic-style error from runtime validation logic."
    if ".append(" in stripped:
        return "Adds a new item to a list, such as output, parameters, arguments, or evaluated list values."
    if ".pop(" in stripped:
        return "Removes and returns an item from a list/dictionary, often during cleanup."
    if ".emit(" in stripped:
        return "Sends an event to the frontend/UI, such as output text or an input request."
    if "self.enter_scope()" in stripped:
        return "Creates a new local runtime scope."
    if "self.exit_scope()" in stripped:
        return "Leaves/removes the current local runtime scope."
    if "self.enter_loop" in stripped:
        return "Marks that execution is now inside a loop."
    if "self.exit_loop" in stripped:
        return "Leaves the current loop and resets loop control flags."
    if "self.lookup_variable" in stripped:
        return "Finds a variable record in the active runtime scopes."
    if "self.set_variable" in stripped:
        return "Updates a variable's stored runtime value."
    if "self.declare_variable" in stripped:
        return "Creates a variable record in the current scope."
    if "self.declare_function" in stripped:
        return "Registers a function in the interpreter's function table."
    if "self.lookup_function" in stripped:
        return "Finds a function record in the function table."
    if "self.interpret(" in stripped:
        return "Recursively sends another AST node back to the interpreter dispatcher."
    if "node.children" in stripped and "=" in stripped:
        return "Reads a specific child from the AST node. Child indexes represent parts such as type, name, condition, target, or value."
    if "node.value" in stripped and "=" in stripped:
        return "Reads the stored value of the AST node, such as an operator, identifier name, or literal text."
    if " = {" in stripped or stripped.endswith("= {}"):
        return "Creates a dictionary used for key-value runtime storage."
    if " = [" in stripped or stripped.endswith("= []"):
        return "Creates a list used to collect runtime values."
    for state_name, state_meaning in STATE_EXPLANATIONS.items():
        if stripped.startswith(state_name):
            return f"Stores/updates {state_meaning}."
    if "=" in stripped:
        left_side = stripped.split("=", 1)[0].strip()
        return f"Assigns/computes a value and stores it in `{left_side}` for later use in this method."
    if stripped.endswith(")") or stripped.endswith(");"):
        return "Calls a function/method to perform part of the runtime operation."
    return "Runtime support statement used by the interpreter."


def block_flow(name):
    if name == "Module Setup":
        return [
            "Imports AST node classes and runtime error classes.",
            "Configures recursion and eventlet input handling.",
            "Prepares the Interpreter class dependencies.",
        ]
    if name == "Interpreter Class":
        return [
            "Defines the runtime object that will hold all interpreter state.",
            "All eval methods below belong to this class.",
        ]
    if name == "interpret":
        return [
            "Receive an AST node.",
            "Check its class or node_type.",
            "Call the correct eval method.",
            "Return the eval method result.",
        ]
    if name == "eval_program":
        return [
            "Loop through top-level AST children.",
            "Register declarations/functions.",
            "Create a FunctionCallNode for root().",
            "Interpret that root call to start program execution.",
        ]
    if name == "eval_function_call":
        return [
            "Evaluate call arguments.",
            "Look up the called function.",
            "Create a new scope.",
            "Bind parameters to argument values.",
            "Execute the function body block.",
            "Catch ReturnValue from reclaim and return it.",
        ]
    if name == "eval_assignment":
        return [
            "Evaluate the right-hand value.",
            "Find the target variable/list/member.",
            "Convert value when needed based on GAL type.",
            "Store the final value.",
        ]
    if name == "eval_binary_op":
        return [
            "Evaluate left and right operands.",
            "Read the operator.",
            "Run the matching operator branch.",
            "Return the computed result.",
        ]
    if name == "eval_input":
        return [
            "Find where water() is being used.",
            "Determine target variable name and expected GAL type.",
            "Ask frontend for input.",
            "Wait for the input value.",
            "Convert the input string into the expected runtime type.",
            "Return the converted input.",
        ]
    if name == "eval_print":
        return [
            "Evaluate plant() arguments.",
            "Use placeholder formatting if the first string has {}.",
            "Otherwise join multiple arguments with spaces.",
            "Emit the final output text.",
        ]
    if name == "eval_if_statement":
        return [
            "Evaluate the spring condition.",
            "If true, run the spring block.",
            "If false, check bud blocks.",
            "If no bud is true, run wither block if present.",
            "Exit conditional scopes after execution.",
        ]
    if name in {"eval_for_loop", "eval_while_loop", "eval_do_while_loop"}:
        return [
            "Enter loop tracking.",
            "Evaluate loop condition/control values.",
            "Execute block while rules allow.",
            "Handle skip/prune flags.",
            "Exit loop tracking and scope.",
        ]
    if name == "eval_return":
        return [
            "Evaluate the reclaim expression if one exists.",
            "Raise ReturnValue so the current function stops immediately.",
        ]
    return [FUNCTION_PURPOSE.get(name, "Executes/supports part of the interpreter runtime.")]


def make_line_table(lines, start, end, current_function):
    rows = []
    for line_no in range(start, end + 1):
        code_text = lines[line_no - 1]
        explanation = code_line_meaning(code_text, current_function)
        rows.append([line_no, f"`{escape_md(code_text)}`", explanation])
    return table(["Line", "Code", "Explanation"], rows)


def make_code_block(lines, start, end):
    body = "\n".join(f"{idx:04d}: {lines[idx - 1]}" for idx in range(start, end + 1))
    return f"```python\n{body}\n```"


def build_markdown():
    lines = read_source()
    blocks = extract_blocks(lines)

    parts = []
    parts.append("# Interpreter Line-by-Line Code Explanation")
    parts.append(
        "Detailed explanation of `Backend/interpreter/interpreter.py`, organized for defense study.\n\n"
        f"Project path: `{ROOT}`\n\n"
        f"Source file: `{SOURCE_REL}`\n\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    parts.append("## 1. How To Study The Interpreter")
    parts.append(
        "The interpreter executes the AST after lexer, parser, builder, and semantic analyzer have already succeeded. "
        "The most important method is `interpret()`, because it is the dispatcher that receives an AST node and sends it to the correct evaluator.\n\n"
        "General runtime flow:\n\n"
        "```text\n"
        "interp.interpret(ast)\n"
        "-> interpret(ProgramNode)\n"
        "-> eval_program()\n"
        "-> register declarations/functions\n"
        "-> create and call root()\n"
        "-> eval_function_call(root)\n"
        "-> eval_block(root body)\n"
        "-> each statement goes back to interpret()\n"
        "-> matching eval method executes the statement\n"
        "```"
    )

    parts.append("## 2. Python Concepts Used In This File")
    parts.append(
        table(
            ["Concept", "Meaning In This Interpreter"],
            [
                ["`self`", "The current Interpreter object. It lets a method access shared runtime state like `self.scopes`, `self.functions`, and `self.socketio`."],
                ["`node`", "The current AST node being executed."],
                ["`node.children`", "List of child AST nodes. Example: assignment has target child and value child."],
                ["`node.value`", "Stored value of a node, such as an identifier name, operator, or literal."],
                ["`isinstance(node, X)`", "Checks what kind of AST node is being interpreted."],
                ["`dict {}`", "Key-value storage, used for variables, functions, scopes, and bundles."],
                ["`list []`", "Ordered collection, used for output, scopes, parameters, arguments, and loop stack."],
                ["`raise ReturnValue`", "Special control flow for `reclaim`."],
                ["`try/finally`", "Ensures cleanup such as leaving scopes and loops."],
            ],
        )
    )

    parts.append("## 3. Function Map")
    map_rows = []
    for name, start, end in blocks:
        if name in {"Module Setup", "Interpreter Class"}:
            purpose = "File/class setup."
        else:
            purpose = FUNCTION_PURPOSE.get(name, "Runtime helper/evaluator.")
        map_rows.append([name, f"{start}-{end}", purpose])
    parts.append(table(["Section/Function", "Lines", "Purpose"], map_rows))

    parts.append("## 4. Example Program Simulation")
    sample = """root() {
    seed a;
    seed b;
    seed sum;

    plant("Enter 1st number: ");
    water(a);

    plant("Enter 2nd number: ");
    water(b);

    sum = a + b;

    plant("Sum: {}", sum);
    reclaim;
}"""
    parts.append("Example input values: `a = 5`, `b = 7`.\n\n")
    parts.append(f"```gal\n{sample}\n```")
    parts.append(
        table(
            ["Step", "Interpreter route", "Runtime effect"],
            [
                ["1", "`interpret(ProgramNode)` -> `eval_program()`", "Registers root and creates a call to `root()`."],
                ["2", "`eval_function_call(root)`", "Creates root scope and executes root block."],
                ["3", "`eval_variable_declaration()` for `seed a;`", "Stores `a = 0`."],
                ["4", "`eval_variable_declaration()` for `seed b;`", "Stores `b = 0`."],
                ["5", "`eval_variable_declaration()` for `seed sum;`", "Stores `sum = 0`."],
                ["6", "`eval_print()`", "Outputs `Enter 1st number:`."],
                ["7", "`eval_input()` + `eval_assignment()`", "Receives input `5`, converts to seed, stores `a = 5`."],
                ["8", "`eval_print()`", "Outputs `Enter 2nd number:`."],
                ["9", "`eval_input()` + `eval_assignment()`", "Receives input `7`, converts to seed, stores `b = 7`."],
                ["10", "`eval_assignment()` + `eval_binary_op()`", "Computes `a + b`, so `5 + 7 = 12`, then stores `sum = 12`."],
                ["11", "`eval_print()`", "Outputs `Sum: 12`."],
                ["12", "`eval_return()`", "`reclaim;` stops root and ends the program."],
            ],
        )
    )
    parts.append(
        "Variable table during the sample:\n\n"
        + table(
            ["Moment", "a", "b", "sum"],
            [
                ["After declarations", "0", "0", "0"],
                ["After `water(a)`", "5", "0", "0"],
                ["After `water(b)`", "5", "7", "0"],
                ["After `sum = a + b`", "5", "7", "12"],
            ],
        )
    )

    parts.append("## 5. Line-By-Line Explanation Of interpreter.py")
    for name, start, end in blocks:
        parts.append(f"### {name} - Lines {start}-{end}")
        parts.append("Purpose: " + (FUNCTION_PURPOSE.get(name, "File/class setup or runtime support block.")))
        parts.append("Process flow:")
        for step in block_flow(name):
            parts.append(f"- {step}")
        parts.append("Code:")
        parts.append(make_code_block(lines, start, end))
        parts.append("Line-by-line notes:")
        parts.append(make_line_table(lines, start, end, name))

    parts.append("## 6. Defense Shortcut")
    parts.append(
        "Use this if you need a short spoken explanation:\n\n"
        "```text\n"
        "Sa interpreter, ang pinaka-start is interpret(ast). Usually ProgramNode muna ang node, kaya papasok siya sa eval_program. "
        "Doon nireregister muna ang top-level declarations and functions, then automatic niyang tatawagin ang root(). "
        "Pag root function call na, eval_function_call ang bahala gumawa ng scope, mag-bind ng parameters kung meron, at patakbuhin ang body gamit eval_block. "
        "Sa eval_block, every statement pinapasa ulit sa interpret(), kaya automatic siyang napupunta sa eval_variable_declaration, eval_assignment, eval_print, eval_input, eval_if_statement, eval_binary_op, or eval_return depende sa AST node type. "
        "Basically, interpret() ang dispatcher, and bawat eval_* method ang actual executor ng specific language feature.\n"
        "```"
    )

    return "\n\n".join(parts)


def clean(text):
    return text.replace("`", "")


def set_cell_shading(cell, fill):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    except Exception:
        pass


def build_docx(markdown):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 3"].font.size = Pt(10)

    def add_code_block(text):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        for line in text.splitlines():
            run = paragraph.add_run(line + "\n")
            run.font.name = "Consolas"
            run.font.size = Pt(6.5)

    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block("\n".join(code_lines))
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1

            rows = []
            for row_idx, table_line in enumerate(table_lines):
                if row_idx == 1:
                    continue
                rows.append([clean(cell.strip().replace("<br>", "\n")) for cell in table_line.strip("|").split("|")])

            if rows:
                max_cols = max(len(row) for row in rows)
                tbl = doc.add_table(rows=len(rows), cols=max_cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(max_cols):
                        cell = tbl.cell(r_idx, c_idx)
                        text = row[c_idx] if c_idx < len(row) else ""
                        cell.text = text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(6.2 if max_cols >= 3 else 7.5)
                                if c_idx == 1 and max_cols >= 3:
                                    run.font.name = "Consolas"
                        if r_idx == 0:
                            set_cell_shading(cell, "D9EAD3")
                doc.add_paragraph()
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = clean(line[level:].strip())
            if level == 1:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                doc.add_heading(text, 1)
            elif level == 3:
                doc.add_heading(text, 2)
            else:
                doc.add_heading(text, 3)
            i += 1
            continue

        if line.startswith("- "):
            doc.add_paragraph(clean(line[2:].strip()), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
            i += 1
            continue

        para = doc.add_paragraph(clean(line.strip()))
        para.paragraph_format.space_after = Pt(3)
        i += 1

    doc.save(DOCX_PATH)


def export_pdf():
    try:
        import win32com.client

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DOCX_PATH.resolve()))
        doc.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return True, str(PDF_PATH)
    except Exception as exc:
        return False, str(exc)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = build_markdown()
    MD_PATH.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    ok, detail = export_pdf()
    print(f"markdown={MD_PATH}")
    print(f"docx={DOCX_PATH}")
    print(f"pdf={'ok ' + detail if ok else 'failed ' + detail}")
