from pathlib import Path
from datetime import datetime
import re

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Docus & mps"
MD_PATH = OUT_DIR / "SUM_PROGRAM_INTERPRETER_WALKTHROUGH.md"
DOCX_PATH = OUT_DIR / "SUM_PROGRAM_INTERPRETER_WALKTHROUGH.docx"
PDF_PATH = OUT_DIR / "SUM_PROGRAM_INTERPRETER_WALKTHROUGH.pdf"


PROGRAM = """root() {
    seed a;
    seed b;
    seed sum;

    plant("Enter 1st number: ");
    water(a);

    plant("Enter 2nd number: ");
    water(b);

    sum = a + b;

    plant("Sum: ", sum);
    reclaim;
}"""


def table(headers, rows):
    out = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join(["---"] * len(headers)) + "|",
    ]
    for row in rows:
        out.append("| " + " | ".join(str(col).replace("\n", "<br>") for col in row) + " |")
    return "\n".join(out)


def build_markdown():
    parts = []
    parts.append("# Exact Interpreter Walkthrough Pattern")
    parts.append(
        "This document explains the exact step-by-step interpreter simulation pattern for the simple sum program.\n\n"
        f"Project path: `{ROOT}`\n\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    parts.append("## Program Example")
    parts.append(f"```gal\n{PROGRAM}\n```")
    parts.append("Sample input used in this walkthrough: `a = 5`, `b = 7`.")

    parts.append("## Short Rule")
    parts.append(
        "For a complete program, the interpreter always starts at `interpret(ProgramNode)`. "
        "Then it goes to `eval_program()`. `eval_program()` first stores top-level functions, then creates a manual call to `root()`."
    )

    parts.append("## Exact Walking Pattern With Lines")
    parts.append(
        table(
            ["Step", "File / Line", "Code Route", "What Happens"],
            [
                ["1", "`interpreter.py:121`", "`def interpret(self, node):`", "Start here. The first node is the whole AST, a `ProgramNode`."],
                ["2", "`interpreter.py:122-123`", "`if isinstance(node, ProgramNode): return self.eval_program(node)`", "Because the node is `ProgramNode`, the interpreter goes to `eval_program()`."],
                ["3", "`interpreter.py:210`", "`def eval_program(self, node):`", "You are now inside the evaluator for the whole program."],
                ["4", "`interpreter.py:211-212`", "`for child in node.children: self.interpret(child)`", "The program has one top-level child: `FunctionDeclarationNode root`. The interpreter sends that child back to `interpret()`."],
                ["5", "`interpreter.py:121`", "`interpret(child)`", "Now the current node is the root function declaration."],
                ["6", "`interpreter.py:140-141`", "`elif isinstance(node, FunctionDeclarationNode): return self.eval_function_declaration(node)`", "Because root is a function declaration, it goes to `eval_function_declaration()`."],
                ["7", "`interpreter.py:716`", "`def eval_function_declaration(self, node):`", "Root is now being registered/saved, not executed yet."],
                ["8", "`interpreter.py:717-719`", "`return_type = ...; parameters_node = ...; func_name = ...`", "Reads root information. `func_name = 'root'`, parameters are empty, return type is root/empty type."],
                ["9", "`interpreter.py:721-729`", "`params = []` and parameter loop", "Because `root()` has no parameters, `params` stays empty."],
                ["10", "`interpreter.py:731`", "`self.declare_function(func_name, return_type, params, node)`", "Calls helper method to save the root function."],
                ["11", "`interpreter.py:95-98`", "`self.functions[name] = {...}`", "Root is saved in `self.functions['root']`. The saved node contains the root body."],
                ["12", "`interpreter.py:733`", "`return None`", "`eval_function_declaration()` ends and returns to the caller."],
                ["13", "`interpreter.py:212`", "`self.interpret(child)` finishes", "The root declaration child is done. Control returns to `eval_program()`."],
                ["14", "`interpreter.py:214`", "`main_call = FunctionCallNode('root', [], node.line)`", "Creates a fake/manual AST node that means `root()` should be called."],
                ["15", "`interpreter.py:215`", "`return self.interpret(main_call)`", "Sends the root call back to `interpret()`."],
                ["16", "`interpreter.py:121`", "`interpret(main_call)`", "Now the current node is `FunctionCallNode root`."],
                ["17", "`interpreter.py:150-151`", "`elif isinstance(node, FunctionCallNode): return self.eval_function_call(node)`", "Because this is a function call, go to `eval_function_call()`."],
                ["18", "`interpreter.py:856`", "`def eval_function_call(self, node):`", "Root will now actually run."],
                ["19", "`interpreter.py:857-858`", "`function_name = node.value; args = [...]`", "For root, `function_name = 'root'` and `args = []`."],
                ["20", "`interpreter.py:860-865`", "`lookup_function`, `expected_params`, `function_node`", "Finds the saved root function in `self.functions` and gets its body node."],
                ["21", "`interpreter.py:873`", "`self.enter_scope()`", "Creates a new runtime scope for variables inside root."],
                ["22", "`interpreter.py:883-884`", "`self.eval_block(function_node.children[2])`", "Starts executing the root body block."],
                ["23", "`interpreter.py:735-737`", "`for statement in block_node.children: self.interpret(statement)`", "Runs each root statement one by one."],
            ],
        )
    )

    parts.append("## Root Body Statement Pattern")
    parts.append(
        table(
            ["Step", "GAL Statement", "Interpreter Route", "Runtime Effect"],
            [
                ["24", "`seed a;`", "`interpret(VariableDeclarationNode)` -> `eval_variable_declaration()` line `218`", "Declares `a` with default seed value `0`."],
                ["25", "`seed b;`", "`interpret(VariableDeclarationNode)` -> `eval_variable_declaration()`", "Declares `b = 0`."],
                ["26", "`seed sum;`", "`interpret(VariableDeclarationNode)` -> `eval_variable_declaration()`", "Declares `sum = 0`."],
                ["27", "`plant(\"Enter 1st number: \");`", "`interpret(PrintNode)` -> `eval_print()` line `752`", "Outputs `Enter 1st number:`."],
                ["28", "`water(a);`", "`interpret(AssignmentNode)` -> `eval_assignment()` line `346` -> `eval_input()` line `1348`", "Asks input for `a`; if user enters `5`, stores `a = 5`."],
                ["29", "`plant(\"Enter 2nd number: \");`", "`interpret(PrintNode)` -> `eval_print()`", "Outputs `Enter 2nd number:`."],
                ["30", "`water(b);`", "`interpret(AssignmentNode)` -> `eval_assignment()` -> `eval_input()`", "Asks input for `b`; if user enters `7`, stores `b = 7`."],
                ["31", "`sum = a + b;`", "`interpret(AssignmentNode)` -> `eval_assignment()` -> `interpret(BinaryOpNode)` -> `eval_binary_op()` line `510`", "Looks up `a = 5`, `b = 7`, computes `5 + 7 = 12`, stores `sum = 12`."],
                ["32", "`plant(\"Sum: \", sum);`", "`interpret(PrintNode)` -> `eval_print()`", "Prints `Sum:  12`. There are two spaces because the string already has a space and print joins arguments with a space."],
                ["33", "`reclaim;`", "`interpret(ReturnNode)` -> `eval_return()` line `851`", "Raises `ReturnValue(None)` to stop root."],
                ["34", "Return catch", "`eval_function_call()` lines `886-887`", "Catches `ReturnValue` and returns from root."],
                ["35", "Scope cleanup", "`eval_function_call()` lines `891-892`", "Runs `self.exit_scope()`. Root scope is removed. Program ends."],
            ],
        )
    )

    parts.append("## Variable Table During Runtime")
    parts.append(
        table(
            ["Moment", "a", "b", "sum"],
            [
                ["After declarations", "0", "0", "0"],
                ["After `water(a)` input `5`", "5", "0", "0"],
                ["After `water(b)` input `7`", "5", "7", "0"],
                ["After `sum = a + b`", "5", "7", "12"],
                ["After final print", "5", "7", "12"],
            ],
        )
    )

    parts.append("## Final Output")
    parts.append(
        "```text\n"
        "Enter 1st number:\n"
        "Enter 2nd number:\n"
        "Sum:  12\n"
        "```"
    )

    parts.append("## One-Line Pattern To Memorize")
    parts.append(
        "```text\n"
        "interpret(ProgramNode)\n"
        "-> eval_program()\n"
        "-> interpret(FunctionDeclarationNode root)\n"
        "-> eval_function_declaration()\n"
        "-> declare_function() saves root\n"
        "-> back to eval_program()\n"
        "-> create FunctionCallNode('root')\n"
        "-> interpret(FunctionCallNode root)\n"
        "-> eval_function_call()\n"
        "-> eval_block(root body)\n"
        "-> each statement goes back to interpret()\n"
        "-> matching eval_* function runs\n"
        "-> reclaim ends root\n"
        "```"
    )

    parts.append("## Defense Script")
    parts.append(
        "```text\n"
        "Sa interpreter, nagsisimula siya sa interpret gamit ang ProgramNode. Dahil ProgramNode siya, papasok siya sa eval_program. "
        "Sa eval_program, una niyang niloloop ang node.children. Sa sample na ito, ang child lang ay root function declaration. "
        "Pinapasa niya iyon ulit sa interpret, tapos dahil FunctionDeclarationNode siya, pupunta sa eval_function_declaration. "
        "Doon kinukuha ang function name, return type, and parameters, then sine-save ang root sa self.functions gamit declare_function.\n\n"
        "Pag tapos nang ma-save ang root, babalik siya sa eval_program. Tapos gagawa siya ng FunctionCallNode('root'), meaning tatawagin na niya ang root. "
        "Ipapasa ulit iyon sa interpret, then dahil FunctionCallNode na siya, pupunta sa eval_function_call. "
        "Sa eval_function_call, hahanapin niya ang saved root, gagawa ng scope, then tatawagin ang eval_block para patakbuhin isa-isa ang statements sa loob ng root.\n\n"
        "Sa eval_block, bawat statement bumabalik sa interpret. Kaya seed a goes to eval_variable_declaration, plant goes to eval_print, water goes to eval_assignment then eval_input, "
        "sum = a + b goes to eval_assignment then eval_binary_op, and reclaim goes to eval_return para matapos ang root.\n"
        "```"
    )

    return "\n\n".join(parts)


def clean(text):
    return text.replace("`", "")


def set_cell_shading(cell, fill):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    except Exception:
        pass


def build_docx(markdown):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    def add_code_block(text):
        p = doc.add_paragraph()
        for line in text.splitlines():
            run = p.add_run(line + "\n")
            run.font.name = "Consolas"
            run.font.size = Pt(8)

    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block("\n".join(code_lines))
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1

            rows = []
            for idx, table_line in enumerate(table_lines):
                if idx == 1:
                    continue
                rows.append([clean(cell.strip().replace("<br>", "\n")) for cell in table_line.strip("|").split("|")])

            if rows:
                max_cols = max(len(row) for row in rows)
                tbl = doc.add_table(rows=len(rows), cols=max_cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(max_cols):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = row[c_idx] if c_idx < len(row) else ""
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(7.5)
                        if r_idx == 0:
                            set_cell_shading(cell, "D9EAD3")
                doc.add_paragraph()
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = clean(line[level:].strip())
            if level == 1:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                doc.add_heading(text, 1)
            else:
                doc.add_heading(text, 2)
            i += 1
            continue

        if line.startswith("- "):
            doc.add_paragraph(clean(line[2:].strip()), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
            i += 1
            continue

        doc.add_paragraph(clean(line.strip()))
        i += 1

    doc.save(DOCX_PATH)


def export_pdf():
    try:
        import win32com.client

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DOCX_PATH.resolve()))
        doc.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return True, str(PDF_PATH)
    except Exception as exc:
        return False, str(exc)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = build_markdown()
    MD_PATH.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    ok, detail = export_pdf()
    print(f"markdown={MD_PATH}")
    print(f"docx={DOCX_PATH}")
    print(f"pdf={'ok ' + detail if ok else 'failed ' + detail}")
